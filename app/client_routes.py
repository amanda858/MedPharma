"""Client Hub API — auth, claims queue, payments, notes, credentialing, enrollment, EDI, providers, dashboard."""

from __future__ import annotations

import os
import json as _json
import re
import hashlib
import logging
import shutil
import sqlite3
import threading
import uuid
from datetime import datetime, date, timedelta
from typing import Optional
from fastapi import APIRouter, HTTPException, Cookie, Response, Request, UploadFile, File as FastAPIFile, Form, Query
from fastapi.responses import JSONResponse
from pydantic import BaseModel

IS_PROD = bool(os.getenv("PORT"))  # Render sets PORT in production
log = logging.getLogger(__name__)

from app.client_db import (
    get_db,
    authenticate, validate_session, logout_session,
    list_clients, create_client, update_client, delete_client,
    create_user_invite, get_password_setup_token_info, consume_password_setup_token,
    set_must_change_password, change_password_with_current,
    force_set_password,
    get_profile, update_profile,
    get_practice_profiles, upsert_practice_profile, delete_practice_profile,
    list_providers, create_provider, update_provider, delete_provider,
    get_claims, get_claim, create_claim, update_claim, delete_claim,
    get_ar_worklist,
    get_payments, create_payment, delete_payment,
    get_notes, add_note, get_claim_client_ids,
    get_credentialing, create_credentialing, update_credentialing, delete_credentialing,
    get_enrollment, create_enrollment, update_enrollment, delete_enrollment,
    get_eligibility, create_eligibility, update_eligibility, delete_eligibility,
    get_eligibility_one,
    get_edi, create_edi, update_edi, delete_edi,
    get_dashboard, CLAIM_STATUSES,
    list_files, add_file, get_file_record, update_file_record, delete_file_record,
    list_production_logs, add_production_log, delete_production_log, get_production_report,
    log_audit, get_audit_log, auto_flag_sla, get_alerts,
    log_activity, list_activity_events, get_live_users, get_productivity_report,
    global_search, bulk_update_claims, export_claims, export_table,
    get_report_notes, upsert_report_note, delete_report_note, rename_report_note,
    get_user_production_snapshot,
    list_sharefile_links, add_sharefile_link, delete_sharefile_link,
    create_job, append_job_event, set_job_running, update_job_progress,
    complete_job, fail_job, get_job, list_jobs, reset_job_for_retry,
    _load_clients_seed,
    list_rooms_for_user, get_room, create_room, update_room, delete_room,
    list_room_members, add_room_member, remove_room_member,
    user_can_access_room, add_room_message, list_room_messages,
    mark_room_read, chat_unread_total, list_chat_eligible_users,
    get_or_create_dm_room,
    list_room_read_state,
    list_client_access, set_client_access, list_clients_for_user,
    accounts_assigned_to_user,
    create_notification, fanout_notification, list_notifications,
    count_unread_notifications, mark_notification_read,
    mark_all_notifications_read, delete_notification, delete_notifications,
    save_eod_report, list_eod_reports, get_eod_report,
    get_team_activity_rollup,
    set_app_setting, get_app_setting, list_app_settings,
    ALLOWED_SETTING_KEYS,
    list_leads, create_lead, update_lead,
    delete_lead, mark_lead_followed_up, list_leads_due_followup,
    restore_lead, list_deleted_leads, get_leads_pipeline,
)

from app.notifications import (
    notify_activity,
    notify_bulk_activity,
    flush_and_notify,
    send_test_notification,
    get_notification_status,
    get_notification_debug,
    send_daily_account_summary,
)
from rule_intercept import intercept_excel_upload
from app.config import business_today, business_today_iso, business_now

router = APIRouter(prefix="/hub/api")


DATA_IMPORT_CATEGORIES = ("Claims", "Credentialing", "Enrollment", "EDI")


def _send_direct_email(to_email: str, subject: str, text_body: str, html_body: str = "") -> tuple[bool, str]:
    """Send a direct email to a single recipient using SendGrid or SMTP.

    Resolution order for credentials:
      1. In-DB encrypted settings (app_settings table) — set via admin UI.
      2. Environment variables (SENDGRID_API_KEY, SMTP_*).
    This lets the operator paste credentials in the hub without ever
    touching Render env vars.
    """
    to_email = (to_email or "").strip()
    if not to_email:
        return False, "missing recipient"

    # In-DB settings win over env so an admin can override Render config.
    try:
        from app.client_db import get_app_setting as _gs
        db_sg_key = (_gs("SENDGRID_API_KEY") or "").strip()
        db_sg_from = (_gs("SENDGRID_FROM") or "").strip()
        db_smtp_h = (_gs("SMTP_HOST") or "").strip()
        db_smtp_p = (_gs("SMTP_PORT") or "").strip()
        db_smtp_u = (_gs("SMTP_USER") or "").strip()
        db_smtp_pw = (_gs("SMTP_PASS") or "").strip()
    except Exception:
        db_sg_key = db_sg_from = db_smtp_h = db_smtp_p = db_smtp_u = db_smtp_pw = ""

    sg_key = db_sg_key or (os.getenv("SENDGRID_API_KEY") or "").strip()
    sg_from = (db_sg_from
               or (os.getenv("SENDGRID_FROM") or "").strip()
               or db_smtp_u
               or (os.getenv("SMTP_USER") or "").strip()
               or "notifications@medprosc.com")
    smtp_h = db_smtp_h or (os.getenv("SMTP_HOST") or "").strip()
    smtp_p_raw = db_smtp_p or (os.getenv("SMTP_PORT") or "587").strip()
    try:
        smtp_p = int(smtp_p_raw or 587)
    except (TypeError, ValueError):
        smtp_p = 587
    smtp_u = db_smtp_u or (os.getenv("SMTP_USER") or "").strip()
    smtp_pw = db_smtp_pw or (os.getenv("SMTP_PASS") or "").strip()

    # Track the actual SendGrid failure reason so we don't mask it with the
    # generic "no provider configured" message when SMTP also isn't set.
    sendgrid_failure: str | None = None

    if sg_key:
        try:
            import urllib.request
            import urllib.error
            payload = _json.dumps({
                "personalizations": [{"to": [{"email": to_email}]}],
                "from": {"email": sg_from, "name": "MedPharma Hub"},
                "subject": subject,
                "content": [
                    {"type": "text/plain", "value": text_body or "(no content)"},
                    *([{"type": "text/html", "value": html_body}] if html_body else []),
                ],
            }).encode("utf-8")
            req = urllib.request.Request(
                "https://api.sendgrid.com/v3/mail/send",
                data=payload,
                headers={
                    "Authorization": f"Bearer {sg_key}",
                    "Content-Type": "application/json",
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=20) as resp:
                if resp.getcode() in (200, 202):
                    return True, "sendgrid"
                return False, f"sendgrid http {resp.getcode()}"
        except urllib.error.HTTPError as e:
            body = ""
            try:
                body = (e.read() or b"").decode("utf-8", "ignore")[:280]
            except Exception:
                body = ""
            msg = f"sendgrid http {e.code}"
            if body:
                msg = f"{msg}: {body}"
            log.error("invite email sendgrid failed: %s", msg)
            return False, msg
        except Exception as e:
            sendgrid_failure = f"sendgrid error: {e}"
            log.error("invite email sendgrid failed: %s", e)

    if smtp_h and smtp_u and smtp_pw:
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart

            msg = MIMEMultipart("alternative")
            msg["Subject"] = subject
            msg["From"] = smtp_u
            msg["To"] = to_email
            msg.attach(MIMEText(text_body or "(no content)", "plain"))
            if html_body:
                msg.attach(MIMEText(html_body, "html"))

            with smtplib.SMTP(smtp_h, smtp_p, timeout=20) as server:
                server.starttls()
                server.login(smtp_u, smtp_pw)
                server.sendmail(msg["From"], [to_email], msg.as_string())
            return True, "smtp"
        except Exception as e:
            log.error("invite email smtp failed: %s", e)
            return False, f"smtp failed: {e}"

    # No provider succeeded. Surface the real SendGrid failure if we hit one,
    # otherwise tell the operator what env vars are still missing.
    if sendgrid_failure:
        return False, sendgrid_failure
    return False, "email provider not configured (set SENDGRID_API_KEY or SMTP_* env vars)"


def _all_team_user_ids() -> list[int]:
    """All active admin/staff IDs for default client-access seeding."""
    ids: list[int] = []
    for u in list_chat_eligible_users() or []:
        role = (u.get("role") or "").strip().lower()
        if role not in ("admin", "staff"):
            continue
        try:
            uid = int(u.get("id"))
        except (TypeError, ValueError):
            continue
        if uid > 0 and uid not in ids:
            ids.append(uid)
    return ids


def _lookup_users_by_ids(user_ids: list[int]) -> list[dict]:
    """Fetch active user rows (id, username, contact_name, email, role) for
    the given ids. Skips rows without an email so we don't try to send to
    nobody."""
    if not user_ids:
        return []
    from app.client_db import get_db
    placeholders = ",".join("?" * len(user_ids))
    conn = get_db()
    try:
        rows = conn.execute(
            f"SELECT id, username, contact_name, email, role FROM clients "
            f"WHERE id IN ({placeholders}) AND COALESCE(is_active,1)=1",
            tuple(int(i) for i in user_ids),
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


def _email_provider_configured() -> bool:
    """True when a SendGrid key or full SMTP creds are available (DB or env).

    When this is False the hub still works perfectly — members are notified
    in-app via the 💬 badge — so we use this to avoid surfacing scary
    "email failed" errors when email simply isn't set up."""
    try:
        from app.client_db import get_app_setting as _gs
        sg = (_gs("SENDGRID_API_KEY") or "").strip()
        smtp_h = (_gs("SMTP_HOST") or "").strip()
        smtp_u = (_gs("SMTP_USER") or "").strip()
        smtp_pw = (_gs("SMTP_PASS") or "").strip()
    except Exception:
        sg = smtp_h = smtp_u = smtp_pw = ""
    sg = sg or (os.getenv("SENDGRID_API_KEY") or "").strip()
    smtp_h = smtp_h or (os.getenv("SMTP_HOST") or "").strip()
    smtp_u = smtp_u or (os.getenv("SMTP_USER") or "").strip()
    smtp_pw = smtp_pw or (os.getenv("SMTP_PASS") or "").strip()
    if sg:
        return True
    return bool(smtp_h and smtp_u and smtp_pw)


def _send_chat_invite_emails(
    request: Request,
    room_id: int,
    room_name: str,
    user_ids: list[int],
    inviter_name: str,
    skip_user_id: int | None = None,
) -> list[dict]:
    """Email every user that was just added to a chat room. Returns a per-
    user delivery report so the API can surface what actually went out."""
    report: list[dict] = []
    targets = _lookup_users_by_ids(user_ids)
    if not targets:
        return report

    # When no email provider is configured, members are still notified in-app
    # via the 💬 badge — so report success via "in-app notification" instead
    # of failing loudly. Email is strictly optional for chat to work.
    email_on = _email_provider_configured()

    base_url = str(request.base_url).rstrip("/")
    # Deep-link straight into the chat panel for that room.
    setup_link = f"{base_url}/hub?chat={room_id}"
    inviter = (inviter_name or "Your team").strip() or "Your team"
    safe_room = (room_name or "a chat room").strip() or "a chat room"

    for u in targets:
        if skip_user_id and int(u["id"]) == int(skip_user_id):
            # Don't email the person who just created the room about
            # themselves.
            report.append({
                "user_id": u["id"], "username": u["username"],
                "email": u.get("email") or "", "sent": False,
                "via": "skipped (creator)",
            })
            continue
        # No email provider configured → the member is still notified in-app
        # via the 💬 badge. Report it as delivered so the UI stays clean.
        if not email_on:
            report.append({
                "user_id": u["id"], "username": u["username"],
                "email": (u.get("email") or "").strip(), "sent": True,
                "via": "in-app notification",
            })
            continue
        email = (u.get("email") or "").strip()
        if not email or "@" not in email:
            report.append({
                "user_id": u["id"], "username": u["username"],
                "email": "", "sent": True, "via": "in-app notification",
            })
            continue
        display = (u.get("contact_name") or u.get("username") or "there").strip()
        subject = f"💬 You were added to MedPharma chat: {safe_room}"
        text_body = (
            f"Hi {display},\n\n"
            f"{inviter} added you to the chat room \"{safe_room}\" on the "
            f"MedPharma Hub.\n\n"
            f"Open the room directly:\n{setup_link}\n\n"
            f"You'll see new messages with the 💬 badge in the sidebar "
            f"once you sign in.\n\n"
            "If you weren't expecting this, contact your administrator."
        )
        html_body = (
            f"<div style=\"font-family:system-ui,Segoe UI,Arial,sans-serif;"
            f"max-width:540px;margin:0 auto;color:#0f172a\">"
            f"<h2 style=\"margin:0 0 12px;color:#1d4ed8\">💬 You were added "
            f"to a MedPharma chat</h2>"
            f"<p>Hi {display},</p>"
            f"<p><b>{inviter}</b> added you to the chat room "
            f"<b>\"{safe_room}\"</b> on the MedPharma Hub.</p>"
            f"<p style=\"margin:18px 0\">"
            f"<a href=\"{setup_link}\" style=\"display:inline-block;"
            f"padding:10px 22px;background:#1d4ed8;color:#fff;"
            f"text-decoration:none;border-radius:8px;font-weight:600\">"
            f"Open the chat room →</a></p>"
            f"<p style=\"font-size:12px;color:#64748b\">Or copy this link: "
            f"{setup_link}</p>"
            f"<p style=\"font-size:12px;color:#64748b\">You'll see new "
            f"messages with the 💬 badge in the sidebar once you sign in. "
            f"If you weren't expecting this, contact your administrator.</p>"
            f"</div>"
        )
        sent, via = _send_direct_email(email, subject, text_body, html_body)
        report.append({
            "user_id": u["id"], "username": u["username"],
            "email": email, "sent": sent, "via": via,
        })
    return report


def _norm_text(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip().lower())


# Header tokens that unambiguously identify a *claims* spreadsheet, grouped so a
# mis-filed upload (e.g. a "Daily Claims Worklist" saved under "General") is still
# routed into the Claims importer instead of being silently saved as a document.
_CLAIMS_KEY_HEADERS = {
    "claim", "claimkey", "claim key", "claim id", "claimid", "claim number",
    "claim no", "claimno", "claim num", "icn", "tcn", "dcn",
}
# Claim-specific signal columns. Deliberately limited to fields that credentialing
# and enrollment sheets do NOT carry (charges, balances, DOS, CPT, patient,
# claim dates) so broadening the match never misroutes those other data types.
_CLAIMS_SIGNAL_HEADERS = {
    # financials
    "charge", "charge amount", "chargeamount", "charges", "total charge",
    "total charges", "billed", "billed amount", "amount billed",
    "balance", "balance remaining", "ar balance", "amount due", "outstanding",
    "paid", "paid amount", "amount paid", "allowed", "allowed amount",
    "adjustment", "adjustment amount", "write off",
    # clinical / claim identity
    "dos", "date of service", "service date", "cpt", "cpt code", "cptcode",
    "procedure", "procedure code", "hcpcs", "modifier", "modifiers",
    "patient", "patient name", "patient id", "member id",
    # status / denial / claim dates
    "claim status", "denial", "denied", "denial reason", "denial code",
    "denial category", "bill date", "billed date", "date billed",
    "paid date", "denied date", "remit date", "eob date",
}


# Single source of truth for claim column mapping — the "established mapping
# parameters" the importer uses to recognize each spreadsheet's columns and route
# their values into the right claim fields. Promoted to module level so the import
# diagnostic can show EXACTLY what the system recognizes per file (the same glasses
# the importer wears). Keys are normalized headers (see _norm_key); values are the
# claims_master DB column the header feeds. Matching is exact-first, then substring
# (_fuzzy_match_column), so multi-word keys are preferred to avoid false hits.
CLAIMS_COLUMN_MAP = {
    # ── ClaimKey ──
    "claimkey": "ClaimKey", "claim key": "ClaimKey", "claim": "ClaimKey",
    "claim id": "ClaimKey", "claimid": "ClaimKey", "claim number": "ClaimKey",
    "claim no": "ClaimKey", "claimno": "ClaimKey", "claim num": "ClaimKey",
    "claim #": "ClaimKey", "clm": "ClaimKey", "clm id": "ClaimKey",
    "clm no": "ClaimKey", "clm number": "ClaimKey",
    "account": "ClaimKey", "account number": "ClaimKey", "account no": "ClaimKey",
    "acct": "ClaimKey", "acct no": "ClaimKey", "acct number": "ClaimKey",
    "ticket": "ClaimKey", "ticket no": "ClaimKey", "ticket number": "ClaimKey",
    "reference": "ClaimKey", "ref": "ClaimKey", "ref no": "ClaimKey",
    "icn": "ClaimKey", "tcn": "ClaimKey", "dcn": "ClaimKey",
    "control number": "ClaimKey", "claim control number": "ClaimKey",
    "patient control number": "ClaimKey", "patient account": "ClaimKey",
    "patient account number": "ClaimKey", "encounter": "ClaimKey",
    "encounter id": "ClaimKey", "encounter number": "ClaimKey",
    "invoice": "ClaimKey", "invoice number": "ClaimKey", "visit": "ClaimKey",
    "visit number": "ClaimKey", "visit id": "ClaimKey",
    # ── Patient ──
    "patientname": "PatientName", "patient name": "PatientName", "patient": "PatientName",
    "patient full name": "PatientName", "pt name": "PatientName", "ptname": "PatientName",
    "patientid": "PatientID", "patient id": "PatientID", "member id": "PatientID",
    "memberid": "PatientID", "member": "PatientName", "subscriber": "PatientName",
    "subscriber name": "PatientName", "insured name": "PatientName",
    "subscriber id": "PatientID", "member number": "PatientID", "mrn": "PatientID",
    "patient last name": "PatientName", "last name": "PatientName", "name": "PatientName",
    "first name": "PatientName", "patient last": "PatientName", "patient first": "PatientName",
    # ── Provider ──
    "providername": "ProviderName", "provider name": "ProviderName", "provider": "ProviderName",
    "rendering provider": "ProviderName", "rendering": "ProviderName",
    "rendering provider name": "ProviderName", "rendering prov": "ProviderName",
    "servicing provider": "ProviderName", "doctor": "ProviderName", "physician": "ProviderName",
    "doctor name": "ProviderName", "physician name": "ProviderName", "practitioner": "ProviderName",
    "attending": "ProviderName", "attending provider": "ProviderName",
    "billing provider": "ProviderName", "billing prov": "ProviderName",
    "referring provider": "ProviderName", "servicing prov": "ProviderName",
    "npi": "NPI", "provider npi": "NPI", "rendering npi": "NPI",
    # ── Payor / Insurance ──
    "payor": "Payor", "payer": "Payor", "insurance": "Payor",
    "insurance name": "Payor", "insurance company": "Payor", "ins": "Payor",
    "plan": "Payor", "plan name": "Payor", "payer name": "Payor", "payor name": "Payor",
    "carrier": "Payor", "insurance plan": "Payor", "health plan": "Payor",
    "primary insurance": "Payor", "primary payor": "Payor", "primary payer": "Payor",
    "primary payer name": "Payor", "primary insurance name": "Payor",
    "insurance 1": "Payor", "payer 1": "Payor", "ins 1": "Payor", "insurance1": "Payor",
    "ins name": "Payor", "ins company": "Payor", "financial class": "Payor",
    "fc": "Payor", "fin class": "Payor", "payer type": "Payor", "payor type": "Payor",
    "responsible party": "Payor", "guarantor": "Payor",
    # ── DOS / CPT ──
    "dos": "DOS", "date of service": "DOS", "service date": "DOS",
    "svc date": "DOS", "from date": "DOS", "from": "DOS", "date from": "DOS",
    "service from": "DOS", "from dos": "DOS", "dos from": "DOS", "svc dt": "DOS",
    "service dt": "DOS", "date of service from": "DOS", "from service date": "DOS",
    "begin date of service": "DOS", "service start": "DOS", "start date": "DOS",
    "cptcode": "CPTCode", "cpt code": "CPTCode", "cpt": "CPTCode",
    "procedure": "CPTCode", "procedure code": "CPTCode", "proc code": "CPTCode",
    "proc": "CPTCode", "service code": "CPTCode", "hcpcs": "CPTCode",
    "cpt/hcpcs": "CPTCode", "cpt hcpcs": "CPTCode", "procedure/cpt": "CPTCode",
    "proc/cpt": "CPTCode", "cpt4": "CPTCode", "service": "CPTCode",
    "description": "Description", "desc": "Description", "service description": "Description",
    "procedure description": "Description", "proc desc": "Description",
    "modifiers": "Description", "modifier": "Description", "mod": "Description",
    "scrub notes": "DenialReason", "scrub note": "DenialReason",
    "timely filing status": "ClaimStatus", "timely filing": "ClaimStatus",
    # ── Financials ──
    "chargeamount": "ChargeAmount", "charge amount": "ChargeAmount", "charge": "ChargeAmount",
    "billed": "ChargeAmount", "billed amount": "ChargeAmount", "total charge": "ChargeAmount",
    "total charges": "ChargeAmount", "charges": "ChargeAmount", "amount billed": "ChargeAmount",
    "gross charge": "ChargeAmount", "original amount": "ChargeAmount", "fee": "ChargeAmount",
    "charge amt": "ChargeAmount", "chg": "ChargeAmount", "chg amt": "ChargeAmount",
    "bill amt": "ChargeAmount", "billed amt": "ChargeAmount", "billed charges": "ChargeAmount",
    "line charge": "ChargeAmount", "service charge": "ChargeAmount", "claim charge": "ChargeAmount",
    "total billed": "ChargeAmount", "submitted charge": "ChargeAmount", "submitted amount": "ChargeAmount",
    "allowedamount": "AllowedAmount", "allowed amount": "AllowedAmount", "allowed": "AllowedAmount",
    "approved amount": "AllowedAmount", "contracted amount": "AllowedAmount",
    "allowed amt": "AllowedAmount", "contract amount": "AllowedAmount", "expected": "AllowedAmount",
    "expected amount": "AllowedAmount", "expected reimbursement": "AllowedAmount",
    "adjustmentamount": "AdjustmentAmount", "adjustment": "AdjustmentAmount", "adj": "AdjustmentAmount",
    "adjustment amount": "AdjustmentAmount", "adj amount": "AdjustmentAmount", "adj amt": "AdjustmentAmount",
    "write off": "AdjustmentAmount", "writeoff": "AdjustmentAmount", "contractual": "AdjustmentAmount",
    "contractual adjustment": "AdjustmentAmount", "write-off": "AdjustmentAmount",
    "paidamount": "PaidAmount", "paid amount": "PaidAmount", "paid": "PaidAmount",
    "payment": "PaidAmount", "payment amount": "PaidAmount", "payments": "PaidAmount",
    "total paid": "PaidAmount", "total payments": "PaidAmount", "amount paid": "PaidAmount",
    "ins paid": "PaidAmount", "insurance paid": "PaidAmount", "reimbursement": "PaidAmount",
    "paid amt": "PaidAmount", "pmt": "PaidAmount", "pmt amt": "PaidAmount",
    "insurance payment": "PaidAmount", "ins payment": "PaidAmount", "plan paid": "PaidAmount",
    "payer paid": "PaidAmount", "net paid": "PaidAmount",
    "balanceremaining": "BalanceRemaining", "balance": "BalanceRemaining",
    "balance remaining": "BalanceRemaining", "bal": "BalanceRemaining",
    "ar balance": "BalanceRemaining", "outstanding": "BalanceRemaining",
    "amount due": "BalanceRemaining", "total balance": "BalanceRemaining",
    "patient balance": "BalanceRemaining", "ins balance": "BalanceRemaining",
    "remaining": "BalanceRemaining", "net balance": "BalanceRemaining",
    "open balance": "BalanceRemaining", "current balance": "BalanceRemaining",
    "balance due": "BalanceRemaining", "amt due": "BalanceRemaining", "bal due": "BalanceRemaining",
    # ── Status / dates ──
    "claimstatus": "ClaimStatus", "claim status": "ClaimStatus", "status": "ClaimStatus",
    "current status": "ClaimStatus", "ar status": "ClaimStatus", "clm status": "ClaimStatus",
    "claim state": "ClaimStatus", "claim status description": "ClaimStatus",
    "status description": "ClaimStatus", "claim stage": "ClaimStatus",
    "billdate": "BillDate", "bill date": "BillDate", "billed date": "BillDate",
    "date billed": "BillDate", "submission date": "BillDate", "submitted date": "BillDate",
    "date submitted": "BillDate", "date sent": "BillDate", "sent date": "BillDate",
    "claim sent date": "BillDate", "date claim sent": "BillDate", "transmit date": "BillDate",
    "transmission date": "BillDate", "date transmitted": "BillDate", "filed date": "BillDate",
    "date filed": "BillDate", "service line bill date": "BillDate",
    "denieddate": "DeniedDate", "denied date": "DeniedDate", "date denied": "DeniedDate",
    "denial date": "DeniedDate", "date of denial": "DeniedDate",
    "paiddate": "PaidDate", "paid date": "PaidDate", "date paid": "PaidDate",
    "payment date": "PaidDate", "check date": "PaidDate", "remit date": "PaidDate",
    "eob date": "PaidDate", "era date": "PaidDate", "posted date": "PaidDate",
    "denialreason": "DenialReason", "denial reason": "DenialReason",
    "denial": "DenialReason", "reason": "DenialReason", "remark": "DenialReason",
    "remark code": "DenialReason", "carc": "DenialReason", "rarc": "DenialReason",
    "denial code": "DenialReason", "reason code": "DenialReason", "rejection reason": "DenialReason",
    "denial remark": "DenialReason", "adjustment reason": "DenialReason",
    "denialcategory": "DenialCategory", "denial category": "DenialCategory",
    "denial type": "DenialCategory",
    "owner": "Owner", "assigned to": "Owner", "assigned": "Owner", "worked by": "Owner",
    "biller": "Owner", "rep": "Owner", "handled by": "Owner",
    # ── Sub-profile ──
    "sub_profile": "sub_profile", "subprofile": "sub_profile", "sub profile": "sub_profile",
    "profile": "sub_profile", "practice profile": "sub_profile", "practice": "sub_profile",
}


def _claims_structural_match(headers: list[str]) -> dict:
    """Decide whether a spreadsheet's headers describe claims data.

    A file qualifies if it has an explicit claim-id column plus any claim field,
    or several claim fields on its own. Tuned so credentialing/enrollment sheets
    (which carry none of the charge/DOS/CPT/patient signals) never match."""
    norm = {_norm_text(h) for h in headers if str(h or "").strip()}
    has_key = bool(norm & _CLAIMS_KEY_HEADERS)
    signals = sorted(norm & _CLAIMS_SIGNAL_HEADERS)
    is_claims = (has_key and len(signals) >= 1) or (len(signals) >= 3)
    return {"has_claim_key": has_key, "signals": signals, "is_claims": is_claims}


def _is_clearinghouse_ack(headers: list[str]) -> bool:
    """True when a sheet is a clearinghouse acknowledgement / submission report
    rather than an originating charge register.

    Such a file lists claims that were already billed from a real charge register
    and merely echoes their submission state (Received / Accepted / Forwarded …)
    keyed by a clearinghouse control number, so counting its charges again simply
    double-bills the same claims. It is identified by the tell-tale trio a plain
    charge register never carries together: a control-number column, a status
    column, and a processed / acknowledged column."""
    hl = [str(h).strip().lower() for h in headers if str(h or "").strip()]
    has_control = any("control" in h for h in hl)
    has_status = any("status" in h for h in hl)
    has_proc = any(("process" in h) or ("acknowledg" in h) for h in hl)
    return has_control and has_status and has_proc


def _is_batch_transmission_log(headers: list[str]) -> bool:
    """True when a sheet is the SVD DAILY batch-transmission log (DATE / BATCH # /
    NUMBER OF CLAIMS / TOTAL BILLED / CLEARINGHOUSE) rather than a per-claim charge
    register.

    This is Melissa & Susan's COLLECTIVE daily clearinghouse-transmission summary --
    a cumulative recap that values each submitted claim at a flat rate. Those very
    claims are already itemized with their real charges in the per-claim register
    (Susan's "Claim Sent"), so counting its TOTAL BILLED again double-bills the same
    submissions (~$176K instead of the team's true combined ~$91K). It is identified
    by the batch-number + total-billed pairing a per-claim register never carries,
    alongside a claim-count or clearinghouse column (SV's spellings vary)."""
    hl = [str(h).strip().lower() for h in headers if str(h or "").strip()]
    has_batch = any("batch" in h for h in hl)
    has_billed = any("billed" in h for h in hl)
    has_count = any("claim" in h and ("numer" in h or "number" in h or "no" in h or "#" in h or "of" in h) for h in hl)
    has_ch = any(("clearing" in h) or ("claring" in h) for h in hl)
    return has_batch and has_billed and (has_count or has_ch)


def _is_svd_batch_workbook(content: bytes, ext: str) -> bool:
    """True when ANY sheet of an Excel workbook is the SVD DAILY batch-transmission
    log, scanning the whole file rather than only the combined first-sheet headers.

    The SVD DAILY workbook pairs a batch-summary sheet (DATE / BATCH # / NUMBER OF
    CLAIMS / TOTAL BILLED / CLEARINGHOUSE) with a 'LIST BY PT FOR BATCHES' per-claim
    detail sheet that restates the very same submissions at a flat per-claim rate.
    Both sheets describe claims already itemized with their real charges in Susan's
    per-claim register, so importing the workbook double-bills Melissa + Susan
    (~$176K instead of their true combined ~$91K). Because the generic combined
    parser can surface either sheet, a header-only guard is unreliable for the
    multi-sheet copies — so this scans every sheet's first non-empty row and the
    distinctive detail-sheet name to skip the workbook deterministically."""
    if (ext or "").lower() not in (".xlsx", ".xlsm"):
        return False
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception:
        return False
    try:
        for sn in wb.sheetnames:
            if "list by pt for batches" in _norm_text(sn):
                return True
            ws = wb[sn]
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i > 4:
                    break
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells and _is_batch_transmission_log(cells):
                    return True
        return False
    except Exception:
        return False
    finally:
        try:
            wb.close()
        except Exception:
            pass


def _infer_excel_category(content: bytes, ext: str, filename: str = "", description: str = "") -> tuple[Optional[str], dict]:
    """Infer the best import category from Excel headers + filename/description text."""
    scores = {"Claims": 0, "Credentialing": 0, "Enrollment": 0, "EDI": 0}
    keywords = {
        "Claims": [
            "claim", "patient", "dos", "cpt", "charge", "allowed", "paid", "balance", "denial", "remit",
        ],
        "Credentialing": [
            "credential", "recredential", "expiration", "approved", "follow up", "provider enrollment",
        ],
        "Enrollment": [
            "enroll", "effective", "participation", "in network", "in-network", "payer enrollment",
        ],
        "EDI": [
            "edi", "era", "eft", "clearinghouse", "trading partner", "submitter", "receiver", "payer id", "837", "835",
        ],
    }

    headers: list[str] = []
    try:
        rows = _parse_excel_rows(content, ext, combine_sheets=True)
        if rows:
            headers = [str(k or "") for k in rows[0].keys()]
    except Exception:
        headers = []

    blob = " ".join([*headers, filename or "", description or ""]).lower()

    # Rule-intercept (deterministic) gets first shot.
    intercept = intercept_excel_upload(headers=headers, filename=filename, description=description)
    intercepted_category = intercept.get("category")
    if intercepted_category in DATA_IMPORT_CATEGORIES:
        debug = {
            "scores": scores,
            "headers_sample": headers[:20],
            "best_score": None,
            "second_score": None,
            "intercept": intercept,
        }
        return intercepted_category, debug

    # Hardcoded format templates ("the glasses"): some recurring claim exports have
    # no usable headers (headerless / batch-summary / multi-sheet positional), so
    # keyword + structural detection can't see them. If a template matches, the file
    # is unambiguously claims data — route it to Claims so it auto-imports.
    try:
        if _extract_templated_claim_rows(content, ext):
            debug = {
                "scores": scores,
                "headers_sample": headers[:20],
                "best_score": None,
                "second_score": None,
                "intercept": intercept,
                "template": True,
            }
            return "Claims", debug
    except Exception:
        pass

    # Structural detection: if the columns clearly describe claims data, route to
    # Claims even when the keyword score is weak. This stops daily billed
    # spreadsheets (mis-filed under "General") from being saved as inert documents.
    structural = _claims_structural_match(headers)
    if structural["is_claims"]:
        debug = {
            "scores": scores,
            "headers_sample": headers[:20],
            "best_score": None,
            "second_score": None,
            "intercept": intercept,
            "structural": structural,
        }
        return "Claims", debug

    # Heuristic fallback.
    for category, words in keywords.items():
        for word in words:
            if word in blob:
                # Longer keywords get a little more weight to reduce collisions.
                scores[category] += 2 if len(word) >= 7 else 1

    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    best_cat, best_score = ranked[0]
    second_score = ranked[1][1] if len(ranked) > 1 else 0
    inferred = best_cat if best_score >= 2 and (best_score - second_score) >= 1 else None
    debug = {
        "scores": scores,
        "headers_sample": headers[:20],
        "best_score": best_score,
        "second_score": second_score,
        "intercept": intercept,
        "structural": structural,
    }
    return inferred, debug


# ─── Auth helpers ─────────────────────────────────────────────────────────────

def _get_user(hub_session: Optional[str] = None):
    """Return the authenticated user dict, or None if not logged in."""
    if not hub_session:
        return None
    return validate_session(hub_session)


def _require_user(hub_session: Optional[str] = Cookie(None)):
    """Return the authenticated user or raise 401."""
    user = _get_user(hub_session)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user


def _require_admin(hub_session: Optional[str] = Cookie(None)):
    """Return the authenticated admin user or raise 401/403."""
    user = _require_user(hub_session)
    if user["role"] not in ("admin", "staff"):
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


def _require_chat_manager(hub_session: Optional[str] = Cookie(None)):
    """Chat rooms can be created/managed by internal users: admin, staff,
    and business-development (bizdev) — so BizDev can message the team."""
    user = _require_user(hub_session)
    if (user.get("role") or "") not in ("admin", "staff", "bizdev"):
        raise HTTPException(status_code=403, detail="Chat access required")
    return user


def _require_full_admin(hub_session: Optional[str] = Cookie(None)):
    """Return the authenticated full-admin user (role='admin') or raise 403.
    Used for sensitive operations: manage clients, audit log, leads."""
    user = _require_user(hub_session)
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Full admin access required")
    return user


def _is_eric(user: dict) -> bool:
    """True when the authenticated user is Eric. Reporting is the only
    comprehensive (all-client) view that a non-admin may access, and only Eric
    is granted that exception alongside full admins."""
    for field in (user.get("username"), user.get("contact_name")):
        if not field:
            continue
        s = str(field).strip().lower()
        if s == "eric" or "eric" in s.split():
            return True
    return False


def _require_reporting_access(hub_session: Optional[str] = Cookie(None)):
    """Comprehensive reporting is restricted to full admins and Eric only."""
    user = _require_user(hub_session)
    if user.get("role") == "admin" or _is_eric(user):
        return user
    raise HTTPException(status_code=403, detail="Reporting access required")


def _client_account_id(user: dict) -> int:
    """The account whose data a CLIENT login tracks. Most client users are
    assigned to exactly ONE account — the lab whose claims/payments are keyed
    under that account's id — so a dedicated login (e.g. 'Tivany' for SV
    Diagnostics) resolves to that assigned account and sees its data even though
    the user's own row id differs from the account id. Account-owner logins with
    no separate assignment fall back to their own id. Two-or-more assignments
    fall back to own id as well (the caller can pass an explicit client_id)."""
    own = int(user.get("id", 0) or 0)
    try:
        assigned = [int(c) for c in accounts_assigned_to_user(own) if int(c) != own]
    except Exception:
        assigned = []
    if len(assigned) == 1:
        return assigned[0]
    return own


def _client_scope(user: dict) -> Optional[int]:
    """Return client_id filter — None means all (admin/staff sees all data).
    A client login is scoped to the account it tracks (its single assigned
    account, or its own id for account-owner logins)."""
    if user.get("role") in ("admin", "staff"):
        return None
    return _client_account_id(user)


def _owner_identities(user: dict) -> set:
    """Lowercase tokens a claim's free-text ``Owner`` may use to refer to this
    user: their username, email, the local-part of either, and their display
    name plus its first token (e.g. 'Susan Smith' -> {'susan smith', 'susan'}).

    Used to scope the Claims Queue so a biller only sees the claims they
    personally billed/own."""
    idents = set()
    for field in (user.get("username"), user.get("email"), user.get("contact_name")):
        if not field:
            continue
        s = str(field).strip().lower()
        if not s:
            continue
        idents.add(s)
        if "@" in s:
            idents.add(s.split("@", 1)[0])
        parts = s.split()
        if parts:
            idents.add(parts[0])
    return {i for i in idents if i}


def _dashboard_member_scope(user: dict, member: Optional[str] = None) -> Optional[list]:
    """Decide whose uploaded work the dashboard should reflect.

    - A STAFF biller (e.g. susan / melissa / jessica) always sees ONLY the work
      they personally uploaded — their own per-user dashboard.
    - A full admin sees the COMPREHENSIVE totals across everyone, unless they
      explicitly drill into one person via ?member=… .
    - Client / bizdev logins see their account's full totals (no per-user split).

    Returns the list of lowercase uploaded_by identity tokens to scope to, or
    None for the comprehensive (all-members) view."""
    role = (user.get("role") or "").lower()
    if role == "staff":
        return sorted(_owner_identities(user)) or None
    if role == "admin" and member:
        m = str(member).strip().lower()
        if m:
            idents = {m}
            if "@" in m:
                idents.add(m.split("@", 1)[0])
            return sorted(idents)
    return None


def _assert_client_can_view(user: dict, client_id: int) -> None:
    """Reject if the user is not allowed to view this client's data.

    - admin (role='admin'): always allowed.
    - staff: allowed only if the staff user has been explicitly granted
      access to client_id via Manage Clients → Access. Staff with no grant
      for this client are rejected so cross-client data is not visible.
    - client: allowed only for their own account.
    """
    role = (user.get("user_role") or user.get("role") or "").lower()
    if role == "admin":
        return
    if role == "staff":
        granted = set(list_clients_for_user(int(user.get("id", 0) or 0)))
        if int(client_id or 0) in granted:
            return
        raise HTTPException(status_code=403, detail="You don’t have access to that account.")
    # Client login: their own id plus any account explicitly assigned to them.
    allowed = set(_doc_account_ids(user))
    if int(client_id or 0) in allowed:
        return
    raise HTTPException(status_code=403, detail="You can only view your own account.")


def _doc_account_ids(user: dict) -> list[int]:
    """Account ids whose documents/attachments a CLIENT user may see: their own
    account plus every account they've been explicitly assigned to. This makes
    a single admin upload visible to ALL users on the account."""
    own = int(user.get("id", 0) or 0)
    ids = set()
    if own:
        ids.add(own)
    try:
        for cid in accounts_assigned_to_user(own):
            ids.add(int(cid))
    except Exception:
        pass
    return list(ids)


def _doc_scope(user: dict, client_id: Optional[int] = None):
    """Resolve which account(s) of documents a request should read.

    - admin/staff: the explicitly requested ``client_id`` (None => all).
    - client user: their own account PLUS every account assigned to them, so
      everyone on an account shares the same documents and attachments.
    """
    if user.get("role") in ("admin", "staff"):
        return client_id  # None => all
    return _doc_account_ids(user)


def _client_upload_account(user: dict) -> int:
    """Account a CLIENT user's own uploads/links should land in. A user who is
    assigned to exactly one account uploads INTO that shared account so the
    whole team sees it; account-owner logins (no assignment) use their own id."""
    own = int(user.get("id", 0) or 0)
    try:
        assigned = [int(c) for c in accounts_assigned_to_user(own) if int(c) != own]
    except Exception:
        assigned = []
    if len(assigned) == 1:
        return assigned[0]
    return own


def _single_client_account_or(default_id: int) -> int:
    """Resolve the client account a staff/admin import with no account selected
    should land in, so claims are never parked under a personal login id (which
    hides them from every account dashboard). UNIVERSAL — works with any number
    of client accounts. Priority:
      1. the single client account this user is assigned to (client_user_access);
      2. the designated primary client account (SV Diagnostics, else busiest);
      3. the caller's own id, only as a last resort."""
    uid = int(default_id or 0)
    # 1) explicit membership — most precise
    try:
        assigned = [int(c) for c in accounts_assigned_to_user(uid) if int(c) != uid]
        if len(assigned) == 1:
            return assigned[0]
    except Exception:
        pass
    # 2) designated primary / busiest client account — universal default
    try:
        from .client_db import primary_client_account
        acct = primary_client_account()
        if acct:
            return int(acct)
    except Exception as e:
        log.warning("_single_client_account_or: %s", e)
    return uid


# ─── Auth ─────────────────────────────────────────────────────────────────────

class LoginIn(BaseModel):
    username: str
    password: str


class InviteUserIn(BaseModel):
    company: Optional[str] = ""
    contact_name: Optional[str] = ""
    # Email is optional: when an initial_password is set the admin is minting a
    # ready-to-use login, so no setup-link email is required.
    email: Optional[str] = ""
    phone: Optional[str] = ""
    role: Optional[str] = "staff"
    username: Optional[str] = ""
    initial_password: Optional[str] = ""
    # Optional: grant the new user access to these existing account ids at
    # creation time, so an admin can mint a login for "whatever account".
    account_ids: Optional[list[int]] = None
    # When an initial_password is set, control whether the user must change it
    # on first login. Defaults False so admin-set credentials persist as-is
    # (useful for shared/service logins handed off directly).
    require_password_change: Optional[bool] = False


class SetupPasswordIn(BaseModel):
    password: str


class ChangePasswordIn(BaseModel):
    current_password: str
    new_password: str


@router.post("/login")
def login(body: LoginIn, request: Request, response: Response):
    user, token = authenticate(body.username, body.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    response.set_cookie(
        "hub_session", token,
        httponly=True,
        samesite="lax",
        secure=IS_PROD,        # Only require HTTPS in production (Render)
        path="/",              # Explicit root path — available to ALL routes
        max_age=86400 * 30,
    )
    try:
        log_activity(
            user["username"], "login",
            client_id=user.get("id"),
            ip=(request.client.host if request.client else ""),
            user_agent=request.headers.get("user-agent", ""),
            details="hub login",
        )
    except Exception:
        pass
    return {"ok": True, "user": user}


@router.post("/logout")
def logout(request: Request, response: Response, hub_session: Optional[str] = Cookie(None)):
    # Capture user info BEFORE deleting session
    user = _get_user(hub_session) if hub_session else None
    if user:
        try:
            log_activity(
                user["username"], "logout",
                client_id=user.get("id"),
                ip=(request.client.host if request.client else ""),
                user_agent=request.headers.get("user-agent", ""),
                details="hub logout",
            )
        except Exception:
            pass
    # Always delete session + cookie first — this must succeed unconditionally
    if hub_session:
        try:
            logout_session(hub_session)
        except Exception as exc:
            log.error(f"logout_session error (continuing): {exc}")
    response.delete_cookie("hub_session", path="/")
    # Fire progress report in a background thread — non-blocking, non-critical
    if user:
        threading.Thread(
            target=flush_and_notify,
            args=(user["username"],),
            daemon=True,
        ).start()
    return {"ok": True}


@router.get("/me")
def me(hub_session: Optional[str] = Cookie(None)):
    user = _get_user(hub_session)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user


@router.get("/auth/setup-password/{token}")
def check_setup_password_token(token: str):
    info = get_password_setup_token_info(token)
    if not info:
        raise HTTPException(status_code=404, detail="Invalid or expired setup token")
    return {
        "ok": True,
        "username": info.get("username", ""),
        "contact_name": info.get("contact_name", ""),
        "email": info.get("email", ""),
        "company": info.get("company", ""),
        "role": info.get("role", "client"),
        "expires_at": info.get("expires_at", ""),
    }


@router.post("/auth/setup-password/{token}")
def complete_setup_password(token: str, body: SetupPasswordIn):
    pw = (body.password or "").strip()
    if len(pw) < 10:
        raise HTTPException(status_code=400, detail="Password must be at least 10 characters")
    updated = consume_password_setup_token(token, pw)
    if not updated:
        raise HTTPException(status_code=404, detail="Invalid or expired setup token")
    return {"ok": True, "username": updated.get("username", "")}


@router.post("/auth/change-password")
def change_password(body: ChangePasswordIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    current_pw = (body.current_password or "").strip()
    new_pw = (body.new_password or "").strip()
    if len(new_pw) < 10:
        raise HTTPException(status_code=400, detail="Password must be at least 10 characters")
    ok = change_password_with_current(int(user.get("id", 0) or 0), current_pw, new_pw)
    if not ok:
        raise HTTPException(status_code=401, detail="Current password is incorrect")
    return {"ok": True}


# ─── Accounts (for selector screen) ──────────────────────────────────────────

@router.get("/accounts")
def accounts(hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)

    role = (user.get("role") or "").lower()
    uid = int(user.get("id", 0) or 0)

    # Account selector should show client companies only (not internal/admin users),
    # and avoid duplicate cards for the same company.
    if role == "admin":
        # Full admins always see every active client account.
        clients = [
            c for c in list_clients()
            if c.get("role") == "client" and int(c.get("is_active", 0) or 0) == 1
        ]
    elif role == "staff":
        # Staff users see only client accounts they've been explicitly granted
        # access to via the Add/Edit Client picker. Staff with zero grants see
        # an empty selector by design — grants are how MedPharma controls who
        # works which account. (Admins still see every account.)
        granted_ids = set(list_clients_for_user(uid))
        all_active = [
            c for c in list_clients()
            if c.get("role") == "client" and int(c.get("is_active", 0) or 0) == 1
        ]
        clients = [c for c in all_active if int(c.get("id", 0) or 0) in granted_ids]
    else:
        # Client users see the account(s) they're ASSIGNED to — the lab whose
        # claims they actually track — not their own (often empty) login row.
        # A dedicated login like "Tivany" has id 26 but is granted access to
        # account 10 (SV Diagnostics); showing card 26 would load an empty
        # dashboard. Fall back to their own id only when they have no
        # assignment (account-owner logins whose own id holds the data).
        assigned = [int(c) for c in accounts_assigned_to_user(uid) if int(c) != uid]
        target_ids = set(assigned) if assigned else {uid}
        clients = [
            c for c in list_clients()
            if int(c.get("id", 0) or 0) in target_ids
        ]

    deduped: dict[str, dict] = {}
    for c in clients:
        key = str(c.get("company") or "").strip().lower()
        if not key:
            key = f"id:{c.get('id')}"
        prev = deduped.get(key)
        if prev is None or int(c.get("id", 0) or 0) < int(prev.get("id", 0) or 0):
            deduped[key] = c

    return sorted(deduped.values(), key=lambda x: str(x.get("company") or "").lower())


# ─── Clients (admin) ──────────────────────────────────────────────────────────

class ClientIn(BaseModel):
    company: str
    contact_name: Optional[str] = ""
    email: Optional[str] = ""
    phone: Optional[str] = ""
    service_type: Optional[str] = ""   # rcm | payer_contracting | auditing | hybrid
    notes: Optional[str] = ""
    role: Optional[str] = "client"
    # Optional login credentials for the client. If omitted, the server
    # auto-generates a username (slug from company) and a strong random
    # password — both are returned in the response so the admin can hand
    # them to the client. If provided, both are honored as the client's
    # initial sign-in credentials.
    username: Optional[str] = None
    password: Optional[str] = None
    # Optional: pre-grant access to this new client for a list of existing users
    user_ids: Optional[list[int]] = None
    # Optional: enabled module list (defaults to all modules)
    enabled_modules: Optional[list[str]] = None
    # Daily production report controls (defaults: opt-in ON; recipients fall back to email)
    daily_report_optin: Optional[bool] = True
    report_recipients: Optional[list[str]] = None


class ClientUpdate(BaseModel):
    company: Optional[str] = None
    contact_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    role: Optional[str] = None
    is_active: Optional[int] = None
    password: Optional[str] = None


class ProfileUpdate(BaseModel):
    company: Optional[str] = None
    contact_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    tax_id: Optional[str] = None
    group_npi: Optional[str] = None
    individual_npi: Optional[str] = None
    ptan_group: Optional[str] = None
    ptan_individual: Optional[str] = None
    address: Optional[str] = None
    specialty: Optional[str] = None
    notes: Optional[str] = None
    practice_type: Optional[str] = None
    doc_tabs: Optional[list] = None
    report_tabs: Optional[list] = None
    enabled_modules: Optional[list[str]] = None
    module_labels: Optional[dict] = None
    custom_modules: Optional[list] = None
    daily_report_optin: Optional[bool] = None
    report_recipients: Optional[list[str]] = None


class PracticeProfileUpdate(BaseModel):
    practice_type: Optional[str] = None
    specialty: Optional[str] = None
    tax_id: Optional[str] = None
    group_npi: Optional[str] = None
    individual_npi: Optional[str] = None
    ptan_group: Optional[str] = None
    ptan_individual: Optional[str] = None
    address: Optional[str] = None
    contact_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    notes: Optional[str] = None


@router.get("/clients")
def get_clients(hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    return list_clients()


# Core modules a brand-new client gets by default. The advanced payor modules
# (credentialing / enrollment / edi) are intentionally excluded — they must be
# turned on explicitly per client via the intake module picker.
NEW_CLIENT_DEFAULT_MODULES = [
    "dashboard", "profile", "claims", "providers",
    "reporting", "production", "documents", "chat",
]


@router.post("/clients")
def add_client(body: ClientIn, hub_session: Optional[str] = Cookie(None)):
    admin = _require_full_admin(hub_session)
    payload = body.model_dump()
    user_ids = payload.pop("user_ids", None) or []
    user_ids = [int(u) for u in user_ids if str(u).isdigit()]
    # Pop report-related fields so create_client doesn't reject them — we
    # persist these via update_profile right after the insert.
    enabled_modules    = payload.pop("enabled_modules", None)
    daily_report_optin = payload.pop("daily_report_optin", True)
    report_recipients  = payload.pop("report_recipients", None) or []
    # When the intake form doesn't specify modules, default to the CORE set
    # only. The advanced payor modules (credentialing / enrollment / EDI) stay
    # OFF unless the admin explicitly enables them — otherwise every brand-new
    # client would show modules they never asked for.
    if enabled_modules is None:
        enabled_modules = list(NEW_CLIENT_DEFAULT_MODULES)
    # Optional client password — require minimum length when one is provided.
    supplied_pw = (payload.get("password") or "").strip()
    if supplied_pw and len(supplied_pw) < 8:
        raise HTTPException(status_code=400, detail="Client password must be at least 8 characters")
    cid = create_client(payload)
    # Grant the selected staff/admin users access to this newly-created client
    granted = 0
    if not user_ids and admin.get("id"):
        # Safer default for tenant isolation: if no explicit users were picked,
        # grant access only to the creator/admin instead of the full team.
        user_ids = [int(admin.get("id"))]
    if user_ids:
        try:
            granted = set_client_access(cid, user_ids, granted_by=admin.get("username", ""))
        except Exception as e:
            log.warning("client_user_access seed failed for %s: %s", cid, e)
    # Persist module enablement + daily-report opt-in + extra recipients
    try:
        profile_patch: dict = {
            "daily_report_optin": bool(daily_report_optin),
            "report_recipients": report_recipients,
        }
        if enabled_modules is not None:
            profile_patch["enabled_modules"] = enabled_modules
        update_profile(cid, profile_patch)
    except Exception as e:
        log.warning("post-create profile patch failed for %s: %s", cid, e)
    # Auto-trigger: send the client a one-time welcome / preview email so
    # they immediately see what the daily 6:35 PM EST report will look
    # like. Only fires when an email is on file AND opt-in is true.
    welcome_result = None
    primary_email = (payload.get("email") or "").strip().lower()
    if primary_email and daily_report_optin:
        try:
            from app.notifications import send_client_daily_report_demo
            welcome_result = send_client_daily_report_demo(to_email=primary_email)
        except Exception as e:
            log.warning("welcome demo email failed for client %s (%s): %s",
                        cid, primary_email, e)
            welcome_result = {"ok": False, "error": str(e)}
    # Surface the client login credentials exactly once so the admin can
    # hand them off. payload mutation in create_client adds these keys.
    return {
        "id": cid,
        "ok": True,
        "access_granted": granted,
        "welcome_email": welcome_result,
        "login": {
            "username": payload.get("_created_username") or "",
            "password": payload.get("_created_password") or "",
            "url": "/hub",
            "auto_generated": not bool(supplied_pw),
        },
    }


@router.get("/clients/{cid}/access")
def get_client_access(cid: int, hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    return {"client_id": cid, "users": list_client_access(cid)}


class ClientAccessIn(BaseModel):
    user_ids: list[int] = []


def _client_display_name(cid: int) -> str:
    """Human-friendly name for a client account, used in notifications."""
    try:
        for c in list_clients() or []:
            if int(c.get("id", 0) or 0) == int(cid):
                return (c.get("company") or c.get("contact_name")
                        or c.get("username") or f"Client #{cid}").strip()
    except Exception:
        pass
    return f"Client #{cid}"


def _notify_client_access_granted(cid: int, new_user_ids: list[int],
                                  admin: dict, request: Optional[Request]) -> int:
    """Email + in-app notify users who were just granted access to a client.
    Returns the number of emails successfully sent. Never raises."""
    if not new_user_ids:
        return 0
    client_name = _client_display_name(cid)
    access = {int(u.get("id", 0) or 0): u for u in list_client_access(cid)}
    base_url = str(request.base_url).rstrip("/") if request else ""
    hub_link = f"{base_url}/hub" if base_url else "/hub"
    inviter = (admin.get("contact_name") or admin.get("username")
               or "An administrator").strip()
    sent_count = 0
    for uid in new_user_ids:
        u = access.get(int(uid))
        if not u:
            continue
        name = (u.get("contact_name") or u.get("username") or "there").strip()
        try:
            create_notification(
                user_id=int(uid),
                kind="access",
                title="New client access granted",
                body=f"{inviter} gave you access to {client_name}.",
                link="/hub",
                related_type="client",
                related_id=int(cid),
            )
        except Exception:
            log.exception("access in-app notification failed for user %s", uid)
        email = (u.get("email") or "").strip()
        if email and "@" in email:
            subject = f"You've been granted access to {client_name}"
            text_body = (
                f"Hi {name},\n\n"
                f"{inviter} has given you access to the {client_name} account "
                f"in the MedPharma Hub.\n\n"
                f"Log in here: {hub_link}\n\n"
                "If you did not expect this, contact your administrator."
            )
            html_body = (
                f"<p>Hi {name},</p>"
                f"<p>{inviter} has given you access to the <b>{client_name}</b> "
                f"account in the MedPharma Hub.</p>"
                f"<p><a href=\"{hub_link}\" style=\"padding:10px 16px;background:#1d4ed8;"
                f"color:#fff;text-decoration:none;border-radius:6px;\">Open the Hub</a></p>"
                f"<p style=\"font-size:12px;color:#64748b\">Or copy this link: {hub_link}</p>"
            )
            try:
                sent, _via = _send_direct_email(email, subject, text_body, html_body)
                if sent:
                    sent_count += 1
            except Exception:
                log.exception("access email failed for user %s", uid)
    return sent_count


@router.put("/clients/{cid}/access")
def put_client_access(cid: int, body: ClientAccessIn, request: Request,
                      hub_session: Optional[str] = Cookie(None)):
    admin = _require_full_admin(hub_session)
    requested = [int(u) for u in (body.user_ids or []) if str(u).isdigit()]
    # Capture who already had access so we only notify the *newly* added users.
    existing_ids = {int(u.get("id", 0) or 0) for u in list_client_access(cid)}
    count = set_client_access(cid, requested, granted_by=admin.get("username", ""))
    new_ids = [u for u in requested if u not in existing_ids]
    notified = _notify_client_access_granted(cid, new_ids, admin, request)
    return {"ok": True, "count": count, "notified": notified}


@router.get("/admin/users")
def list_admin_users(hub_session: Optional[str] = Cookie(None)):
    """Internal MedPharma team users for admin pickers.

    Includes admin, staff and bizdev (Business Development / Victor) so the
    Team Production user filter can surface Business Development. The
    client-access-grant UI re-filters this list to admin/staff on its own.
    """
    _require_full_admin(hub_session)
    users = [u for u in list_chat_eligible_users()
             if (u.get("role") or "").lower() in ("admin", "staff", "bizdev")]
    return users


@router.post("/admin/users/invite")
def invite_user(body: InviteUserIn, request: Request, hub_session: Optional[str] = Cookie(None)):
    admin = _require_full_admin(hub_session)
    email = (body.email or "").strip().lower()
    initial_password = (body.initial_password or "").strip()
    # Email is optional WHEN a password is set (admin mints a ready-to-use
    # login). Otherwise an email is required so we can send a setup link.
    if email and "@" not in email:
        raise HTTPException(status_code=400, detail="Email address is not valid")
    if not email and not initial_password:
        raise HTTPException(
            status_code=400,
            detail="Enter an email to send a setup link, or set a password so the user can sign in immediately",
        )
    if initial_password and len(initial_password) < 10:
        raise HTTPException(status_code=400, detail="Password must be at least 10 characters")

    payload = body.model_dump()
    payload["email"] = email
    payload.pop("initial_password", None)
    account_ids = [int(a) for a in (payload.pop("account_ids", None) or []) if str(a).isdigit()]
    require_password_change = bool(payload.pop("require_password_change", False))
    if not (payload.get("company") or "").strip():
        payload["company"] = (admin.get("company") or "").strip() or "MedPharma Team"
    # Staff can invite users but cannot create full-admin accounts.
    if admin.get("role") != "admin" and (payload.get("role") or "client") == "admin":
        payload["role"] = "staff"

    try:
        invite = create_user_invite(payload, invited_by=admin.get("username", "admin"), ttl_hours=72)
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=409, detail="Username or email already exists")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not create user invite: {e}")

    new_uid = int(invite.get("client_id", 0) or 0)

    # Grant the new login access to whichever accounts the admin picked. Merge
    # with each account's existing access list so other users are never dropped.
    granted_accounts = 0
    for acct in account_ids:
        try:
            current = [int(u.get("id")) for u in list_client_access(acct) if u.get("id")]
            if new_uid and new_uid not in current:
                current.append(new_uid)
            set_client_access(acct, current, granted_by=admin.get("username", ""))
            granted_accounts += 1
        except Exception as e:
            log.warning("grant access to account %s for new user %s failed: %s", acct, new_uid, e)

    base_url = str(request.base_url).rstrip("/")
    setup_link = f"{base_url}/hub?setup_token={invite['token']}"
    display_name = invite.get("contact_name") or invite.get("username") or "there"
    sent, via = (False, "skipped-no-email")
    if email:
        subject = "MedPharma Hub: Set your password"
        text_body = (
            f"Hi {display_name},\n\n"
            f"Your MedPharma Hub account is ready.\n"
            f"Username: {invite.get('username','')}\n"
            f"Role: {invite.get('role','client')}\n\n"
            f"Set your password using this link (expires in 72 hours):\n{setup_link}\n\n"
            "If you did not expect this email, contact your administrator."
        )
        html_body = (
            f"<p>Hi {display_name},</p>"
            f"<p>Your MedPharma Hub account is ready.</p>"
            f"<p><b>Username:</b> {invite.get('username','')}<br/>"
            f"<b>Role:</b> {invite.get('role','client')}</p>"
            f"<p><a href=\"{setup_link}\" style=\"padding:10px 16px;background:#1d4ed8;color:#fff;text-decoration:none;border-radius:6px;\">Set Password</a></p>"
            f"<p style=\"font-size:12px;color:#64748b\">Or copy this link: {setup_link}</p>"
        )
        sent, via = _send_direct_email(email, subject, text_body, html_body)

    password_set = False
    if initial_password:
        try:
            updated = consume_password_setup_token(invite.get("token", ""), initial_password)
            password_set = bool(updated)
            if password_set:
                set_must_change_password(int(invite.get("client_id", 0) or 0), require_password_change)
        except Exception:
            password_set = False

    # In-app welcome notification so the new user can see they've been
    # invited the moment they log in — independent of email delivery.
    try:
        new_uid = int(invite.get("client_id", 0) or 0)
        if new_uid:
            inviter_display = (admin.get("contact_name")
                               or admin.get("username") or "An administrator").strip()
            create_notification(
                user_id=new_uid,
                kind="welcome",
                title="Welcome to the MedPharma Hub",
                body=f"{inviter_display} invited you to the hub. "
                     f"Username: {invite.get('username','')}. "
                     f"Role: {invite.get('role','client')}.",
                link="/hub",
                related_type="client",
                related_id=new_uid,
            )
    except Exception:
        log.exception("welcome in-app notification failed for user %s",
                      invite.get("client_id"))

    log_audit(
        None,
        admin.get("username", ""),
        "invite_user",
        "client",
        invite.get("client_id"),
        f"Invited user {invite.get('username')} ({email}), email_sent={sent}, via={via}",
    )

    return {
        "ok": True,
        "user_id": invite.get("client_id"),
        "username": invite.get("username"),
        "email": email,
        "email_sent": sent,
        "delivery": via,
        "password_set": password_set,
        "granted_accounts": granted_accounts,
        "setup_link": setup_link,
        "expires_at": invite.get("expires_at"),
    }


@router.put("/clients/{cid}")
def edit_client(cid: int, body: ClientUpdate, hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    update_client(cid, {k: v for k, v in body.model_dump().items() if v is not None})
    return {"ok": True}


class ForcePasswordIn(BaseModel):
    username: str
    new_password: str


@router.post("/admin/users/force-password")
def admin_force_password(
    body: ForcePasswordIn,
    hub_session: Optional[str] = Cookie(None),
):
    """Admin-only: hard reset a user's password without knowing the old one.
    Intended for unlocking accounts whose hashes drifted out of sync with
    expected starter passwords. Logs an audit row."""
    admin = _require_full_admin(hub_session)
    if len(body.new_password) < 6:
        raise HTTPException(status_code=400, detail="new_password must be at least 6 characters")
    result = force_set_password(body.username, body.new_password)
    if not result.get("ok"):
        raise HTTPException(status_code=400, detail=result.get("error", "force-reset failed"))
    log_audit(
        None,
        admin.get("username", ""),
        "force_password_reset",
        "client",
        result.get("user_id"),
        f"Force-reset password for {result.get('username')}",
    )
    return result


@router.get("/admin/diag/users")
def admin_diag_users(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: dump ALL rows in clients (active + inactive) + the
    app_migrations table. Used to diagnose why a seed/ensure step didn't
    create a row (purged, mark-applied-but-row-missing, etc.)."""
    _require_full_admin(hub_session)
    from .client_db import get_db
    import traceback
    conn = get_db()
    try:
        users = []
        try:
            for r in conn.execute(
                "SELECT id, username, role, company, contact_name, email, "
                "COALESCE(is_active,1) AS is_active "
                "FROM clients ORDER BY id"
            ).fetchall():
                users.append({k: r[k] for k in r.keys()})
        except Exception as e:
            return {"error": "users_query", "detail": f"{type(e).__name__}: {e}", "trace": traceback.format_exc()}
        migrations = []
        try:
            for r in conn.execute("SELECT key, applied_at FROM app_migrations").fetchall():
                migrations.append({k: r[k] for k in r.keys()})
        except Exception as e:
            migrations = [{"error": f"{type(e).__name__}: {e}"}]
    finally:
        conn.close()
    return {"users": users, "user_count": len(users), "migrations": migrations}


@router.get("/admin/diag/data-health")
def admin_diag_data_health(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: explain WHY a dashboard total looks frozen/wrong.

    Surfaces three things, per account and overall:
      1. Claims health — how many claims exist, the Outstanding (AR) balance,
         charged / paid totals, and the most recent claim update. If the
         "last claim activity" timestamp is days old while people upload daily,
         their uploads are NOT reaching claims_master.
      2. The smoking gun — spreadsheets uploaded to Documents that were NEVER
         imported as data (Excel/CSV files whose category isn't a data section).
         Each of these is daily work sitting inert: it never moved the numbers.
      3. Recent uploads — the latest files with their category + row counts, so
         we can see what's coming in and whether it imported.
    """
    admin = _require_full_admin(hub_session)
    from .client_db import get_db
    import traceback
    conn = get_db()
    try:
        cur = conn.cursor()
        id_to_name = {}
        try:
            for r in cur.execute(
                "SELECT id, COALESCE(NULLIF(TRIM(company),''), username) AS name "
                "FROM clients"
            ).fetchall():
                id_to_name[r["id"]] = r["name"]
        except Exception:
            id_to_name = {}

        # 1. Per-account claims health.
        accounts = []
        try:
            rows = cur.execute(
                """SELECT client_id,
                          COUNT(*) AS claims,
                          ROUND(COALESCE(SUM(BalanceRemaining),0),2) AS outstanding,
                          ROUND(COALESCE(SUM(ChargeAmount),0),2)     AS charged,
                          ROUND(COALESCE(SUM(PaidAmount),0),2)       AS paid,
                          SUM(CASE WHEN TRIM(COALESCE(BillDate,''))<>'' THEN 1 ELSE 0 END) AS with_bill_date,
                          SUM(CASE WHEN TRIM(COALESCE(BillDate,''))='' THEN 1 ELSE 0 END)  AS missing_bill_date,
                          MAX(updated_at) AS last_update,
                          MAX(created_at) AS last_created
                   FROM claims_master
                   GROUP BY client_id
                   ORDER BY outstanding DESC"""
            ).fetchall()
            for r in rows:
                accounts.append({
                    "account": id_to_name.get(r["client_id"], f"client {r['client_id']}"),
                    "client_id": r["client_id"],
                    "claims": r["claims"],
                    "outstanding": r["outstanding"],
                    "charged": r["charged"],
                    "paid": r["paid"],
                    "with_bill_date": r["with_bill_date"],
                    "missing_bill_date": r["missing_bill_date"],
                    "last_claim_update": r["last_update"],
                    "last_claim_created": r["last_created"],
                })
        except Exception as e:
            accounts = [{"error": f"{type(e).__name__}: {e}"}]

        # 2. Spreadsheets uploaded but never imported (saved as plain documents).
        not_imported = []
        not_imported_count = 0
        not_imported_rows = 0
        try:
            cats = ",".join("?" * len(DATA_IMPORT_CATEGORIES))
            rows = cur.execute(
                f"""SELECT id, client_id, original_name, category, row_count,
                           uploaded_by, created_at
                    FROM client_files
                    WHERE file_type='excel' AND category NOT IN ({cats})
                    ORDER BY created_at DESC""",
                tuple(DATA_IMPORT_CATEGORIES),
            ).fetchall()
            for r in rows:
                not_imported_count += 1
                not_imported_rows += int(r["row_count"] or 0)
                if len(not_imported) < 25:
                    not_imported.append({
                        "account": id_to_name.get(r["client_id"], f"client {r['client_id']}"),
                        "file": r["original_name"],
                        "saved_as_category": r["category"],
                        "rows_in_file": r["row_count"],
                        "uploaded_by": r["uploaded_by"],
                        "uploaded_at": r["created_at"],
                    })
        except Exception as e:
            not_imported = [{"error": f"{type(e).__name__}: {e}"}]

        # Count PDF uploads — these can never auto-import into claims.
        pdf_uploads = 0
        try:
            pdf_uploads = cur.execute(
                "SELECT COUNT(*) FROM client_files WHERE file_type='pdf'"
            ).fetchone()[0]
        except Exception:
            pdf_uploads = None

        # 3. Recent uploads (any type) so we can see what's actually arriving.
        recent_uploads = []
        try:
            rows = cur.execute(
                """SELECT client_id, original_name, file_type, category,
                          row_count, uploaded_by, created_at
                   FROM client_files
                   ORDER BY created_at DESC
                   LIMIT 15"""
            ).fetchall()
            for r in rows:
                imported_flag = (r["file_type"] == "excel"
                                 and r["category"] in DATA_IMPORT_CATEGORIES)
                recent_uploads.append({
                    "account": id_to_name.get(r["client_id"], f"client {r['client_id']}"),
                    "file": r["original_name"],
                    "type": r["file_type"],
                    "category": r["category"],
                    "rows": r["row_count"],
                    "would_import_as_data": imported_flag,
                    "uploaded_by": r["uploaded_by"],
                    "uploaded_at": r["created_at"],
                })
        except Exception as e:
            recent_uploads = [{"error": f"{type(e).__name__}: {e}"}]

        totals = {}
        try:
            t = cur.execute(
                """SELECT COUNT(*) AS claims,
                          ROUND(COALESCE(SUM(BalanceRemaining),0),2) AS outstanding,
                          ROUND(COALESCE(SUM(ChargeAmount),0),2) AS charged,
                          ROUND(COALESCE(SUM(PaidAmount),0),2) AS paid
                   FROM claims_master"""
            ).fetchone()
            totals = {"claims": t["claims"], "outstanding": t["outstanding"],
                      "charged": t["charged"], "paid": t["paid"]}
        except Exception as e:
            totals = {"error": f"{type(e).__name__}: {e}"}

        # 4. MISFILED CLAIMS — the usual cause of a "frozen" account total while
        # the team insists they already imported. A client login that is itself
        # assigned to an account (has a client_user_access row) is a WORKER on
        # that account; their uploads are meant to roll up to it. But an upload
        # is scoped to the account only when the worker is assigned to EXACTLY
        # ONE account — with zero or 2+ assignments the importer falls back to
        # the worker's OWN id, so the claims land under a personal client_id and
        # the account dashboard never reflects them. This finds claims sitting
        # under such worker logins and names the account they should belong to.
        misfiled = []
        try:
            assign_map = {}
            for r in cur.execute(
                """SELECT user_id AS uid, client_id AS acct
                   FROM client_user_access
                   WHERE user_id IN (SELECT DISTINCT client_id FROM claims_master)"""
            ).fetchall():
                assign_map.setdefault(int(r["uid"]), []).append(int(r["acct"]))
            for acc in accounts:
                cid = acc.get("client_id")
                if cid is None or cid not in assign_map:
                    continue
                targets = sorted({t for t in assign_map[cid] if t != cid})
                if targets:
                    misfiled.append({
                        "parked_under_login": acc.get("account"),
                        "parked_client_id": cid,
                        "claims": acc.get("claims"),
                        "charged": acc.get("charged"),
                        "should_roll_up_to": [id_to_name.get(t, f"client {t}") for t in targets],
                        "target_client_ids": targets,
                    })
        except Exception as e:
            misfiled = [{"error": f"{type(e).__name__}: {e}"}]

        # 5. Claims actually CREATED per day (last 12 days), per account — shows
        # exactly where a given day's batch (e.g. 6/28) landed. New claims stamp
        # created_at at insert; if a day's rows show up under a worker login here
        # (not the account) that is the smoking gun. If a day shows up NOWHERE,
        # the file was never imported (see uploaded_but_not_imported above).
        recent_claims = []
        try:
            for r in cur.execute(
                """SELECT client_id, DATE(created_at) AS day, COUNT(*) AS n,
                          ROUND(COALESCE(SUM(ChargeAmount),0),2) AS charged,
                          MAX(uploaded_by) AS by_user
                   FROM claims_master
                   WHERE created_at >= DATE('now','-12 day')
                   GROUP BY client_id, DATE(created_at)
                   ORDER BY day DESC, n DESC"""
            ).fetchall():
                recent_claims.append({
                    "account": id_to_name.get(r["client_id"], f"client {r['client_id']}"),
                    "client_id": r["client_id"],
                    "day": r["day"],
                    "claims_created": r["n"],
                    "charged": r["charged"],
                    "uploaded_by": r["by_user"],
                })
        except Exception as e:
            recent_claims = [{"error": f"{type(e).__name__}: {e}"}]
    except Exception as e:
        conn.close()
        return {"error": "data_health", "detail": f"{type(e).__name__}: {e}",
                "trace": traceback.format_exc()}
    finally:
        try:
            conn.close()
        except Exception:
            pass

    diagnosis = []
    if not_imported_count:
        diagnosis.append(
            f"{not_imported_count} uploaded spreadsheet(s) totalling ~{not_imported_rows} "
            f"row(s) were saved as documents and NEVER imported into claims — this is "
            f"daily work that did not move any number. Re-upload them under the "
            f"\"Claims\" category (or use Claims Queue → Import Excel) to ingest them."
        )
    if pdf_uploads:
        diagnosis.append(
            f"{pdf_uploads} PDF upload(s) exist; PDFs are never auto-imported into "
            f"claims, so any production reported only as PDF is invisible to the totals."
        )
    if misfiled and not (len(misfiled) == 1 and "error" in misfiled[0]):
        _mis_claims = sum(int(m.get("claims") or 0) for m in misfiled)
        _logins = ", ".join(sorted({str(m.get("parked_under_login")) for m in misfiled}))
        _targets = ", ".join(sorted({t for m in misfiled for t in (m.get("should_roll_up_to") or [])}))
        diagnosis.append(
            f"MISFILED CLAIMS: {_mis_claims} claim(s) are parked under worker login(s) "
            f"[{_logins}] instead of the account they belong to [{_targets}]. Their "
            f"imports fell back to the worker's own id because the worker isn't assigned "
            f"to exactly one account. Fix: Manage Clients → Access — assign each of those "
            f"logins to EXACTLY ONE account ([{_targets}]), then re-parent those claims to "
            f"that account so they roll into the dashboard total."
        )
    if not diagnosis:
        diagnosis.append(
            "No un-imported spreadsheets detected. If a total still looks frozen, "
            "compare each account's last_claim_update against when people last "
            "uploaded — a stale timestamp means new uploads aren't reaching that account."
        )

    return {
        "ok": True,
        "viewed_by": admin.get("username"),
        "overall_totals": totals,
        "accounts": accounts,
        "uploaded_but_not_imported": {
            "count": not_imported_count,
            "total_rows": not_imported_rows,
            "files": not_imported,
        },
        "pdf_uploads": pdf_uploads,
        "recent_uploads": recent_uploads,
        "misfiled_claims": misfiled,
        "recent_claims_by_day": recent_claims,
        "diagnosis": diagnosis,
    }


@router.get("/admin/diag/claim-mapping")
def admin_diag_claim_mapping(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: show EXACTLY what the system's mapping 'glasses' recognize in
    every uploaded spreadsheet — the headers it found, which DB column each header
    maps to, which headers it could NOT map, whether the file is recognized as
    claims, and how many rows would actually import. This is the tool to see why a
    file isn't moving the totals: a frozen number almost always means either the
    file's columns aren't mapping (unmapped headers below) or every row collides
    with an existing claim (re-imports of the same claims update in place)."""
    admin = _require_full_admin(hub_session)
    from .client_db import get_db
    conn = get_db()
    try:
        cur = conn.cursor()
        id_to_name = {}
        try:
            for r in cur.execute(
                "SELECT id, COALESCE(NULLIF(TRIM(company),''), username) AS name FROM clients"
            ).fetchall():
                id_to_name[r["id"]] = r["name"]
        except Exception:
            id_to_name = {}
        rows = cur.execute(
            """SELECT id, client_id, filename, original_name, category,
                      row_count, uploaded_by, created_at
               FROM client_files
               WHERE file_type='excel'
               ORDER BY created_at DESC
               LIMIT 40"""
        ).fetchall()
    finally:
        conn.close()

    files = []
    for r in rows:
        info = {
            "file_id": r["id"],
            "account": id_to_name.get(r["client_id"], f"client {r['client_id']}"),
            "file": r["original_name"],
            "saved_as_category": r["category"],
            "rows_in_file_record": r["row_count"],
            "uploaded_by": r["uploaded_by"],
            "uploaded_at": r["created_at"],
        }
        path = os.path.join(UPLOAD_DIR, r["filename"])
        if not os.path.isfile(path):
            info["status"] = "missing_on_disk"
            files.append(info)
            continue
        ext = os.path.splitext(r["original_name"] or "")[1].lower()
        if ext not in (".xlsx", ".xls", ".csv", ".ods", ".odf"):
            ext = os.path.splitext(r["filename"])[1].lower()
        try:
            with open(path, "rb") as fh:
                content = fh.read()
            # Hardcoded template recognition takes precedence — it's how headerless
            # / batch-summary / multi-sheet SV exports are read.
            tpl = _extract_templated_claim_rows(content, ext)
            if tpl:
                tpl_name, tpl_rows = tpl
                from collections import Counter as _Counter
                by_status = dict(_Counter(str(rw.get("Claim Status", "")) for rw in tpl_rows))
                info.update({
                    "status": "ok",
                    "recognized_via": f"template:{tpl_name}",
                    "rows_parsed": len(tpl_rows),
                    "template_status_breakdown": by_status,
                    "recognized_as_claims": True,
                    "has_money_column": True,
                    "has_claim_key_column": True,
                })
                files.append(info)
                continue
            parsed = _parse_excel_rows(content, ext, combine_sheets=True)
            headers = list(parsed[0].keys()) if parsed else []
            mapped = {}
            unmapped = []
            for h in headers:
                col = _fuzzy_match_column(h, CLAIMS_COLUMN_MAP)
                if col:
                    mapped[str(h)] = col
                else:
                    unmapped.append(str(h))
            match = _claims_structural_match(headers)
            mapped_targets = set(mapped.values())
            info.update({
                "status": "ok",
                "recognized_via": "headers" if match["is_claims"] else "none",
                "rows_parsed": len(parsed),
                "headers_detected": [str(h) for h in headers],
                "mapped_columns": mapped,
                "unmapped_headers": unmapped,
                "recognized_as_claims": match["is_claims"],
                "has_money_column": bool(mapped_targets & {"ChargeAmount", "PaidAmount", "BalanceRemaining"}),
                "has_claim_key_column": "ClaimKey" in mapped_targets,
            })
        except Exception as e:
            info["status"] = f"parse_error: {type(e).__name__}: {e}"
        files.append(info)

    return {
        "ok": True,
        "viewed_by": admin.get("username"),
        "how_to_read": (
            "For each file: 'mapped_columns' is what the glasses recognized; "
            "'unmapped_headers' are columns the system ignored. If a money column "
            "(charge/paid/balance) is unmapped, that file's dollars won't compute. "
            "'recognized_as_claims=false' means the file was never treated as claims. "
            "If a file looks fully mapped but still didn't move totals, its rows are "
            "re-imports of existing claims (same claim numbers) and updated in place."
        ),
        "mapping_targets": sorted(set(CLAIMS_COLUMN_MAP.values())),
        "files": files,
    }


@router.post("/admin/diag/reimport-all-claims")
def admin_diag_reimport_all_claims(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: force a fresh import pass over EVERY stored claim spreadsheet
    (any category), not just the ones still filed as documents. Idempotent — the
    importer upserts by (client_id, ClaimKey), so re-running never double-counts.
    Use this to recompute totals after a mapping change without re-uploading."""
    admin = _require_full_admin(hub_session)
    result = reimport_all_claim_files()
    result["viewed_by"] = admin.get("username")
    return result


def reimport_all_claim_files() -> dict:
    """Re-import EVERY stored claim spreadsheet (any category). Idempotent — the
    importer upserts by (client_id, ClaimKey), so re-running never double-counts.
    Shared by the admin diagnostic endpoint AND the daily scheduler so totals
    refresh automatically without anyone clicking a button or logging in."""
    from .client_db import get_db
    conn = get_db()
    try:
        cur = conn.cursor()
        rows = cur.execute(
            """SELECT id, client_id, filename, original_name, category, uploaded_by
               FROM client_files WHERE file_type='excel'"""
        ).fetchall()
    finally:
        conn.close()

    results = []
    total_imported = 0
    for r in rows:
        path = os.path.join(UPLOAD_DIR, r["filename"])
        if not os.path.isfile(path):
            results.append({"file": r["original_name"], "imported": 0, "note": "missing_on_disk"})
            continue
        ext = os.path.splitext(r["original_name"] or "")[1].lower()
        if ext not in (".xlsx", ".xls", ".csv", ".ods", ".odf"):
            ext = os.path.splitext(r["filename"])[1].lower()
        try:
            with open(path, "rb") as fh:
                content = fh.read()
            templated = bool(_extract_templated_claim_rows(content, ext))
            if not templated:
                parsed = _parse_excel_rows(content, ext, combine_sheets=True)
                headers = list(parsed[0].keys()) if parsed else []
                if not _claims_structural_match(headers)["is_claims"]:
                    results.append({"file": r["original_name"], "imported": 0,
                                    "note": "not recognized as claims"})
                    continue
            imported, errors = _import_claims_from_excel(
                content, ext, int(r["client_id"]), uploaded_by=str(r["uploaded_by"] or ""))
            total_imported += int(imported or 0)
            try:
                update_file_record(int(r["id"]), {"category": "Claims", "status": "Imported"},
                                   int(r["client_id"]))
            except Exception:
                pass
            results.append({"file": r["original_name"], "imported": imported,
                            "errors": (errors or [])[:3]})
        except Exception as e:
            results.append({"file": r["original_name"], "imported": 0,
                            "note": f"{type(e).__name__}: {e}"})

    return {
        "ok": True,
        "files_processed": len(results),
        "total_rows_imported": total_imported,
        "results": results,
    }


@router.post("/admin/diag/rebuild-client-claims")
def admin_diag_rebuild_client_claims(client_id: int, hub_session: Optional[str] = Cookie(None)):
    """Admin-only: rebuild a client's claims so every claim is counted exactly ONCE.

    Wipes the client's claims_master + payments, then re-imports the genuine
    per-claim registers PLUS the SVD DAILY batch log (Melissa's distinct daily
    clearinghouse sends, credited to her). Skips only aggregate recaps that
    re-describe billing already counted elsewhere — clearinghouse acknowledgement
    lists and verification worklists — so Billed Out reflects each claim once
    instead of the inflated sum of every representation. Returns before/after
    totals AND a per-user tally (billed grouped by the uploader) so each person's
    number is visible. Idempotent: always reconstructs from the stored source
    files."""
    admin = _require_full_admin(hub_session)
    from .client_db import get_db

    conn = get_db()
    try:
        cur = conn.cursor()
        before = cur.execute(
            "SELECT COUNT(*), COALESCE(SUM(ChargeAmount),0) FROM claims_master WHERE client_id=?",
            (client_id,)).fetchone()
        before_rows, before_billed = int(before[0]), float(before[1] or 0)
        files = cur.execute(
            """SELECT id, filename, original_name, category, uploaded_by
               FROM client_files WHERE client_id=? AND file_type='excel'""",
            (client_id,)).fetchall()
        # Wipe so the re-import reconstructs cleanly — clears stale rows left by
        # older keying passes and any previously-expanded batch phantoms.
        cur.execute("DELETE FROM claims_master WHERE client_id=?", (client_id,))
        cur.execute("DELETE FROM payments WHERE client_id=?", (client_id,))
        conn.commit()
    finally:
        conn.close()

    per_file = []
    for r in files:
        path = os.path.join(UPLOAD_DIR, r["filename"])
        if not os.path.isfile(path):
            per_file.append({"file": r["original_name"], "imported": 0,
                             "skipped": "missing_on_disk"})
            continue
        ext = os.path.splitext(r["original_name"] or "")[1].lower()
        if ext not in (".xlsx", ".xls", ".csv", ".ods", ".odf"):
            ext = os.path.splitext(r["filename"])[1].lower()
        try:
            with open(path, "rb") as fh:
                content = fh.read()
            # Gate: genuine per-claim registers import. A recognized template
            # (svd_denials / lims_payments) qualifies; otherwise require a structural
            # claims match WITH a money column so ack lists, worklists and batch-summary
            # recaps are skipped. The SVD DAILY transmission log is Melissa & Susan's
            # COLLECTIVE clearinghouse recap of the same claims already itemized (with
            # real charges) in Susan's per-claim register, so counting its TOTAL BILLED
            # again double-bills the pair (~$176K instead of their true combined ~$91K).
            tpl = _extract_templated_claim_rows(content, ext)
            ok, reason = False, ""
            if tpl:
                ok, reason = True, f"template:{tpl[0]}"
            else:
                parsed = _parse_excel_rows(content, ext, combine_sheets=True)
                headers = list(parsed[0].keys()) if parsed else []
                match = _claims_structural_match(headers)
                mapped_targets = {_fuzzy_match_column(h, CLAIMS_COLUMN_MAP) for h in headers}
                has_money = bool(mapped_targets & {"ChargeAmount", "PaidAmount", "BalanceRemaining"})
                if _is_clearinghouse_ack(headers):
                    # Submission acknowledgement echoing an already-billed register.
                    ok, reason = False, "clearinghouse acknowledgement (submission recap skipped)"
                elif _is_batch_transmission_log(headers) or _is_svd_batch_workbook(content, ext):
                    # SVD DAILY collective transmission recap -- same claims as the
                    # per-claim register; counting it again double-bills the submissions.
                    ok, reason = False, "batch transmission log (SVD DAILY recap skipped)"
                elif match["is_claims"] and has_money:
                    ok, reason = True, "headers"
                else:
                    reason = "not a per-claim register (recap/worklist skipped)"
            if not ok:
                per_file.append({"file": r["original_name"], "uploaded_by": r["uploaded_by"],
                                 "imported": 0, "skipped": reason})
                continue
            imported, errors = _import_claims_from_excel(
                content, ext, int(client_id), uploaded_by=str(r["uploaded_by"] or ""))
            per_file.append({"file": r["original_name"], "uploaded_by": r["uploaded_by"],
                             "imported": int(imported or 0), "via": reason,
                             "errors": (errors or [])[:2]})
        except Exception as e:
            per_file.append({"file": r["original_name"], "imported": 0,
                             "skipped": f"{type(e).__name__}: {e}"})

    conn = get_db()
    try:
        cur = conn.cursor()
        after = cur.execute(
            "SELECT COUNT(*), COALESCE(SUM(ChargeAmount),0) FROM claims_master WHERE client_id=?",
            (client_id,)).fetchone()
        after_rows, after_billed = int(after[0]), float(after[1] or 0)
        per_user = [
            {"uploaded_by": (row[0] or "(unattributed)"),
             "claims": int(row[1]),
             "billed": round(float(row[2] or 0), 2)}
            for row in cur.execute(
                """SELECT uploaded_by, COUNT(*), COALESCE(SUM(ChargeAmount),0)
                   FROM claims_master WHERE client_id=?
                   GROUP BY uploaded_by ORDER BY SUM(ChargeAmount) DESC""",
                (client_id,)).fetchall()
        ]
    finally:
        conn.close()

    return {
        "ok": True,
        "viewed_by": admin.get("username"),
        "client_id": client_id,
        "before": {"claims": before_rows, "billed": round(before_billed, 2)},
        "after": {"claims": after_rows, "billed": round(after_billed, 2)},
        "removed_inflation": round(before_billed - after_billed, 2),
        "per_user_tally": per_user,
        "files": per_file,
    }


@router.post("/admin/diag/ensure-team")
def admin_diag_ensure_team(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: re-run _ensure_medpharma_team_accounts immediately on the
    live DB so we don't have to wait for a restart to seed missing rows."""
    _require_full_admin(hub_session)
    from .client_db import get_db, _ensure_medpharma_team_accounts
    import traceback
    conn = get_db()
    try:
        cur = conn.cursor()
        before = cur.execute("SELECT COUNT(*) FROM clients").fetchone()[0]
        err = None
        trace = None
        try:
            _ensure_medpharma_team_accounts(cur)
            conn.commit()
        except Exception as e:
            err = f"{type(e).__name__}: {e}"
            trace = traceback.format_exc()
            conn.rollback()
        after = cur.execute("SELECT COUNT(*) FROM clients").fetchone()[0]
        usernames = [
            r[0] for r in cur.execute(
                "SELECT username FROM clients WHERE company='MedPharma SC' ORDER BY username"
            ).fetchall()
        ]
    finally:
        conn.close()
    return {"ok": err is None, "error": err, "trace": trace, "before": before, "after": after, "team_usernames": usernames}


@router.get("/admin/diag/email")
def admin_diag_email(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: report the live email + chat-encryption configuration so
    operators can see at a glance why invites aren't going out (missing key,
    wrong From, etc.) without grepping logs.

    Never returns the API key itself — only whether it's set, its prefix, and
    where the From address resolves from."""
    _require_full_admin(hub_session)
    # In-DB (admin-pasted) values win, env vars fall back.
    db = list_app_settings()
    db_key = (get_app_setting("SENDGRID_API_KEY") or "").strip()
    db_from = (get_app_setting("SENDGRID_FROM") or "").strip()
    db_smtp_h = (get_app_setting("SMTP_HOST") or "").strip()
    db_smtp_u = (get_app_setting("SMTP_USER") or "").strip()
    env_sg_key = (os.getenv("SENDGRID_API_KEY") or "").strip()
    env_sg_from = (os.getenv("SENDGRID_FROM") or "").strip()
    env_smtp_h = (os.getenv("SMTP_HOST") or "").strip()
    env_smtp_u = (os.getenv("SMTP_USER") or "").strip()
    notify_emails = (get_app_setting("NOTIFY_EMAILS")
                     or os.getenv("NOTIFY_EMAILS") or "").strip()
    eod_recipients = (get_app_setting("EOD_REPORT_EMAIL")
                      or os.getenv("EOD_REPORT_EMAIL") or "").strip()
    effective_key = db_key or env_sg_key
    effective_from = (db_from or env_sg_from or db_smtp_u
                      or env_smtp_u or "notifications@medprosc.com")
    effective_smtp_h = db_smtp_h or env_smtp_h
    effective_smtp_u = db_smtp_u or env_smtp_u
    sg_key_prefix = effective_key[:6] + "…" if effective_key else ""
    has_email_provider = bool(effective_key or (effective_smtp_h and effective_smtp_u))
    try:
        from app.security import encryption_status
        chat_enc = encryption_status()
    except Exception as e:
        chat_enc = {"encryption": "unknown", "ready": False, "error": str(e)}
    return {
        "email": {
            "sendgrid_key_set": bool(effective_key),
            "sendgrid_key_source": ("db" if db_key
                                    else ("env" if env_sg_key else "")),
            "sendgrid_key_prefix": sg_key_prefix,
            "sendgrid_from": effective_from,
            "sendgrid_from_source": ("db" if db_from
                                     else ("env" if env_sg_from else "default")),
            "smtp_host_set": bool(effective_smtp_h),
            "smtp_user_set": bool(effective_smtp_u),
            "smtp_source": ("db" if (db_smtp_h or db_smtp_u)
                            else ("env" if (env_smtp_h or env_smtp_u) else "")),
            "notify_emails": notify_emails,
            "eod_report_email": eod_recipients,
            "ready": has_email_provider,
            "guidance": (
                "ok" if has_email_provider else
                "No email provider configured. Use Admin → Email Settings in "
                "the hub to paste your SendGrid API key (or SMTP credentials) "
                "— no Render dashboard required."
            ),
        },
        "db_settings": db,  # which keys are set, masked previews only
        "chat_encryption": chat_enc,
    }


# ── Admin-pasted email credentials (stored Fernet-encrypted in DB) ─────────

class EmailSettingsIn(BaseModel):
    sendgrid_api_key: Optional[str] = None
    sendgrid_from: Optional[str] = None
    smtp_host: Optional[str] = None
    smtp_port: Optional[str] = None
    smtp_user: Optional[str] = None
    smtp_pass: Optional[str] = None
    notify_emails: Optional[str] = None
    eod_report_email: Optional[str] = None


@router.get("/admin/email/settings")
def admin_email_settings_get(hub_session: Optional[str] = Cookie(None)):
    """List which credentials are set (never returns full secrets)."""
    _require_full_admin(hub_session)
    return {"settings": list_app_settings(),
            "allowed_keys": sorted(ALLOWED_SETTING_KEYS)}


@router.put("/admin/email/settings")
def admin_email_settings_put(body: EmailSettingsIn,
                             hub_session: Optional[str] = Cookie(None)):
    """Save / clear admin-pasted email credentials.

    Empty string clears a key. Sending null leaves the key unchanged.
    Stored Fernet-encrypted via app.security.
    """
    admin = _require_full_admin(hub_session)
    mapping = {
        "SENDGRID_API_KEY": body.sendgrid_api_key,
        "SENDGRID_FROM":    body.sendgrid_from,
        "SMTP_HOST":        body.smtp_host,
        "SMTP_PORT":        body.smtp_port,
        "SMTP_USER":        body.smtp_user,
        "SMTP_PASS":        body.smtp_pass,
        "NOTIFY_EMAILS":    body.notify_emails,
        "EOD_REPORT_EMAIL": body.eod_report_email,
    }
    saved: list[str] = []
    cleared: list[str] = []
    for key, val in mapping.items():
        if val is None:
            continue  # untouched
        v = (val or "").strip()
        ok = set_app_setting(key, v, updated_by=admin.get("username", ""))
        if not ok:
            raise HTTPException(400, f"Could not save {key}")
        (cleared if not v else saved).append(key)
    log_audit(None, admin.get("username", ""),
              "email_settings_update", "app_settings", None,
              f"saved={saved}, cleared={cleared}")
    return {"ok": True, "saved": saved, "cleared": cleared,
            "settings": list_app_settings()}


class EmailTestIn(BaseModel):
    to: str
    subject: Optional[str] = None
    body: Optional[str] = None


@router.post("/admin/email/test")
def admin_email_test(body: EmailTestIn,
                     hub_session: Optional[str] = Cookie(None)):
    """Fire a one-shot test email using whatever credentials are configured
    right now (DB > env). Returns the EXACT provider message on failure so
    the admin can debug without grepping logs."""
    admin = _require_full_admin(hub_session)
    to = (body.to or "").strip()
    if not to or "@" not in to:
        raise HTTPException(400, "Valid 'to' email address is required")
    subject = (body.subject or "MedPharma Hub — email test").strip()
    text_body = (body.body or
                 f"This is a test email from the MedPharma Hub.\n"
                 f"Triggered by: {admin.get('username','admin')} at "
                 f"{datetime.now().isoformat(timespec='seconds')}.\n\n"
                 f"If you got this, outbound email is working — chat invites "
                 f"and EOD reports will now reach inboxes.")
    html_body = (
        f"<div style='font-family:system-ui,Arial,sans-serif'>"
        f"<h2 style='color:#1d4ed8'>✅ MedPharma Hub email test</h2>"
        f"<p>{text_body.replace(chr(10), '<br/>')}</p>"
        f"</div>"
    )
    sent, via = _send_direct_email(to, subject, text_body, html_body)
    log_audit(None, admin.get("username", ""),
              "email_test", "email", None,
              f"to={to}, sent={sent}, via={via}")
    return {"ok": sent, "sent": sent, "via": via, "to": to}


@router.get("/admin/reports/eod/preview")
def admin_eod_preview(report_date: Optional[str] = None,
                      hub_session: Optional[str] = Cookie(None)):
    """Admin-only: dry-run the EOD report aggregator without sending email.

    Useful to debug what would land in lexi@/eric@'s inbox before the
    scheduler fires at 6:30 PM EST.
    """
    _require_full_admin(hub_session)
    from .client_db import get_eod_team_report
    return get_eod_team_report(report_date)


@router.post("/admin/reports/eod/send-now")
def admin_eod_send_now(report_date: Optional[str] = None,
                       force: bool = True,
                       demo: bool = False,
                       hub_session: Optional[str] = Cookie(None)):
    """Admin-only: dispatch the end-of-day team report immediately.

    Sends to EOD_REPORT_EMAIL (defaults to lexi@medprosc.com).
    Returns the delivery report so you can see who it went to.

    Set ``demo=true`` to send a fully-populated showcase email with
    fabricated activity so the recipient can see what the real report
    will look like once the team is using the hub (handy on quiet days
    when the actual report would just say "no activity").
    """
    _require_full_admin(hub_session)
    from app.notifications import send_eod_team_report, send_eod_team_report_demo
    if demo:
        return send_eod_team_report_demo()
    return send_eod_team_report(report_date=report_date, force=bool(force))


# ─── Per-client daily production report (sent TO the client) ───────────

@router.get("/admin/reports/client/{cid}/preview")
def admin_client_report_preview(cid: int, report_date: Optional[str] = None,
                                hub_session: Optional[str] = Cookie(None)):
    """Admin-only: dry-run the per-client production report aggregator.
    Returns the structured dict the email layer would render.
    """
    _require_full_admin(hub_session)
    from .client_db import get_client_daily_report
    return get_client_daily_report(cid, report_date)


@router.post("/admin/reports/client/{cid}/send-now")
def admin_client_report_send_now(cid: int,
                                 report_date: Optional[str] = None,
                                 force: bool = True,
                                 demo: bool = False,
                                 to_email: Optional[str] = None,
                                 hub_session: Optional[str] = Cookie(None)):
    """Admin-only: dispatch a per-client production report right now.

    The email goes to the client's primary email + any extras configured
    in their profile under ``report_recipients``. Comes with an Excel
    attachment containing the full row-level detail.

    Set ``demo=true`` to fabricate the data (useful for verifying the
    layout in your own inbox). Set ``to_email=...`` to override the
    recipient while previewing.
    """
    _require_full_admin(hub_session)
    from app.notifications import send_client_daily_report, send_client_daily_report_demo
    if demo:
        return send_client_daily_report_demo(to_email=to_email)
    return send_client_daily_report(client_id=cid, report_date=report_date,
                                    force=bool(force), demo=False)


@router.post("/admin/reports/client/send-all")
def admin_client_report_send_all(report_date: Optional[str] = None,
                                 force: bool = False,
                                 hub_session: Optional[str] = Cookie(None)):
    """Admin-only: fan out per-client production reports to every
    opted-in client. This is what the 6:35 PM EST scheduler triggers
    automatically each evening; exposed here so an admin can re-fire on
    demand without waiting for the cron tick.
    """
    _require_full_admin(hub_session)
    from app.notifications import send_all_client_daily_reports
    return send_all_client_daily_reports(report_date=report_date, force=bool(force))


@router.post("/leads-followups/send-reminders")
def api_send_followup_reminders(hub_session: Optional[str] = Cookie(None)):
    """Fire the BizDev follow-up reminders now (in-app + email). This is what
    the 9 AM EST scheduler runs automatically every day; exposed so BizDev or
    an admin can trigger it on demand."""
    _require_leads_access(hub_session)
    from app.notifications import send_bizdev_followup_reminders
    return send_bizdev_followup_reminders()



@router.delete("/clients/{cid}")
def remove_client(cid: int, hub_session: Optional[str] = Cookie(None)):
    admin = _require_full_admin(hub_session)
    # Don't let an admin nuke their own account out from under their session.
    if int(admin.get("id") or 0) == int(cid):
        raise HTTPException(status_code=400, detail="You cannot remove the account you are signed in as.")
    try:
        delete_client(cid)
    except Exception as exc:
        import logging
        logging.getLogger(__name__).exception("delete_client failed for cid=%s", cid)
        raise HTTPException(status_code=500, detail=f"Failed to remove account: {exc}")
    return {"ok": True}


# ─── Business-development leads (BizDev / Victor) ─────────────────────────────

class LeadIn(BaseModel):
    practice_name: Optional[str] = ""
    contact_name: Optional[str] = ""
    contact_email: Optional[str] = ""
    contact_phone: Optional[str] = ""
    service_rcm: Optional[bool] = False
    service_payor: Optional[bool] = False
    service_workflow: Optional[bool] = False
    service_compliance: Optional[bool] = False
    status: Optional[str] = "New"
    est_value: Optional[float] = 0
    owner: Optional[str] = ""
    notes: Optional[str] = ""


class LeadUpdateIn(BaseModel):
    practice_name: Optional[str] = None
    contact_name: Optional[str] = None
    contact_email: Optional[str] = None
    contact_phone: Optional[str] = None
    service_rcm: Optional[bool] = None
    service_payor: Optional[bool] = None
    service_workflow: Optional[bool] = None
    service_compliance: Optional[bool] = None
    status: Optional[str] = None
    est_value: Optional[float] = None
    owner: Optional[str] = None
    notes: Optional[str] = None


def _require_leads_access(hub_session: Optional[str]):
    """Business Development (leads pipeline) is restricted to full admins and
    Eric only — no other user (including the bizdev/staff roles) may view or
    work the leads pipeline."""
    user = _require_user(hub_session)
    if user.get("role") == "admin" or _is_eric(user):
        return user
    raise HTTPException(status_code=403, detail="Business Development access required")


@router.get("/leads")
def api_list_leads(category: Optional[str] = "all",
                   hub_session: Optional[str] = Cookie(None)):
    _require_leads_access(hub_session)
    return {"category": (category or "all"), "leads": list_leads(category)}


@router.post("/leads")
def api_create_lead(body: LeadIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_leads_access(hub_session)
    data = body.model_dump()
    if not data.get("owner"):
        data["owner"] = user.get("contact_name") or user.get("username") or ""
    lead_id = create_lead(data)
    return {"ok": True, "id": lead_id}


@router.put("/leads/{lead_id}")
def api_update_lead(lead_id: int, body: LeadUpdateIn,
                    hub_session: Optional[str] = Cookie(None)):
    _require_leads_access(hub_session)
    changes = {k: v for k, v in body.model_dump().items() if v is not None}
    ok = update_lead(lead_id, changes)
    if not ok:
        raise HTTPException(status_code=400, detail="No valid fields to update")
    return {"ok": True}


@router.delete("/leads/{lead_id}")
def api_delete_lead(lead_id: int, hub_session: Optional[str] = Cookie(None)):
    _require_leads_access(hub_session)
    if not delete_lead(lead_id):
        raise HTTPException(status_code=404, detail="Lead not found")
    return {"ok": True}


@router.get("/leads-deleted")
def api_list_deleted_leads(hub_session: Optional[str] = Cookie(None)):
    """Archived (soft-deleted) leads, so a stray delete can be recovered."""
    _require_leads_access(hub_session)
    return {"leads": list_deleted_leads()}


@router.post("/leads/{lead_id}/restore")
def api_restore_lead(lead_id: int, hub_session: Optional[str] = Cookie(None)):
    """Bring a soft-deleted lead back into the active pipeline."""
    _require_leads_access(hub_session)
    if not restore_lead(lead_id):
        raise HTTPException(status_code=404, detail="Deleted lead not found")
    return {"ok": True}


@router.get("/leads-pipeline")
def api_leads_pipeline(hub_session: Optional[str] = Cookie(None)):
    """Weighted sales-pipeline forecast for the Business Development view."""
    _require_leads_access(hub_session)
    return get_leads_pipeline()


@router.get("/leads-followups-due")
def api_leads_followups_due(hub_session: Optional[str] = Cookie(None)):
    """Open leads that haven't been contacted in 2+ days — the BizDev
    follow-up queue shown in the Leads view."""
    _require_leads_access(hub_session)
    due = list_leads_due_followup()
    return {"count": len(due), "leads": due}


@router.post("/leads/{lead_id}/follow-up")
def api_mark_lead_followup(lead_id: int, hub_session: Optional[str] = Cookie(None)):
    """Log a follow-up for a lead — resets its 2-day reminder clock."""
    _require_leads_access(hub_session)
    if not mark_lead_followed_up(lead_id):
        raise HTTPException(status_code=404, detail="Lead not found")
    return {"ok": True}


# ─── Profile (own client profile) ──────────────────────────────────────────────────────────

@router.get("/profile")
def get_my_profile(hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    return get_profile(cid)


@router.get("/profile/{cid}")
def get_client_profile(cid: int, hub_session: Optional[str] = Cookie(None)):
    # Staff (not just full admins) need to READ a client's profile so the hub
    # can hide the modules the admin disabled. Writes stay admin-only (PUT).
    user = _require_admin(hub_session)
    _assert_client_can_view(user, cid)
    return get_profile(cid)


@router.put("/profile")
def update_my_profile(body: ProfileUpdate, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    data = {k: v for k, v in body.model_dump().items() if v is not None and k not in ("doc_tabs", "report_tabs", "enabled_modules")}
    if body.doc_tabs is not None:
        data["doc_tab_names"] = _json.dumps(body.doc_tabs)
    if body.report_tabs is not None:
        data["report_tab_names"] = _json.dumps(body.report_tabs)
    if body.enabled_modules is not None:
        data["enabled_modules"] = _json.dumps(body.enabled_modules)
    update_profile(cid, data)
    _notify_profile_change(cid, user, body, scope_label="self-service")
    return {"ok": True}


@router.put("/profile/{cid}")
def update_client_profile(cid: int, body: ProfileUpdate, hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    data = {k: v for k, v in body.model_dump().items() if v is not None and k not in ("doc_tabs", "report_tabs", "enabled_modules")}
    if body.doc_tabs is not None:
        data["doc_tab_names"] = _json.dumps(body.doc_tabs)
    if body.report_tabs is not None:
        data["report_tab_names"] = _json.dumps(body.report_tabs)
    if body.enabled_modules is not None:
        data["enabled_modules"] = _json.dumps(body.enabled_modules)
    update_profile(cid, data)
    admin = _get_user(hub_session) or {}
    _notify_profile_change(cid, admin, body, scope_label="admin")
    return {"ok": True}


def _notify_profile_change(client_id: int, actor: dict, body: ProfileUpdate, scope_label: str = ""):
    """Send an audit + admin notification whenever a profile or its module
    opt-outs change. Module changes are highlighted because they affect what
    the client sees in the hub."""
    try:
        changed_fields = [k for k, v in body.model_dump().items() if v is not None]
        details_parts = []
        if body.enabled_modules is not None:
            mods = ", ".join(sorted(body.enabled_modules)) or "(none)"
            details_parts.append(f"enabled_modules=[{mods}]")
        if changed_fields:
            details_parts.append("fields=" + ",".join(sorted(changed_fields)))
        details = " | ".join(details_parts) or "no changes"
        actor_name = actor.get("username", "?") if actor else "?"
        log_audit(client_id, actor_name, "profile_update",
                  "clients", client_id, details)
        notify_activity(actor_name, "updated", "Client Profile",
                        f"client #{client_id} ({scope_label}) — {details}")
        # Note: module opt-in/out changes are recorded in the audit trail and
        # activity feed only — no email is sent (toggling modules is routine).
    except Exception:
        log.exception("profile change notification failed")


# ─── Report Notes (custom report tab content) ──────────────────────────────────

class ReportNoteBody(BaseModel):
    tab_name: str
    content: str = ""

class ReportNoteRenameBody(BaseModel):
    old_name: str
    new_name: str

@router.get("/report-notes/{cid}")
def get_client_report_notes(cid: int, tab_name: Optional[str] = None,
                             hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    _assert_client_can_view(user, cid)
    notes = get_report_notes(cid, tab_name)
    return {"notes": notes}

@router.put("/report-notes/{cid}")
def save_report_note(cid: int, body: ReportNoteBody,
                     hub_session: Optional[str] = Cookie(None)):
    user = _require_full_admin(hub_session)
    upsert_report_note(cid, body.tab_name, body.content, user.get("username", ""))
    return {"ok": True}

@router.delete("/report-notes/{cid}/{tab_name}")
def remove_report_note(cid: int, tab_name: str,
                       hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    delete_report_note(cid, tab_name)
    return {"ok": True}

@router.put("/report-notes/{cid}/rename")
def rename_report_note_endpoint(cid: int, body: ReportNoteRenameBody,
                                hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    rename_report_note(cid, body.old_name, body.new_name)
    return {"ok": True}


# ─── Practice Sub-Profiles ─────────────────────────────────────────────────────────────

@router.get("/practice-profiles")
def list_practice_profiles(hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    return {"profiles": get_practice_profiles(cid)}


@router.get("/practice-profiles/{cid}")
def list_practice_profiles_admin(cid: int, hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    return {"profiles": get_practice_profiles(cid)}


@router.put("/practice-profiles/{profile_name}")
def save_practice_profile(profile_name: str, body: PracticeProfileUpdate,
                          hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    upsert_practice_profile(cid, profile_name, body.model_dump(exclude_none=True))
    return {"ok": True}


@router.put("/practice-profiles/{cid}/{profile_name}")
def save_practice_profile_admin(cid: int, profile_name: str, body: PracticeProfileUpdate,
                                hub_session: Optional[str] = Cookie(None)):
    _require_full_admin(hub_session)
    upsert_practice_profile(cid, profile_name, body.model_dump(exclude_none=True))
    return {"ok": True}


@router.delete("/practice-profiles/{pp_id}")
def remove_practice_profile(pp_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    delete_practice_profile(pp_id, cid)
    return {"ok": True}


# ─── Providers ────────────────────────────────────────────────────────────────

class ProviderIn(BaseModel):
    client_id: int
    ProviderName: str
    NPI: Optional[str] = ""
    Specialty: Optional[str] = ""
    TaxID: Optional[str] = ""
    Email: Optional[str] = ""
    Phone: Optional[str] = ""
    Status: Optional[str] = "Active"
    StartDate: Optional[str] = ""
    Notes: Optional[str] = ""
    sub_profile: Optional[str] = ""


class ProviderUpdate(BaseModel):
    ProviderName: Optional[str] = None
    NPI: Optional[str] = None
    Specialty: Optional[str] = None
    TaxID: Optional[str] = None
    Email: Optional[str] = None
    Phone: Optional[str] = None
    Status: Optional[str] = None
    StartDate: Optional[str] = None
    Notes: Optional[str] = None
    sub_profile: Optional[str] = None


@router.get("/providers")
def get_providers(client_id: Optional[int] = None, sub_profile: Optional[str] = None,
                 hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    return list_providers(scope, sub_profile=sub_profile)


@router.post("/providers")
def add_provider(body: ProviderIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    if user["role"] not in ("admin", "staff"):
        data["client_id"] = user["id"]
    pid = create_provider(data)
    return {"id": pid, "ok": True}


@router.put("/providers/{pid}")
def edit_provider(pid: int, body: ProviderUpdate, hub_session: Optional[str] = Cookie(None)):
    _require_user(hub_session)
    update_provider(pid, {k: v for k, v in body.model_dump().items() if v is not None})
    return {"ok": True}


@router.delete("/providers/{pid}")
def remove_provider(pid: int, hub_session: Optional[str] = Cookie(None)):
    _require_user(hub_session)
    delete_provider(pid)
    return {"ok": True}


# ─── Claims ───────────────────────────────────────────────────────────────────

class ClaimIn(BaseModel):
    client_id: int
    ClaimKey: str
    PatientID: Optional[str] = ""
    PatientName: Optional[str] = ""
    Payor: Optional[str] = ""
    ProviderName: Optional[str] = ""
    NPI: Optional[str] = ""
    DOS: Optional[str] = ""
    CPTCode: Optional[str] = ""
    Description: Optional[str] = ""
    ChargeAmount: Optional[float] = 0
    AllowedAmount: Optional[float] = 0
    AdjustmentAmount: Optional[float] = 0
    PaidAmount: Optional[float] = 0
    BalanceRemaining: Optional[float] = 0
    ClaimStatus: Optional[str] = "Intake"
    BillDate: Optional[str] = ""
    DeniedDate: Optional[str] = ""
    PaidDate: Optional[str] = ""
    Owner: Optional[str] = ""
    NextAction: Optional[str] = ""
    NextActionDueDate: Optional[str] = ""
    SLABreached: Optional[int] = 0
    DenialCategory: Optional[str] = ""
    DenialReason: Optional[str] = ""
    AppealDate: Optional[str] = ""
    AppealStatus: Optional[str] = ""
    sub_profile: Optional[str] = ""


class ClaimUpdate(BaseModel):
    PatientID: Optional[str] = None
    PatientName: Optional[str] = None
    Payor: Optional[str] = None
    ProviderName: Optional[str] = None
    NPI: Optional[str] = None
    DOS: Optional[str] = None
    CPTCode: Optional[str] = None
    Description: Optional[str] = None
    ChargeAmount: Optional[float] = None
    AllowedAmount: Optional[float] = None
    AdjustmentAmount: Optional[float] = None
    PaidAmount: Optional[float] = None
    BalanceRemaining: Optional[float] = None
    ClaimStatus: Optional[str] = None
    BillDate: Optional[str] = None
    DeniedDate: Optional[str] = None
    PaidDate: Optional[str] = None
    Owner: Optional[str] = None
    NextAction: Optional[str] = None
    NextActionDueDate: Optional[str] = None
    SLABreached: Optional[int] = None
    DenialCategory: Optional[str] = None
    DenialReason: Optional[str] = None
    AppealDate: Optional[str] = None
    AppealStatus: Optional[str] = None
    sub_profile: Optional[str] = None


@router.get("/claims")
def get_claims_list(status: Optional[str] = None, client_id: Optional[int] = None,
                   sub_profile: Optional[str] = None,
                   hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    role = (user.get("role") or "").lower()

    if role in ("admin", "staff"):
        # Internal users may scope to a specific account (None => all accounts).
        claims = get_claims(client_id, status, sub_profile=sub_profile)
        if role == "staff":
            # Billers see the claims they personally own/billed, plus any
            # unassigned claims (no Owner yet) so they can pick up new work.
            # Full admins keep the cross-account view used by the admin report.
            idents = _owner_identities(user)
            claims = [c for c in claims
                      if (c.get("Owner") or "").strip().lower() in idents
                      or not (c.get("Owner") or "").strip()]
    else:
        # Account (client) login: locked to the account(s) they belong to so a
        # forged client_id can't expose another lab's claims. They see every
        # claim on their own account regardless of which biller owns it.
        allowed = set(_doc_account_ids(user))
        try:
            requested = int(client_id) if client_id is not None else None
        except (TypeError, ValueError):
            requested = None
        scope = requested if (requested in allowed) else _client_account_id(user)
        claims = get_claims(scope, status, sub_profile=sub_profile)
    return {"claims": claims}


@router.get("/claims/statuses")
def claim_statuses():
    return CLAIM_STATUSES


@router.get("/claims/ar-worklist")
def claims_ar_worklist(client_id: Optional[int] = None, owner: Optional[str] = None,
                       bucket: Optional[str] = None, sub_profile: Optional[str] = None,
                       limit: int = 300, hub_session: Optional[str] = Cookie(None)):
    """Prioritized Accounts-Receivable worklist — highest-recovery open claims
    first, with aging-bucket rollups."""
    user = _require_user(hub_session)
    scope = client_id if client_id is not None else _client_scope(user)
    # Non-staff/admin users are always scoped to the account they track, so a
    # forged client_id can't expose another lab's worklist.
    if user["role"] not in ("admin", "staff"):
        scope = _client_account_id(user)
    return get_ar_worklist(scope, owner=owner, bucket=bucket,
                           sub_profile=sub_profile, limit=limit)


@router.get("/claims/{claim_id}")
def get_single_claim(claim_id: int, hub_session: Optional[str] = Cookie(None)):
    _require_user(hub_session)
    c = get_claim(claim_id)
    if not c:
        raise HTTPException(404, "Claim not found")
    return c


@router.post("/claims")
def add_claim(body: ClaimIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    if user["role"] not in ("admin", "staff"):
        data["client_id"] = user["id"]
    cid = create_claim(data)
    notify_activity(user["username"], "created", "Claims",
                    f"Patient: {data.get('PatientName','')}, Payor: {data.get('Payor','')}")
    return {"id": cid, "ok": True}


@router.put("/claims/{claim_id}")
def edit_claim(claim_id: int, body: ClaimUpdate, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    changes = {k: v for k, v in body.model_dump().items() if v is not None}
    update_claim(claim_id, changes)
    notify_activity(user["username"], "updated", "Claims",
                    f"Claim #{claim_id}, fields: {', '.join(changes.keys())}")
    return {"ok": True}


@router.delete("/claims/{claim_id}")
def remove_claim(claim_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    delete_claim(claim_id)
    notify_activity(user["username"], "deleted", "Claims", f"Claim #{claim_id}")
    return {"ok": True}


# ─── Payments ─────────────────────────────────────────────────────────────────

class PaymentIn(BaseModel):
    ClaimKey: str
    PostDate: Optional[str] = ""
    PaymentAmount: Optional[float] = 0
    AdjustmentAmount: Optional[float] = 0
    PayerType: Optional[str] = "Primary"
    CheckNumber: Optional[str] = ""
    ERA: Optional[str] = ""
    Notes: Optional[str] = ""


@router.get("/claims/{claim_key}/payments")
def list_payments(claim_key: str, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    return get_payments(cid, claim_key)


@router.post("/claims/{claim_key}/payments")
def add_payment(claim_key: str, body: PaymentIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    cid = scope if scope is not None else user["id"]
    data = body.model_dump()
    data["client_id"] = cid
    data["ClaimKey"] = claim_key
    # Attribute the posting to the logged-in user so payment posting shows up in
    # the production report / EOD tally (this is part of how the team is paid).
    data["PostedBy"] = user.get("username") or ""
    pid = create_payment(data)
    notify_activity(user["username"], "posted payment", "Payments", f"{claim_key}")
    return {"id": pid, "ok": True}


@router.delete("/payments/{payment_id}")
def remove_payment(payment_id: int, hub_session: Optional[str] = Cookie(None)):
    _require_user(hub_session)
    delete_payment(payment_id)
    return {"ok": True}


# ─── Notes ────────────────────────────────────────────────────────────────────

class NoteIn(BaseModel):
    ClaimKey: Optional[str] = ""
    Module: Optional[str] = "Claim"
    RefID: Optional[int] = 0
    Note: str
    Author: Optional[str] = ""


def _resolve_note_client_id(user: dict, claim_key: Optional[str]) -> int:
    """Pick the client_id a claim's notes live under. Notes are keyed by
    client_id, but admins/staff browse claims across every account and a
    biller's claims live under the lab's account (not the biller's own login),
    so scoping notes to the viewer's own id hides them. When a claim_key is
    given, resolve the claim's real owning account and authorize it:
      • admin/staff may read any claim's notes,
      • an account login only its own account(s).
    Falls back to the legacy scope when the claim can't be resolved."""
    role = (user.get("role") or "").lower()
    if claim_key:
        owners = get_claim_client_ids(claim_key)
        if owners:
            if role in ("admin", "staff"):
                return owners[0]
            allowed = set(_doc_account_ids(user)) | {user["id"]}
            for oc in owners:
                if oc in allowed:
                    return oc
    scope = _client_scope(user)
    return scope if scope is not None else user["id"]


@router.get("/notes")
def list_notes(claim_key: Optional[str] = None, module: Optional[str] = None,
               ref_id: Optional[int] = None, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    cid = _resolve_note_client_id(user, claim_key)
    return get_notes(cid, claim_key, module, ref_id)


@router.post("/notes")
def post_note(body: NoteIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    cid = _resolve_note_client_id(user, data.get("ClaimKey"))
    data["client_id"] = cid
    if not data.get("Author"):
        data["Author"] = user.get("username", "")
    nid = add_note(data)
    return {"id": nid, "ok": True}


# ─── Credentialing ────────────────────────────────────────────────────────────

class CredIn(BaseModel):
    client_id: int
    provider_id: Optional[int] = None
    ProviderName: Optional[str] = ""
    Payor: Optional[str] = ""
    CredType: Optional[str] = "Initial"
    Status: Optional[str] = "Not Started"
    SubmittedDate: Optional[str] = ""
    FollowUpDate: Optional[str] = ""
    ApprovedDate: Optional[str] = ""
    ExpirationDate: Optional[str] = ""
    Owner: Optional[str] = ""
    Notes: Optional[str] = ""
    sub_profile: Optional[str] = ""


class CredUpdate(BaseModel):
    ProviderName: Optional[str] = None
    Payor: Optional[str] = None
    CredType: Optional[str] = None
    Status: Optional[str] = None
    SubmittedDate: Optional[str] = None
    FollowUpDate: Optional[str] = None
    ApprovedDate: Optional[str] = None
    ExpirationDate: Optional[str] = None
    Owner: Optional[str] = None
    Notes: Optional[str] = None
    sub_profile: Optional[str] = None


@router.get("/credentialing")
def list_cred(status: Optional[str] = None, client_id: Optional[int] = None,
             sub_profile: Optional[str] = None,
             hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    return get_credentialing(scope, status, sub_profile=sub_profile)


@router.post("/credentialing")
def add_cred(body: CredIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    if user["role"] not in ("admin", "staff"):
        data["client_id"] = user["id"]
    rid = create_credentialing(data)
    notify_activity(user["username"], "created", "Credentialing",
                    f"Provider: {data.get('ProviderName','')}, Payor: {data.get('Payor','')}")
    return {"id": rid, "ok": True}


@router.put("/credentialing/{rid}")
def edit_cred(rid: int, body: CredUpdate, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    changes = {k: v for k, v in body.model_dump().items() if v is not None}
    update_credentialing(rid, changes)
    notify_activity(user["username"], "updated", "Credentialing",
                    f"Record #{rid}, fields: {', '.join(changes.keys())}")
    return {"ok": True}


@router.delete("/credentialing/{rid}")
def remove_cred(rid: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    delete_credentialing(rid)
    notify_activity(user["username"], "deleted", "Credentialing", f"Record #{rid}")
    return {"ok": True}


# ─── Enrollment ───────────────────────────────────────────────────────────────

class EnrollIn(BaseModel):
    client_id: int
    provider_id: Optional[int] = None
    ProviderName: Optional[str] = ""
    Payor: Optional[str] = ""
    EnrollType: Optional[str] = "Enrollment"
    Status: Optional[str] = "Not Started"
    SubmittedDate: Optional[str] = ""
    FollowUpDate: Optional[str] = ""
    ApprovedDate: Optional[str] = ""
    EffectiveDate: Optional[str] = ""
    Owner: Optional[str] = ""
    Notes: Optional[str] = ""
    sub_profile: Optional[str] = ""


class EnrollUpdate(BaseModel):
    ProviderName: Optional[str] = None
    Payor: Optional[str] = None
    EnrollType: Optional[str] = None
    Status: Optional[str] = None
    SubmittedDate: Optional[str] = None
    FollowUpDate: Optional[str] = None
    ApprovedDate: Optional[str] = None
    EffectiveDate: Optional[str] = None
    Owner: Optional[str] = None
    Notes: Optional[str] = None
    sub_profile: Optional[str] = None


@router.get("/enrollment")
def list_enroll(status: Optional[str] = None, client_id: Optional[int] = None,
               sub_profile: Optional[str] = None,
               hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    return get_enrollment(scope, status, sub_profile=sub_profile)


@router.post("/enrollment")
def add_enroll(body: EnrollIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    if user["role"] not in ("admin", "staff"):
        data["client_id"] = user["id"]
    eid = create_enrollment(data)
    notify_activity(user["username"], "created", "Enrollment",
                    f"Provider: {data.get('ProviderName','')}, Payor: {data.get('Payor','')}")
    return {"id": eid, "ok": True}


@router.put("/enrollment/{rid}")
def edit_enroll(rid: int, body: EnrollUpdate, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    changes = {k: v for k, v in body.model_dump().items() if v is not None}
    update_enrollment(rid, changes)
    notify_activity(user["username"], "updated", "Enrollment",
                    f"Record #{rid}, fields: {', '.join(changes.keys())}")
    return {"ok": True}


@router.delete("/enrollment/{rid}")
def remove_enroll(rid: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    delete_enrollment(rid)
    notify_activity(user["username"], "deleted", "Enrollment", f"Record #{rid}")
    return {"ok": True}


# ─── Eligibility / Benefits Verification ──────────────────────────────────────

class EligIn(BaseModel):
    client_id: Optional[int] = None
    PatientName: Optional[str] = ""
    DOB: Optional[str] = ""
    Payor: Optional[str] = ""
    MemberID: Optional[str] = ""
    PlanGroup: Optional[str] = ""
    Status: Optional[str] = "Pending"
    EffectiveDate: Optional[str] = ""
    TermDate: Optional[str] = ""
    Copay: Optional[str] = ""
    Deductible: Optional[str] = ""
    Coinsurance: Optional[str] = ""
    OOPMax: Optional[str] = ""
    PriorAuthRequired: Optional[str] = ""
    AuthNumber: Optional[str] = ""
    VerifiedBy: Optional[str] = ""
    VerifiedDate: Optional[str] = ""
    NextReverifyDate: Optional[str] = ""
    Notes: Optional[str] = ""
    sub_profile: Optional[str] = ""
    Stage: Optional[str] = "Received"


class EligUpdate(BaseModel):
    PatientName: Optional[str] = None
    DOB: Optional[str] = None
    Payor: Optional[str] = None
    MemberID: Optional[str] = None
    PlanGroup: Optional[str] = None
    Status: Optional[str] = None
    EffectiveDate: Optional[str] = None
    TermDate: Optional[str] = None
    Copay: Optional[str] = None
    Deductible: Optional[str] = None
    Coinsurance: Optional[str] = None
    OOPMax: Optional[str] = None
    PriorAuthRequired: Optional[str] = None
    AuthNumber: Optional[str] = None
    VerifiedBy: Optional[str] = None
    VerifiedDate: Optional[str] = None
    NextReverifyDate: Optional[str] = None
    Notes: Optional[str] = None
    sub_profile: Optional[str] = None
    Stage: Optional[str] = None
    IntakeFileId: Optional[int] = None
    IntakeFileName: Optional[str] = None
    ReportFileId: Optional[int] = None
    ReportFileName: Optional[str] = None
    CompletedBy: Optional[str] = None
    CompletedAt: Optional[str] = None


@router.get("/eligibility")
def list_elig(status: Optional[str] = None, client_id: Optional[int] = None,
              sub_profile: Optional[str] = None,
              hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    return get_eligibility(scope, status, sub_profile=sub_profile)


@router.post("/eligibility")
def add_elig(body: EligIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    # Eligibility is a SHARED, account-level board: every PCR eligibility login
    # must read/write the SAME account's records. So scope a client login's
    # writes to the account it tracks (its assigned account), NOT its own row id
    # — otherwise each verifier would only see records they personally entered.
    if user["role"] not in ("admin", "staff"):
        data["client_id"] = _client_account_id(user)
    elif not data.get("client_id"):
        data["client_id"] = _client_scope(user)
    if not data.get("client_id"):
        raise HTTPException(status_code=400, detail="client_id required")
    data["uploaded_by"] = (user.get("email") or user.get("username") or "").strip().lower()
    eid = create_eligibility(data)
    notify_activity(user["username"], "created", "Eligibility",
                    f"Patient: {data.get('PatientName','')}, Payor: {data.get('Payor','')}")
    return {"id": eid, "ok": True}


@router.put("/eligibility/{rid}")
def edit_elig(rid: int, body: EligUpdate, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    changes = {k: v for k, v in body.model_dump().items() if v is not None}
    update_eligibility(rid, changes)
    notify_activity(user["username"], "updated", "Eligibility",
                    f"Record #{rid}, fields: {', '.join(changes.keys())}")
    return {"ok": True}


@router.delete("/eligibility/{rid}")
def remove_elig(rid: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    delete_eligibility(rid)
    notify_activity(user["username"], "deleted", "Eligibility", f"Record #{rid}")
    return {"ok": True}


@router.post("/eligibility/{rid}/file")
async def upload_elig_file(
    rid: int,
    kind: str = Form("intake"),
    file: UploadFile = FastAPIFile(...),
    hub_session: Optional[str] = Cookie(None),
):
    """Attach a document to an eligibility record.

    kind='intake'  → a client-uploaded intake document.
    kind='report'  → the MedPharma-completed report; also marks the record
                     Completed and stamps who/when. The file is stored in the
                     same account's Documents so the existing download route and
                     account scoping apply unchanged.
    """
    user = _require_user(hub_session)
    rec = get_eligibility_one(rid)
    if not rec:
        raise HTTPException(404, "Eligibility record not found")
    kind = (kind or "intake").strip().lower()
    if kind not in ("intake", "report"):
        raise HTTPException(400, "kind must be 'intake' or 'report'")
    # Only MedPharma (admin/staff) may file the completed report.
    if kind == "report" and user.get("role") not in ("admin", "staff"):
        raise HTTPException(403, "Only MedPharma staff can upload the completed report")

    ext = os.path.splitext(file.filename or "")[1].lower()
    ALLOWED = (".xlsx", ".xls", ".csv", ".ods", ".odf", ".odt", ".odp",
               ".pdf", ".doc", ".docx", ".ppt", ".pptx", ".txt", ".rtf",
               ".png", ".jpg", ".jpeg")
    if ext not in ALLOWED:
        raise HTTPException(400, "Unsupported file type")
    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(413, "File too large. Maximum is 50MB")
    unique_name = f"{uuid.uuid4().hex}{ext}"
    with open(os.path.join(UPLOAD_DIR, unique_name), "wb") as f:
        f.write(content)

    scope = int(rec.get("client_id") or 0)
    category = "Eligibility Report" if kind == "report" else "Eligibility Intake"
    file_id = add_file(
        client_id=scope, filename=unique_name, original_name=file.filename or "file",
        file_type=("pdf" if ext == ".pdf" else "document"), file_size=len(content),
        category=category, description=f"{category} — {rec.get('PatientName', '')}",
        row_count=0, uploaded_by=user["username"],
    )
    if kind == "report":
        update_eligibility(rid, {
            "ReportFileId": file_id, "ReportFileName": file.filename or "report",
            "Stage": "Completed",
            "CompletedBy": (user.get("username") or ""),
            "CompletedAt": datetime.now().isoformat(),
        })
        notify_activity(user["username"], "completed report", "Eligibility",
                        f"{rec.get('PatientName', '')} — {file.filename or ''}")
    else:
        update_eligibility(rid, {
            "IntakeFileId": file_id, "IntakeFileName": file.filename or "intake",
        })
        notify_activity(user["username"], "uploaded intake", "Eligibility",
                        f"{rec.get('PatientName', '')} — {file.filename or ''}")
    return {"ok": True, "file_id": file_id,
            "stage": "Completed" if kind == "report" else rec.get("Stage", "Received")}


# ─── EDI Setup ────────────────────────────────────────────────────────────────

class EDIIn(BaseModel):
    client_id: int
    provider_id: Optional[int] = None
    ProviderName: Optional[str] = ""
    Payor: Optional[str] = ""
    EDIStatus: Optional[str] = "Not Started"
    ERAStatus: Optional[str] = "Not Started"
    EFTStatus: Optional[str] = "Not Started"
    SubmittedDate: Optional[str] = ""
    GoLiveDate: Optional[str] = ""
    PayerID: Optional[str] = ""
    Owner: Optional[str] = ""
    Notes: Optional[str] = ""
    sub_profile: Optional[str] = ""


class EDIUpdate(BaseModel):
    ProviderName: Optional[str] = None
    Payor: Optional[str] = None
    EDIStatus: Optional[str] = None
    ERAStatus: Optional[str] = None
    EFTStatus: Optional[str] = None
    SubmittedDate: Optional[str] = None
    GoLiveDate: Optional[str] = None
    PayerID: Optional[str] = None
    Owner: Optional[str] = None
    Notes: Optional[str] = None
    sub_profile: Optional[str] = None


@router.get("/edi")
def list_edi(client_id: Optional[int] = None, sub_profile: Optional[str] = None,
            hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    return get_edi(scope, sub_profile=sub_profile)


@router.post("/edi")
def add_edi(body: EDIIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    data = body.model_dump()
    if user["role"] not in ("admin", "staff"):
        data["client_id"] = user["id"]
    eid = create_edi(data)
    notify_activity(user["username"], "created", "EDI Setup",
                    f"Provider: {data.get('ProviderName','')}, Payor: {data.get('Payor','')}")
    return {"id": eid, "ok": True}


@router.put("/edi/{rid}")
def edit_edi(rid: int, body: EDIUpdate, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    changes = {k: v for k, v in body.model_dump().items() if v is not None}
    update_edi(rid, changes)
    notify_activity(user["username"], "updated", "EDI Setup",
                    f"Record #{rid}, fields: {', '.join(changes.keys())}")
    return {"ok": True}


@router.delete("/edi/{rid}")
def remove_edi(rid: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    delete_edi(rid)
    notify_activity(user["username"], "deleted", "EDI Setup", f"Record #{rid}")
    return {"ok": True}


# ─── Dashboard ────────────────────────────────────────────────────────────────

@router.get("/dashboard")
def dashboard(hub_session: Optional[str] = Cookie(None), member: Optional[str] = None):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    try:
        if isinstance(scope, int):
            auto_import_pending_claim_files(scope)
    except Exception as _e:
        log.warning("auto-import on dashboard failed: %s", _e)
    member_idents = _dashboard_member_scope(user, member)
    data = get_dashboard(scope, member_idents=member_idents)
    # Per-user dashboards already ARE one person's work, and the comprehensive
    # view is intentionally totals-only — so drop the per-member breakdown.
    data["billed_by_member"] = []
    data["scoped_member"] = (member_idents[0] if member_idents else None)
    data["user"] = user
    return data


@router.get("/dashboard/client/{client_id}")
def dashboard_for_client(client_id: int, sub_profile: Optional[str] = None,
                        member: Optional[str] = None,
                        hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    _assert_client_can_view(user, client_id)
    try:
        auto_import_pending_claim_files(client_id)
    except Exception as _e:
        log.warning("auto-import on client dashboard failed: %s", _e)
    member_idents = _dashboard_member_scope(user, member)
    data = get_dashboard(client_id, sub_profile=sub_profile, member_idents=member_idents)
    data["billed_by_member"] = []
    data["scoped_member"] = (member_idents[0] if member_idents else None)
    data["user"] = user
    return data


# ─── File Uploads ───────────────────────────────────────────────────────────

_DATA_ROOT = "/data" if os.path.isdir("/data") else "data"
UPLOAD_DIR = os.path.join(_DATA_ROOT, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ─── Team Production ──────────────────────────────────────────────────────────

class ProductionLogIn(BaseModel):
    client_id: Optional[int] = None
    work_date: str
    category: str = ""
    task_description: str = ""
    quantity: int = 0
    time_spent: float = 0
    notes: str = ""
    attachment_file_id: Optional[int] = None
    attachment_name: str = ""


class ProductionRelinkIn(BaseModel):
    source_client_ids: Optional[list[int]] = None
    usernames: Optional[list[str]] = None
    dry_run: bool = False
    max_rows: int = 5000


class ProductionReportJobIn(BaseModel):
    client_id: Optional[int] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None


def _import_rows_to_production(
    rows: list[dict],
    client_id: int,
    default_username: str,
    dry_run: bool = False,
    progress_cb=None,
) -> dict:
    # Tiered column finder: exact > starts-with > contains.
    def _find_col(headers: list[str], *candidates: str) -> Optional[str]:
        hl = [h.lower().strip() for h in headers]
        for c in candidates:
            cl = c.lower()
            for i, h in enumerate(hl):
                if h == cl:
                    return headers[i]
        for c in candidates:
            cl = c.lower()
            for i, h in enumerate(hl):
                if h.startswith(cl):
                    return headers[i]
        for c in candidates:
            cl = c.lower()
            for i, h in enumerate(hl):
                if cl in h:
                    return headers[i]
        return None

    headers = list(rows[0].keys()) if rows else []
    col_date = _find_col(headers, "work date", "work_date", "date", "day")
    col_username = _find_col(headers, "username", "user name", "user", "agent", "rep", "employee", "staff", "technician", "tech")
    col_category = _find_col(headers, "category", "task type", "activity type", "work type", "type")
    col_task = _find_col(headers, "task description", "task", "description", "work performed", "work done", "activity", "detail")
    col_qty = _find_col(headers, "quantity", "qty", "count", "units", "items")
    col_hours = _find_col(headers, "hours", "time spent", "duration", "hrs")
    col_notes = _find_col(headers, "notes", "comments", "comment", "additional", "remarks")

    used: set[str] = set()

    def _claim(col: Optional[str]) -> Optional[str]:
        if col is None or col in used:
            return None
        used.add(col)
        return col

    col_date = _claim(col_date)
    col_task = _claim(col_task)
    col_username = col_username if col_username not in used else None
    col_category = col_category if col_category not in used else None
    col_qty = col_qty if col_qty not in used else None
    col_hours = col_hours if col_hours not in used else None
    col_notes = col_notes if col_notes not in used else None

    if not col_task:
        remaining = [h for h in headers if h not in used]
        col_task = _find_col(remaining, "task", "description", "work", "notes", "detail", "activity")
        if col_task:
            used.add(col_task)

    if not col_date or not col_task:
        missing = []
        if not col_date:
            missing.append("Date")
        if not col_task:
            missing.append("Task/Description")
        raise ValueError(f"Cannot find required column(s): {', '.join(missing)}. Headers found: {headers}")

    date_formats = ["%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%d/%m/%Y", "%Y/%m/%d", "%Y-%m-%d %H:%M:%S"]

    def _parse_date(val) -> Optional[str]:
        if val is None or str(val).strip() in ("", "None"):
            return None
        if hasattr(val, "strftime"):
            return val.strftime("%Y-%m-%d")
        if isinstance(val, (int, float)):
            try:
                from datetime import datetime as _dt, timedelta as _td
                dt = _dt(1899, 12, 30) + _td(days=float(val))
                if 1970 <= dt.year <= 2100:
                    return dt.strftime("%Y-%m-%d")
            except (ValueError, OverflowError):
                pass
            return None
        s = str(val).strip()
        for fmt in date_formats:
            try:
                from datetime import datetime as _dt
                return _dt.strptime(s, fmt).strftime("%Y-%m-%d")
            except ValueError:
                continue
        return None

    imported = 0
    skipped = 0
    errors: list[str] = []
    preview: list[dict] = []
    total = len(rows)

    for i, row in enumerate(rows, start=2):
        try:
            work_date = _parse_date(row.get(col_date, ""))
            task_desc = str(row.get(col_task, "")).strip()
            if not work_date or not task_desc:
                skipped += 1
                continue

            username = str(row.get(col_username, "") or default_username).strip() or default_username
            category = str(row.get(col_category, "") or "General").strip() or "General"
            notes = str(row.get(col_notes, "") or "").strip()

            raw_qty = row.get(col_qty, "")
            try:
                quantity = int(float(str(raw_qty))) if raw_qty not in (None, "") else 1
            except (ValueError, TypeError):
                quantity = 1

            raw_hrs = row.get(col_hours, "")
            try:
                time_spent = float(str(raw_hrs)) if raw_hrs not in (None, "") else 0.0
            except (ValueError, TypeError):
                time_spent = 0.0

            row_payload = {
                "client_id": client_id,
                "work_date": work_date,
                "username": username,
                "category": category,
                "task_description": task_desc,
                "quantity": quantity,
                "time_spent": time_spent,
                "notes": notes,
            }

            if dry_run:
                if len(preview) < 25:
                    preview.append(row_payload)
            else:
                add_production_log(row_payload)
            imported += 1
        except Exception as exc:
            errors.append(f"Row {i}: {str(exc)[:120]}")

        if progress_cb and total > 0 and (imported + skipped) % 25 == 0:
            progress_cb(min(95, int(((imported + skipped) / total) * 100)))

    return {
        "imported": imported,
        "skipped": skipped,
        "errors": errors,
        "preview": preview,
    }


def _start_production_report_job(job_id: str, user: dict):
    def _runner():
        try:
            job = get_job(job_id)
            if not job:
                return
            payload = job.get("payload") or {}
            requested_client_id = payload.get("client_id")
            start_date = payload.get("start_date")
            end_date = payload.get("end_date")

            set_job_running(job_id, progress=5)
            append_job_event(job_id, "start", "Starting production report build")

            role = (payload.get("requested_by_role") or user.get("role") or "").lower()
            # Admins/Eric → comprehensive roll-up for the (optional) selected
            # client. Everyone else → their own self-view across all accounts.
            if role == "admin" or _is_eric(user):
                report = get_production_report(requested_client_id, start_date, end_date)
            else:
                report = get_production_report(None, start_date, end_date,
                                               username=user.get("username"))

            update_job_progress(job_id, 90)
            result = {
                "report": report,
                "selected_client_id": requested_client_id,
                "start_date": start_date,
                "end_date": end_date,
                "generated_at": datetime.now().isoformat(),
            }
            complete_job(job_id, result=result)
            append_job_event(job_id, "done", "Production report ready")
        except Exception as exc:
            fail_job(job_id, str(exc))
            append_job_event(job_id, "error", f"Report build failed: {str(exc)[:200]}", "error")

    threading.Thread(target=_runner, daemon=True).start()


@router.get("/production")
def get_production(client_id: Optional[int] = None,
                   start_date: Optional[str] = None,
                   end_date: Optional[str] = None,
                   username: Optional[str] = None,
                   hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    role = (user.get("role") or "").lower()
    uname = (username or "").strip() or None
    if role == "admin":
        # Team Production Report: admin can filter per client, per user, or
        # combined (both). No filter => the whole team across every account.
        logs = list_production_logs(client_id, start_date, end_date, username=uname)
        return {"logs": logs, "fallback_all_clients": False, "selected_client_id": client_id}
    if role == "staff":
        # User Production: a staff user only ever sees their OWN logged work,
        # across whichever accounts they're working.
        logs = list_production_logs(None, start_date, end_date, username=user.get("username"))
        return {"logs": logs, "fallback_all_clients": False, "selected_client_id": client_id}
    # Client account: production logged against their own account.
    scope = client_id or _client_scope(user)
    logs = list_production_logs(scope, start_date, end_date, username=None)
    return {"logs": logs, "fallback_all_clients": False, "selected_client_id": client_id}


@router.post("/production")
def create_production_log(body: ProductionLogIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = body.client_id if body.client_id is not None else (_client_scope(user) if _client_scope(user) is not None else user["id"])
    data = body.model_dump()
    data["client_id"] = scope
    data["username"] = user["username"]
    log_id = add_production_log(data)
    notify_activity(user["username"], "logged production", "Time Tracking",
                    f"{data.get('hours',0)}h — {data.get('task_type','')}: {data.get('description','')[:60]}")
    return {"id": log_id, "ok": True}


@router.delete("/production/{log_id}")
def remove_production_log(log_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    if user.get("role") in ("admin", "staff"):
        deleted = delete_production_log(log_id)
    else:
        deleted = delete_production_log(log_id, username=user.get("username", ""))
    if not deleted:
        raise HTTPException(status_code=404, detail="Production entry not found")
    notify_activity(user["username"], "deleted", "Time Tracking", f"Log #{log_id}")
    return {"ok": True}


@router.post("/production/import")
async def import_production_excel(
    client_id: Optional[int] = Query(None),
    file: UploadFile = FastAPIFile(...),
    dry_run: bool = Query(False, description="If true, parse and preview without saving."),
    async_job: bool = Query(False, description="If true, run import as a tracked background job."),
    hub_session: Optional[str] = Cookie(None),
):
    """Import production log entries from Excel / CSV / structured PDF files."""
    user = _require_user(hub_session)
    if not client_id:
        raise HTTPException(status_code=422, detail="client_id is required")

    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in {"xlsx", "xls", "csv", "ods", "odf", "pdf", "doc", "docx"}:
        raise HTTPException(status_code=422, detail="File must be .xlsx, .xls, .csv, .ods, .odf, .pdf, .doc, or .docx")

    content = await file.read()
    try:
        if ext == "pdf":
            rows = _parse_pdf_rows(content)
        elif ext in ("doc", "docx"):
            rows = _parse_docx_rows(content)
        else:
            rows = _parse_excel_rows(content, f".{ext}")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not parse file: {exc}")
    if not rows:
        raise HTTPException(status_code=422, detail="No data rows found in file")

    if async_job and not dry_run:
        job = create_job(
            account_id=client_id,
            job_type="production_import",
            created_by=user.get("username", ""),
            payload={
                "client_id": client_id,
                "source_type": ext,
                "filename": file.filename or "upload",
                "total_rows": len(rows),
            },
        )
        append_job_event(job["id"], "queued", f"Queued import of {len(rows)} rows")

        def _runner():
            try:
                set_job_running(job["id"], progress=5)
                append_job_event(job["id"], "start", "Parsing and importing rows")

                def _progress_cb(pct: int):
                    update_job_progress(job["id"], pct)

                outcome = _import_rows_to_production(
                    rows=rows,
                    client_id=client_id,
                    default_username=user["username"],
                    dry_run=False,
                    progress_cb=_progress_cb,
                )
                result = {
                    "source_type": ext,
                    "filename": file.filename or "upload",
                    "total_rows": len(rows),
                    "imported": outcome["imported"],
                    "skipped": outcome["skipped"],
                    "errors": outcome["errors"],
                }
                complete_job(job["id"], result=result)
                append_job_event(job["id"], "done", f"Imported {outcome['imported']} rows")
                notify_activity(user["username"], "imported", "Time Tracking",
                                f"{outcome['imported']} entries for client #{client_id}")
            except Exception as exc:
                fail_job(job["id"], str(exc))
                append_job_event(job["id"], "error", f"Import failed: {str(exc)[:200]}", "error")

        threading.Thread(target=_runner, daemon=True).start()
        return {
            "ok": True,
            "job_id": job["id"],
            "status": "queued",
            "async_job": True,
            "source_type": ext,
            "total_rows": len(rows),
        }

    outcome = _import_rows_to_production(
        rows=rows,
        client_id=client_id,
        default_username=user["username"],
        dry_run=dry_run,
    )

    if not dry_run:
        notify_activity(user["username"], "imported", "Time Tracking",
                        f"{outcome['imported']} entries for client #{client_id}")

    return {
        "ok": True,
        "dry_run": dry_run,
        "source_type": ext,
        "imported": outcome["imported"],
        "skipped": outcome["skipped"],
        "errors": outcome["errors"],
        "preview": outcome["preview"],
    }


@router.post("/jobs/production-report")
def create_production_report_job(body: ProductionReportJobIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    role = user.get("role")
    requested_client_id = body.client_id
    if role not in ("admin", "staff"):
        requested_client_id = user["id"]

    job = create_job(
        account_id=requested_client_id,
        job_type="production_report_pack",
        created_by=user.get("username", ""),
        payload={
            "client_id": requested_client_id,
            "start_date": body.start_date,
            "end_date": body.end_date,
            "requested_by_role": role,
        },
    )
    append_job_event(job["id"], "queued", "Queued production report pack")
    _start_production_report_job(job["id"], user)
    return {"ok": True, "job_id": job["id"], "status": "queued"}


@router.get("/jobs")
def jobs_list(
    status: Optional[str] = None,
    job_type: Optional[str] = None,
    limit: int = Query(50, ge=1, le=200),
    hub_session: Optional[str] = Cookie(None),
):
    user = _require_user(hub_session)
    account_scope = _client_scope(user)
    rows = list_jobs(
        account_id=account_scope,
        status=(status or "").strip(),
        job_type=(job_type or "").strip(),
        limit=limit,
    )
    return {"jobs": rows}


@router.get("/jobs/{job_id}")
def jobs_get(job_id: str, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    job = get_job(job_id, include_events=True)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    scope = _client_scope(user)
    if scope is not None and int(job.get("account_id") or 0) != int(scope):
        raise HTTPException(status_code=403, detail="Forbidden")
    return job


@router.post("/jobs/{job_id}/retry")
def jobs_retry(job_id: str, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    scope = _client_scope(user)
    if scope is not None and int(job.get("account_id") or 0) != int(scope):
        raise HTTPException(status_code=403, detail="Forbidden")

    if job.get("status") != "error":
        raise HTTPException(status_code=409, detail="Only failed jobs can be retried")

    if job.get("job_type") != "production_report_pack":
        raise HTTPException(status_code=409, detail="Retry currently supported for production report jobs")

    reset = reset_job_for_retry(job_id)
    if not reset:
        raise HTTPException(status_code=404, detail="Job not found")
    append_job_event(job_id, "queued", f"Retry requested by {user.get('username', '')}")
    _start_production_report_job(job_id, user)
    return {"ok": True, "job_id": job_id, "status": "queued"}


@router.get("/production/report")
def production_report(client_id: Optional[int] = None,
                      start_date: Optional[str] = None,
                      end_date: Optional[str] = None,
                      hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    # Admins and Eric get the comprehensive roll-up (every biller, combined,
    # optionally narrowed to one client). Every other user (Susan / Melissa /
    # Jessica) gets a self-view: only what THEY billed, posted, and were paid,
    # across whichever accounts they work.
    if user.get("role") == "admin" or _is_eric(user):
        report = get_production_report(client_id, start_date, end_date)
    else:
        report = get_production_report(None, start_date, end_date,
                                       username=user.get("username"))
    report["fallback_all_clients"] = False
    report["selected_client_id"] = client_id
    return report


@router.get("/production/report/download")
def download_production_report(
    client_id: Optional[int] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    hub_session: Optional[str] = Cookie(None),
):
    """Return a branded, printable HTML production report with MedPharma logo."""
    user = _require_user(hub_session)
    # Same scoping as the on-screen report: admins/Eric get the combined
    # roll-up; a biller gets only their own production.
    if user.get("role") == "admin" or _is_eric(user):
        data = get_production_report(client_id, start_date, end_date)
    else:
        data = get_production_report(None, start_date, end_date,
                                     username=user.get("username"))

    from html import escape as _esc

    period_label = f"{start_date or 'All time'} — {end_date or 'today'}"
    by_user  = data.get("by_user", [])
    by_cat   = data.get("by_category", [])
    details  = data.get("details", [])
    pay_details = data.get("payment_details", [])
    flags    = data.get("time_management_flags", [])

    # ── Team summary rows ──────────────────────────────────────────────
    def _user_rows():
        if not by_user:
            return "<tr><td colspan='6' style='text-align:center;color:#9ca3af'>No team data for this period</td></tr>"
        return "".join(
            f"<tr><td><strong>{_esc(str(u.get('username','')))}</strong></td>"
            f"<td>{u.get('claims_billed',0)}</td>"
            f"<td>${u.get('claims_billed_amount',0):,.2f}</td>"
            f"<td>{u.get('claims_denied',0)}</td>"
            f"<td>{u.get('payments_posted',0)}</td>"
            f"<td>${u.get('claims_paid_amount', u.get('payments_amount',0)):,.2f}</td></tr>"
            for u in by_user
        )

    def _cat_rows():
        if not by_cat:
            return "<tr><td colspan='4' style='text-align:center;color:#9ca3af'>No data</td></tr>"
        return "".join(
            f"<tr><td>{_esc(str(c.get('category','Uncategorized')))}</td>"
            f"<td>{c.get('total_entries',0)}</td>"
            f"<td>{c.get('total_quantity',0)}</td>"
            f"<td>{c.get('total_hours',0)}h</td></tr>"
            for c in by_cat
        )

    def _detail_rows():
        if not details:
            return "<tr><td colspan='7' style='text-align:center;color:#9ca3af'>No entries in this period</td></tr>"
        return "".join(
            f"<tr><td>{_esc(str(d.get('work_date','')))}</td>"
            f"<td>{_esc(str(d.get('username','')))}</td>"
            f"<td>{_esc(str(d.get('category','')))}</td>"
            f"<td>{_esc(str(d.get('task_description','')))}</td>"
            f"<td style='text-align:center'>{d.get('quantity',0)}</td>"
            f"<td style='text-align:center'>{d.get('time_spent',0)}h</td>"
            f"<td style='font-size:12px;color:#6b7280'>{_esc(str(d.get('notes','') or ''))}</td></tr>"
            for d in details
        )

    def _payment_rows():
        if not pay_details:
            return "<tr><td colspan='5' style='text-align:center;color:#9ca3af'>No payments posted in this period</td></tr>"
        return "".join(
            f"<tr><td>{_esc(str(pd.get('post_date','')))}</td>"
            f"<td>{_esc(str(pd.get('username','')))}</td>"
            f"<td>{_esc(str(pd.get('ClaimKey','')))}</td>"
            f"<td>{_esc(str(pd.get('PayerType','') or ''))}</td>"
            f"<td style='text-align:right'>${float(pd.get('PaymentAmount') or 0):,.2f}</td></tr>"
            for pd in pay_details
        )

    def _flag_section():
        if not flags:
            return ""
        rows = "".join(
            f"<tr><td>{_esc(str(f.get('username','')))}</td>"
            f"<td>{f.get('avg_hours_per_day',0)}h/day</td>"
            f"<td>{f.get('days_worked',0)}</td>"
            f"<td style='color:#dc2626'>{_esc(str(f.get('recommendation','')))}</td></tr>"
            for f in flags
        )
        return f"""
        <section class="section">
          <h2 style="color:#dc2626">⚠️ Time Management Alerts</h2>
          <table><thead><tr><th>Team Member</th><th>Avg Hrs/Day</th><th>Days Worked</th><th>Recommendation</th></tr></thead>
          <tbody>{rows}</tbody></table>
        </section>"""

    from datetime import date as _date
    generated = business_today().strftime("%B %d, %Y")

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>MedPharma Production Report — {_esc(period_label)}</title>
<style>
  * {{ box-sizing:border-box; margin:0; padding:0; }}
  body {{ font-family:'Segoe UI',Arial,sans-serif; font-size:13px; color:#1f2937; background:#fff; padding:0; }}
  .page {{ max-width:960px; margin:0 auto; padding:32px 24px; }}
  .header {{ display:flex; align-items:center; gap:20px; border-bottom:3px solid #1d4ed8; padding-bottom:20px; margin-bottom:28px; }}
  .logo-img {{ height:64px; width:auto; object-fit:contain; }}
  .header-text h1 {{ font-size:22px; font-weight:800; color:#1d4ed8; letter-spacing:-0.5px; }}
  .header-text p {{ font-size:13px; color:#6b7280; margin-top:4px; }}
  .meta-bar {{ display:flex; gap:28px; background:#f0f4ff; border-radius:8px; padding:12px 18px; margin-bottom:24px; font-size:13px; }}
  .meta-bar b {{ color:#1d4ed8; }}
  .section {{ margin-bottom:28px; }}
  .section h2 {{ font-size:15px; font-weight:700; color:#374151; margin-bottom:10px; padding-bottom:6px; border-bottom:1px solid #e5e7eb; }}
  table {{ width:100%; border-collapse:collapse; font-size:12.5px; }}
  th {{ background:#1d4ed8; color:#fff; text-align:left; padding:8px 10px; font-weight:600; font-size:12px; }}
  td {{ padding:7px 10px; border-bottom:1px solid #f3f4f6; }}
  tr:nth-child(even) td {{ background:#f9fafb; }}
  .footer {{ margin-top:36px; padding-top:14px; border-top:1px solid #e5e7eb; font-size:11px; color:#9ca3af; display:flex; justify-content:space-between; }}
  @media print {{
    body {{ padding:0; }}
    .page {{ padding:20px 16px; }}
    .no-print {{ display:none; }}
  }}
</style>
</head>
<body>
<div class="page">
  <div class="header">
    <img class="logo-img" src="https://medpharmasc.com/wp-content/uploads/2024/11/IMG_2392.png" alt="MedPharma Logo" crossorigin="anonymous">
    <div class="header-text">
      <h1>MedPharma Internal Hub</h1>
    <p>Work Production Report &nbsp;|&nbsp; {_esc(period_label)}</p>
    </div>
  </div>

  <div class="meta-bar">
    <span><b>Period:</b> {_esc(period_label)}</span>
    <span><b>Claims Submitted:</b> {data.get('billed_total_count', 0)} (${data.get('billed_total_amount', 0):,.2f})</span>
    <span><b>Paid:</b> ${data.get('paid_total_amount', data.get('payments_total_amount', 0)):,.2f}</span>
    <span><b>Prior Denials &amp; Rebill:</b> {data.get('denied_total_count', 0)} (${data.get('denied_total_amount', 0):,.2f}) <i style="color:#6b7280">— submitted claims, incl. in Submitted</i></span>
    <span><b>Posted:</b> {data.get('payments_total_count', 0)}</span>
    <span><b>Rolling AR (pre-{_esc(str(data.get('rolling_ar_cutoff','')))}):</b> ${data.get('rolling_ar', 0):,.2f}</span>
    <span><b>Generated:</b> {generated}</span>
  </div>

  {_flag_section()}

  <section class="section">
    <h2>👥 Production by User</h2>
    <table>
      <thead><tr><th>Team Member</th><th>Claims Submitted</th><th>$ Submitted</th><th>Prior Denials &amp; Rebill</th><th>Posted</th><th>$ Paid</th></tr></thead>
      <tbody>{_user_rows()}</tbody>
    </table>
    <p style="font-size:11px;color:#6b7280;margin-top:8px">
      <b>Note:</b> Every claim here was <b>submitted</b>. The Prior Denials &amp;
      Rebill column flags how many submitted claims had a prior denial and were
      rebilled/resubmitted — a rebilled claim is still a submitted, billed claim
      (the prior denial only reflects what it was before it was reworked and
      rebilled), so it stays counted in Claims Submitted. Posted reflects payments
      posted (same dollars as Paid); each claim's posting date is recorded in its
      claim notes.
    </p>
  </section>

  <section class="section">
    <h2>💵 Payments Posted by User</h2>
    <table>
      <thead><tr><th>Post Date</th><th>Posted By</th><th>Claim</th><th>Payer</th><th>Amount</th></tr></thead>
      <tbody>{_payment_rows()}</tbody>
    </table>
  </section>

  <section class="section">
    <h2>📊 Work by Category</h2>
    <table>
      <thead><tr><th>Category</th><th>Entries</th><th>Items</th><th>Hours</th></tr></thead>
      <tbody>{_cat_rows()}</tbody>
    </table>
  </section>

  <section class="section">
    <h2>📋 Detailed Work Production</h2>
    <table>
      <thead><tr><th>Date</th><th>User</th><th>Category</th><th>Task Description</th><th>Qty</th><th>Hours</th><th>Notes</th></tr></thead>
      <tbody>{_detail_rows()}</tbody>
    </table>
  </section>

  <div class="footer">
    <span>MedPharma Internal Hub &nbsp;|&nbsp; <a href="https://medpharmasc.com">medpharmasc.com</a></span>
    <span class="no-print"><button onclick="window.print()" style="background:#1d4ed8;color:#fff;border:none;padding:6px 18px;border-radius:6px;cursor:pointer;font-size:13px">🖨️ Print / Save PDF</button></span>
    <span>Generated {generated}</span>
  </div>
</div>
</body>
</html>"""

    return Response(content=html, media_type="text/html")


@router.get("/notifications/status")
def notifications_status_endpoint(hub_session: Optional[str] = Cookie(None)):
    """Return the live notification channel configuration (full admin only)."""
    _require_full_admin(hub_session)
    return get_notification_status()


@router.post("/notifications/test")
def notifications_test_endpoint(hub_session: Optional[str] = Cookie(None)):
    """Fire a real test notification through configured email/SMS channels.

    Full admin only. Returns the per-channel delivery results plus current status.
    """
    user = _require_full_admin(hub_session)
    return send_test_notification(triggered_by=user.get("username") or "admin")


@router.get("/notifications/debug")
def notifications_debug_endpoint(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: return notification runtime/config debug details."""
    _require_admin(hub_session)
    return get_notification_debug()


@router.post("/notifications/daily-report")
def send_daily_report_now(hub_session: Optional[str] = Cookie(None)):
    """Admin-only: immediately send the daily account summary report (email + SMS)."""
    _require_admin(hub_session)
    try:
        send_daily_account_summary()
        return {"ok": True, "message": "Daily report sent to configured recipients"}
    except Exception as e:
        raise HTTPException(500, f"Failed to send daily report: {e}")


@router.get("/production/snapshot")
def production_snapshot(work_date: Optional[str] = None, hub_session: Optional[str] = Cookie(None)):
    """Admin-only: get per-user production snapshot for a date (defaults to today)."""
    _require_admin(hub_session)
    return get_user_production_snapshot(work_date)


@router.post("/admin/production/relink-kindercare")
def relink_kindercare_production(body: ProductionRelinkIn, hub_session: Optional[str] = Cookie(None)):
    """Safely copy legacy production rows into the KinderCare account.

    This endpoint is idempotent: rows that already exist for KinderCare are skipped.
    """
    _require_full_admin(hub_session)

    usernames = [str(u or "").strip().lower() for u in (body.usernames or []) if str(u or "").strip()]
    if not usernames:
        usernames = ["mike", "sarah"]

    source_ids = [int(x) for x in (body.source_client_ids or []) if int(x) > 0]
    max_rows = max(1, min(int(body.max_rows or 5000), 20000))

    conn = get_db()
    conn.row_factory = sqlite3.Row
    try:
        cur = conn.cursor()
        target = cur.execute(
            """
            SELECT id, username, company
            FROM clients
            WHERE lower(username)='kindercare' OR lower(company) LIKE '%kindercare%'
            ORDER BY id DESC
            LIMIT 1
            """
        ).fetchone()
        if not target:
            raise HTTPException(status_code=404, detail="KinderCare account not found")

        target_id = int(target["id"])
        target_existing = int(
            cur.execute("SELECT COUNT(*) FROM team_production WHERE client_id=?", (target_id,)).fetchone()[0]
        )

        conditions = ["client_id <> ?"]
        params: list[object] = [target_id]
        if source_ids:
            conditions.append("client_id IN ({})".format(",".join(["?"] * len(source_ids))))
            params.extend(source_ids)
        if usernames:
            conditions.append("lower(username) IN ({})".format(",".join(["?"] * len(usernames))))
            params.extend(usernames)

        where = " AND ".join(conditions)
        rows = cur.execute(
            f"""
            SELECT id, client_id, work_date, username, category, task_description, quantity, time_spent, notes
            FROM team_production
            WHERE {where}
            ORDER BY id ASC
            LIMIT ?
            """,
            [*params, max_rows],
        ).fetchall()

        copied = 0
        skipped_existing = 0
        sample = []
        for row in rows:
            exists = cur.execute(
                """
                SELECT 1
                FROM team_production
                WHERE client_id=?
                  AND work_date=?
                  AND lower(username)=lower(?)
                  AND category=?
                  AND task_description=?
                  AND IFNULL(quantity, 0)=IFNULL(?, 0)
                  AND ABS(IFNULL(time_spent, 0)-IFNULL(?, 0)) < 0.0001
                  AND IFNULL(notes, '')=IFNULL(?, '')
                LIMIT 1
                """,
                (
                    target_id,
                    row["work_date"],
                    row["username"],
                    row["category"],
                    row["task_description"],
                    row["quantity"],
                    row["time_spent"],
                    row["notes"],
                ),
            ).fetchone()
            if exists:
                skipped_existing += 1
                continue

            if not body.dry_run:
                cur.execute(
                    """
                    INSERT INTO team_production
                    (client_id, work_date, username, category, task_description, quantity, time_spent, notes)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        target_id,
                        row["work_date"],
                        row["username"],
                        row["category"],
                        row["task_description"],
                        row["quantity"],
                        row["time_spent"],
                        row["notes"],
                    ),
                )
            copied += 1
            if len(sample) < 10:
                sample.append(
                    {
                        "source_row_id": int(row["id"]),
                        "source_client_id": int(row["client_id"]),
                        "work_date": row["work_date"],
                        "username": row["username"],
                        "category": row["category"],
                    }
                )

        if not body.dry_run:
            conn.commit()

        target_after = int(
            cur.execute("SELECT COUNT(*) FROM team_production WHERE client_id=?", (target_id,)).fetchone()[0]
        )
        return {
            "ok": True,
            "dry_run": bool(body.dry_run),
            "target": {"id": target_id, "username": target["username"], "company": target["company"]},
            "target_existing_before": target_existing,
            "source_candidates": len(rows),
            "copied": copied,
            "skipped_existing": skipped_existing,
            "target_after": target_after,
            "sample": sample,
        }
    finally:
        conn.close()


# ─── Files ────────────────────────────────────────────────────────────────────

@router.get("/files")
def get_files(client_id: Optional[int] = None, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id if client_id is not None else _doc_scope(user)
    # Seamless ingestion: before listing, auto-import any claim-shaped spreadsheets
    # that were saved as plain documents, so the team never has to click "import".
    try:
        if isinstance(scope, int):
            auto_import_pending_claim_files(scope)
    except Exception as _e:
        log.warning("auto-import on file list failed: %s", _e)
    files = list_files(scope)
    return {"files": files}


@router.post("/files/upload")
async def upload_file(
    file: UploadFile = FastAPIFile(...),
    category: str = Form("General"),
    description: str = Form(""),
    client_id: Optional[int] = Form(None),
    sub_profile: Optional[str] = Form(None),
    hub_session: Optional[str] = Cookie(None),
):
    user = _require_user(hub_session)
    if client_id is not None:
        scope = client_id
    elif user.get("role") in ("admin", "staff"):
        scope = _client_scope(user)
        if scope is None:
            # Staff have no single scope; route data to the one client account
            # instead of the staff member's own id (which orphans the claims).
            scope = _single_client_account_or(user["id"])
    else:
        # Client sub-users upload into the shared account they belong to so the
        # rest of the team on that account sees the file too.
        scope = _client_upload_account(user)

    # Validate type
    ext = os.path.splitext(file.filename or "")[1].lower()
    ALLOWED_UPLOAD_EXTS = (
        ".xlsx", ".xls", ".csv", ".ods", ".odf", ".odt", ".odp",
        ".pdf", ".doc", ".docx", ".ppt", ".pptx", ".txt", ".rtf",
    )
    if ext not in ALLOWED_UPLOAD_EXTS:
        raise HTTPException(400, "Unsupported file type. Allowed: Word (.doc/.docx), Excel (.xlsx/.xls/.csv), OpenDocument (.odt/.ods/.odp/.odf), PDF, PowerPoint (.ppt/.pptx), .txt, .rtf")

    spreadsheet_exts = (".xlsx", ".xls", ".csv", ".ods", ".odf")
    if ext in spreadsheet_exts:
        file_type = "excel"
    elif ext == ".pdf":
        file_type = "pdf"
    else:
        file_type = "document"
    unique_name = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join(UPLOAD_DIR, unique_name)

    content = await file.read()
    file_size = len(content)

    # Enforce upload size limit (50 MB)
    MAX_UPLOAD_SIZE = 50 * 1024 * 1024
    if file_size > MAX_UPLOAD_SIZE:
        raise HTTPException(413, f"File too large. Maximum is {MAX_UPLOAD_SIZE // (1024*1024)}MB")

    with open(dest, "wb") as f:
        f.write(content)

    # Count rows for Excel/CSV
    row_count = 0
    if file_type == "excel":
        try:
            import csv, io
            if ext == ".csv":
                reader = csv.reader(io.StringIO(content.decode("utf-8", errors="replace")))
                row_count = max(0, sum(1 for _ in reader) - 1)
            else:
                row_count = len(_parse_excel_rows(content, ext, combine_sheets=True))
        except Exception:
            row_count = 0

    requested_category = (category or "General").strip() or "General"
    effective_category = requested_category
    category_source = "requested"
    infer_debug = None
    claims_fallback = False

    if file_type == "excel" and requested_category not in DATA_IMPORT_CATEGORIES:
        inferred, infer_debug = _infer_excel_category(content, ext, file.filename or "", description or "")
        if inferred in DATA_IMPORT_CATEGORIES:
            effective_category = inferred
            category_source = "auto"
        else:
            # Unclassified spreadsheet: do NOT silently shelve it as an inert
            # document. Attempt a Claims import so every uploaded worklist
            # computes even when its headers don't score. The importer skips
            # rows that don't map to claim columns, so a genuinely non-claims
            # sheet imports 0 rows and is reverted to a document below.
            effective_category = "Claims"
            category_source = "auto-claims-fallback"
            claims_fallback = True
    elif file_type in ("pdf", "document") and requested_category not in DATA_IMPORT_CATEGORIES:
        # Universal read: the team routinely sends daily worklists / ERAs as PDF
        # or Word, not just spreadsheets. Don't shelve them as inert documents —
        # attempt a Claims read so the rows compute. The importer extracts table
        # rows and skips any that don't map to claim columns, so a non-claims PDF
        # imports 0 rows and is reverted to a plain document below.
        effective_category = "Claims"
        category_source = "auto-claims-fallback"
        claims_fallback = True

    file_id = add_file(
        client_id=scope,
        filename=unique_name,
        original_name=file.filename or "file",
        file_type=file_type,
        file_size=file_size,
        category=effective_category,
        description=description,
        row_count=row_count,
        uploaded_by=user["username"],
    )

    # ── Auto-import data when category matches a known section and file is Excel/CSV ──
    imported = 0
    import_errors = []
    import_category = None
    if effective_category in DATA_IMPORT_CATEGORIES and file_type in ("excel", "pdf", "document"):
        import_category = effective_category
        _uploader = user.get("username") or ""
        _sub_profile = (sub_profile or "").strip()
        try:
            if effective_category == "Claims":
                imported, import_errors = _import_claims_from_excel(
                    content, ext, scope, uploaded_by=_uploader,
                    default_sub_profile=_sub_profile)
            elif file_type == "excel" and effective_category == "Credentialing":
                imported, import_errors = _import_credentialing_from_excel(content, ext, scope, uploaded_by=_uploader)
                if import_errors and any('header' in e.lower() or 'no rows' in e.lower() for e in import_errors):
                    import_errors.append("Required headers: Provider, Payor, Type, Status, Submitted, Follow Up, Approved, Expiration, Owner, Notes, Sub Profile")
            elif file_type == "excel" and effective_category == "Enrollment":
                imported, import_errors = _import_enrollment_from_excel(content, ext, scope, uploaded_by=_uploader)
            elif file_type == "excel" and effective_category == "EDI":
                imported, import_errors = _import_edi_from_excel(content, ext, scope, uploaded_by=_uploader)
        except Exception as e:
            import_errors = [str(e)]

    # ── Funnel notification to admin (every upload, regardless of category) ──
    # Claims-fallback that found nothing claim-like → it wasn't claims data.
    # Revert to a plain document under the originally requested category so it
    # isn't mis-shelved under Claims.
    if claims_fallback and imported == 0:
        effective_category = requested_category
        category_source = "document"
        import_category = None
        try:
            update_file_record(file_id, {"category": requested_category}, scope)
        except Exception as _e:
            log.warning("upload claims-fallback revert failed: %s", _e)

    # Surface the most common silent failure: a claims-looking PDF / Word doc
    # cannot be imported. Tell the uploader to send the spreadsheet instead so
    # the numbers actually move, instead of letting it sit inert.
    data_warning = None
    if file_type in ("pdf", "document") and effective_category not in DATA_IMPORT_CATEGORIES:
        _fn = (file.filename or "").lower()
        if any(k in _fn for k in ("daily", "claim", "worklist", "billed", "remit",
                                  "era", "deposit", "svd", "ledger", "aging", "payment")):
            data_warning = (
                f"\u201c{file.filename}\u201d looks like claims data. We tried to read it "
                f"automatically but couldn't extract claim rows from this "
                f"{(ext.lstrip('.') or 'file')}. For reliable totals, re-upload it as "
                "Excel/CSV (.xlsx or .csv) with claim columns."
            )

    try:
        size_kb = max(1, file_size // 1024)
        nice_name = file.filename or unique_name
        detail_bits = [f'"{nice_name}" ({size_kb} KB, {ext.lstrip(".") or "file"})',
                       f"category={effective_category}"]
        if row_count:
            detail_bits.append(f"{row_count} rows")
        if imported:
            detail_bits.append(f"auto-imported {imported} into {import_category}")
        if import_errors:
            detail_bits.append(f"{len(import_errors)} import warning(s)")
        if data_warning:
            detail_bits.append("NOT IMPORTED (needs Excel/CSV)")
        notify_activity(user["username"], "uploaded", "Documents", " · ".join(detail_bits))
        if imported and import_category:
            notify_bulk_activity(user["username"], "imported", import_category, imported,
                                 f'from upload "{nice_name}"')
    except Exception as _e:
        log.warning("upload notify failed: %s", _e)

    return {
        "id": file_id,
        "filename": unique_name,
        "original_name": file.filename,
        "requested_category": requested_category,
        "effective_category": effective_category,
        "category_source": category_source,
        "category_inference": infer_debug,
        "row_count": row_count,
        "imported": imported,
        "import_category": import_category,
        "import_errors": import_errors[:5],
        "data_warning": data_warning,
    }


# ─── Credentialing Excel Template ─────────────────────────────────────────────
import io, csv
from fastapi.responses import StreamingResponse

@router.get("/files/template/credentialing")
def download_credentialing_template():
    headers = ["Provider", "Payor", "Type", "Status", "Submitted", "Follow Up", "Approved", "Expiration", "Owner", "Notes", "Sub Profile"]
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(headers)
    writer.writerow(["John Smith", "Aetna", "Initial", "Submitted", "2026-02-20", "2026-02-27", "", "", "Jane Admin", "", "MHP"])
    output.seek(0)
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=credentialing_template.csv"}
    )


# ─── Client Report ────────────────────────────────────────────────────────────

def _build_section_data(conn, client_id, sub_profile=None, period=None):
    """Build claims/cred/enroll/edi/payments data for one filter set."""
    # Build parameterized filters
    sp_clause = ""
    sp_params = []
    if sub_profile:
        sp_clause = " AND sub_profile=?"
        sp_params = [sub_profile]

    date_clause = ""
    date_params = []
    if period == "mtd":
        from datetime import date as _d
        date_clause = " AND date(DOS) >= ?"
        date_params = [_d.today().replace(day=1).isoformat()]
    elif period == "ytd":
        from datetime import date as _d
        date_clause = " AND date(DOS) >= ?"
        date_params = [_d.today().replace(month=1, day=1).isoformat()]

    # Claims
    base = f"SELECT * FROM claims_master WHERE client_id=?{sp_clause}{date_clause}"
    claims = [dict(r) for r in conn.execute(base, [client_id] + sp_params + date_params).fetchall()]
    total_charged = sum(float(c.get("ChargeAmount") or 0) for c in claims)
    total_paid = sum(float(c.get("PaidAmount") or 0) for c in claims)
    total_balance = sum(float(c.get("BalanceRemaining") or 0) for c in claims)

    status_agg = {}
    for c in claims:
        st = c.get("ClaimStatus") or "Unknown"
        if st not in status_agg:
            status_agg[st] = {"count": 0, "charged": 0, "paid": 0}
        status_agg[st]["count"] += 1
        status_agg[st]["charged"] += float(c.get("ChargeAmount") or 0)
        status_agg[st]["paid"] += float(c.get("PaidAmount") or 0)
    by_status = [{"status": k, **v} for k, v in status_agg.items()]

    # Billed activity — charge dollars actually billed, grouped by BillDate window.
    # Computed independently of the DOS period filter so the admin always sees how
    # much was billed recently (today / yesterday / last 7 days / this month),
    # respecting the sub_profile scope.
    from datetime import date as _ba_date
    _today = business_today()
    _yesterday = _today.fromordinal(_today.toordinal() - 1)
    _week_start = _today.fromordinal(_today.toordinal() - 6)  # inclusive last 7 days
    _month_start = _today.replace(day=1)
    billing_activity = {
        "today": {"count": 0, "charged": 0.0},
        "yesterday": {"count": 0, "charged": 0.0},
        "this_week": {"count": 0, "charged": 0.0},
        "this_month": {"count": 0, "charged": 0.0},
        "all_time": {"count": 0, "charged": 0.0},
    }
    ba_sql = (f"SELECT BillDate, ChargeAmount FROM claims_master "
              f"WHERE client_id=?{sp_clause} AND COALESCE(BillDate,'')!=''")
    for r in conn.execute(ba_sql, [client_id] + sp_params).fetchall():
        bd = str(r["BillDate"] or "").strip()[:10]
        try:
            d = _ba_date.fromisoformat(bd)
        except Exception:
            continue
        amt = float(r["ChargeAmount"] or 0)
        # All-time billed since inception — every claim line with a Bill Date.
        billing_activity["all_time"]["count"] += 1
        billing_activity["all_time"]["charged"] += amt
        if d == _today:
            billing_activity["today"]["count"] += 1
            billing_activity["today"]["charged"] += amt
        if d == _yesterday:
            billing_activity["yesterday"]["count"] += 1
            billing_activity["yesterday"]["charged"] += amt
        if d >= _week_start:
            billing_activity["this_week"]["count"] += 1
            billing_activity["this_week"]["charged"] += amt
        if d >= _month_start:
            billing_activity["this_month"]["count"] += 1
            billing_activity["this_month"]["charged"] += amt
    for _k in billing_activity:
        billing_activity[_k]["charged"] = round(billing_activity[_k]["charged"], 2)

    denial_agg = {}
    for c in claims:
        dc = c.get("DenialCategory") or ""
        if dc:
            denial_agg[dc] = denial_agg.get(dc, 0) + 1
    top_denials = sorted([{"category": k, "count": v} for k, v in denial_agg.items()], key=lambda x: -x["count"])[:10]

    # Credentialing
    cred_base = f"SELECT * FROM credentialing WHERE client_id=?{sp_clause}"
    cred_rows = [dict(r) for r in conn.execute(cred_base, [client_id] + sp_params).fetchall()]
    cred_summary = {}
    for r in cred_rows:
        st = r.get("Status") or "Unknown"
        cred_summary[st] = cred_summary.get(st, 0) + 1
    cred_detail = [{"provider": r.get("ProviderName",""), "payor": r.get("Payor",""), "type": r.get("CredType",""),
                    "status": r.get("Status",""), "submitted": r.get("SubmittedDate",""), "approved": r.get("ApprovedDate",""),
                    "expires": r.get("ExpirationDate",""), "owner": r.get("Owner","")} for r in cred_rows]

    # Enrollment
    enr_base = f"SELECT * FROM enrollment WHERE client_id=?{sp_clause}"
    enr_rows = [dict(r) for r in conn.execute(enr_base, [client_id] + sp_params).fetchall()]
    enr_summary = {}
    for r in enr_rows:
        st = r.get("Status") or "Unknown"
        enr_summary[st] = enr_summary.get(st, 0) + 1
    enr_detail = [{"provider": r.get("ProviderName",""), "payor": r.get("Payor",""), "type": r.get("EnrollType",""),
                   "status": r.get("Status",""), "submitted": r.get("SubmittedDate",""),
                   "effective": r.get("EffectiveDate",""), "owner": r.get("Owner","")} for r in enr_rows]

    # EDI
    edi_base = f"SELECT * FROM edi_setup WHERE client_id=?{sp_clause}"
    edi_rows = [dict(r) for r in conn.execute(edi_base, [client_id] + sp_params).fetchall()]
    edi_summary = {}
    for r in edi_rows:
        st = r.get("EDIStatus") or "Unknown"
        edi_summary[st] = edi_summary.get(st, 0) + 1
    edi_detail = [{"provider": r.get("ProviderName",""), "payor": r.get("Payor",""), "payer_id": r.get("PayerID",""),
                   "edi": r.get("EDIStatus",""), "era": r.get("ERAStatus",""), "eft": r.get("EFTStatus",""),
                   "submitted": r.get("SubmittedDate",""), "go_live": r.get("GoLiveDate",""),
                   "owner": r.get("Owner","")} for r in edi_rows]

    # Payments
    pay_rows = conn.execute("SELECT COALESCE(SUM(PaymentAmount),0) as total, COUNT(*) as cnt FROM payments WHERE client_id=?", (client_id,)).fetchone()
    payments = {"total": float(pay_rows["total"]) if pay_rows else 0, "count": int(pay_rows["cnt"]) if pay_rows else 0}

    # ── Claim lifecycle buckets (Billed / Denied / Paid / Posted) ──
    # Computed identically to the dashboard (get_dashboard in client_db) so the
    # report's headline numbers always MATCH the dashboard the team monitors.
    # "Claims Out" == Billed (the superset of every claim line with a Bill Date).
    # Denied / Paid / Posted are overlapping status sub-views of that superset —
    # a denied claim is still a billed claim, it just shows its current state.
    # These ignore the DOS period filter (all-time, like the dashboard buckets)
    # but respect the same client_id + sub_profile scope.
    ROLLING_AR_START = "2026-06-18"  # first day the team began entering data
    bk_base = [client_id] + sp_params

    def _bk_one(sql, params):
        row = conn.execute(sql, params).fetchone()
        return row[0] if row else 0

    claim_buckets = {
        "billed": {
            "count": int(_bk_one(f"SELECT COUNT(*) FROM claims_master WHERE client_id=?{sp_clause}", bk_base)),
            "amount": round(float(_bk_one(f"SELECT COALESCE(SUM(ChargeAmount),0) FROM claims_master WHERE client_id=?{sp_clause}", bk_base)), 2),
        },
        # Denied = real denials only (status Denied/Appeals or a Denied Date),
        # NOT any claim carrying a DenialReason remark/adjustment code — those
        # appear on paid/adjusted claims too and inflated denials to the whole
        # billed book. Kept identical to the dashboard + production report.
        "denied": {
            "count": int(_bk_one(f"SELECT COUNT(*) FROM claims_master WHERE client_id=?{sp_clause} AND (ClaimStatus IN ('Denied','Appeals') OR TRIM(COALESCE(DeniedDate,''))!='')", bk_base)),
            "amount": round(float(_bk_one(f"SELECT COALESCE(SUM(ChargeAmount),0) FROM claims_master WHERE client_id=?{sp_clause} AND (ClaimStatus IN ('Denied','Appeals') OR TRIM(COALESCE(DeniedDate,''))!='')", bk_base)), 2),
        },
        "paid": {
            "count": int(_bk_one(f"SELECT COUNT(*) FROM claims_master WHERE client_id=?{sp_clause} AND COALESCE(PaidAmount,0)>0", bk_base)),
            "amount": round(float(_bk_one(f"SELECT COALESCE(SUM(PaidAmount),0) FROM claims_master WHERE client_id=?{sp_clause} AND COALESCE(PaidAmount,0)>0", bk_base)), 2),
        },
        "posted": {
            "count": int(_bk_one(f"SELECT COUNT(*) FROM payments WHERE client_id=?{sp_clause}", bk_base)),
            "amount": round(float(_bk_one(f"SELECT COALESCE(SUM(PaymentAmount),0) FROM payments WHERE client_id=?{sp_clause}", bk_base)), 2),
        },
    }
    rolling_ar = round(float(_bk_one(
        f"SELECT COALESCE(SUM(BalanceRemaining),0) FROM claims_master WHERE client_id=?{sp_clause} AND BillDate >= ?",
        bk_base + [ROLLING_AR_START])), 2)

    return {
        "claims": {"total": len(claims), "total_charged": round(total_charged,2), "total_paid": round(total_paid,2),
                    "total_balance": round(total_balance,2), "by_status": by_status, "top_denials": top_denials,
                    "billing_activity": billing_activity},
        "claim_buckets": claim_buckets,
        "rolling_ar": rolling_ar,
        "rolling_ar_start": ROLLING_AR_START,
        "credentialing": {"summary": [{"status":k,"count":v} for k,v in cred_summary.items()], "detail": cred_detail},
        "enrollment": {"summary": [{"status":k,"count":v} for k,v in enr_summary.items()], "detail": enr_detail},
        "edi": {"summary": [{"status":k,"count":v} for k,v in edi_summary.items()], "detail": edi_detail},
        "payments": payments,
    }


@router.get("/report/{client_id}")
def get_report(client_id: int, period: str = "all", sub_profile: Optional[str] = None,
               hub_session: Optional[str] = Cookie(None)):
    """Generate a comprehensive cross-section report for CSV / print, with sub-profile breakdowns."""
    user = _require_reporting_access(hub_session)
    from app.client_db import get_db
    from datetime import date, datetime

    conn = get_db()
    conn.row_factory = sqlite3.Row

    # Client info (including practice_type)
    client_row = conn.execute("SELECT company,contact_name,email,phone,practice_type FROM clients WHERE id=?", (client_id,)).fetchone()
    client_info = dict(client_row) if client_row else {}
    practice_type = client_info.get("practice_type", "") or ""

    try:
        # Build overall data — pass sub_profile and period as params (no f-string SQL)
        overall = _build_section_data(conn, client_id, sub_profile=sub_profile, period=period)

        # Build per-sub-profile breakdowns if MHP+OMT
        sub_profiles = {}
        if practice_type == "MHP+OMT" and not sub_profile:
            for sp_name in ["OMT", "MHP"]:
                sub_profiles[sp_name] = _build_section_data(conn, client_id, sub_profile=sp_name, period=period)
    finally:
        conn.close()

    result = {
        "generated_at": business_today_iso(),
        "period": period,
        "client": client_info,
        "practice_type": practice_type,
        **overall,
    }
    if sub_profiles:
        result["sub_profiles"] = sub_profiles

    return result


# ─── Direct Excel Import (per-section) ───────────────────────────────────────

@router.post("/import-excel")
async def import_excel(
    file: UploadFile = FastAPIFile(...),
    category: str = Form("Claims"),
    client_id: Optional[int] = Form(None),
    sub_profile: Optional[str] = Form(None),
    hub_session: Optional[str] = Cookie(None),
):
    """Import an Excel/CSV file directly into a data table (Claims, Credentialing, Enrollment, EDI).
    Also saves a copy of the file in Documents under the appropriate category."""
    user = _require_user(hub_session)
    scope = client_id if client_id is not None else (_client_scope(user) if _client_scope(user) is not None else _single_client_account_or(user["id"]))

    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in (".xlsx", ".xls", ".csv", ".ods", ".odf"):
        raise HTTPException(400, "Only .xlsx, .xls, .csv, .ods, .odf files supported for import")

    content = await file.read()

    # ── Save a copy of the file in Documents ──
    unique_name = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join(UPLOAD_DIR, unique_name)
    with open(dest, "wb") as f:
        f.write(content)
    file_size = len(content)
    row_count = 0
    try:
        import csv as _csv, io as _io
        if ext == ".csv":
            reader = _csv.reader(_io.StringIO(content.decode("utf-8", errors="replace")))
            row_count = max(0, sum(1 for _ in reader) - 1)
        else:
            import openpyxl
            wb = openpyxl.load_workbook(_io.BytesIO(content), read_only=True, data_only=True)
            row_count = max(0, sum(ws.max_row - 1 for ws in wb.worksheets if ws.max_row))
            wb.close()
    except Exception:
        pass
    file_id = add_file(
        client_id=scope, filename=unique_name, original_name=file.filename or "file",
        file_type="excel", file_size=file_size, category=category,
        description=f"{category} import — {file.filename}",
        row_count=row_count, uploaded_by=user["username"],
    )

    imported = 0
    errors = []

    try:
        if category == "Claims":
            imported, errors = _import_claims_from_excel(
                content, ext, scope, uploaded_by=user["username"],
                default_sub_profile=(sub_profile or "").strip())
        elif category == "Credentialing":
            imported, errors = _import_credentialing_from_excel(content, ext, scope, uploaded_by=user["username"])
        elif category == "Enrollment":
            imported, errors = _import_enrollment_from_excel(content, ext, scope, uploaded_by=user["username"])
        elif category == "EDI":
            imported, errors = _import_edi_from_excel(content, ext, scope, uploaded_by=user["username"])
        else:
            raise HTTPException(400, f"Unknown category: {category}")
    except HTTPException:
        raise
    except Exception as e:
        errors = [str(e)]

    # Notify admin of team imports
    if imported > 0:
        notify_bulk_activity(user["username"], "imported", category, imported,
                             f"File: {file.filename}")

    return {
        "category": category,
        "imported": imported,
        "errors": errors[:10],
        "original_name": file.filename,
        "file_id": file_id,
    }


def _parse_excel_rows(content: bytes, ext: str, combine_sheets: bool = True):
    """Parse Excel/CSV bytes into list of dict rows with smart header detection.
    If combine_sheets=True and multiple sheets share the same header structure,
    rows from all matching sheets are combined (useful for multi-tab claim files).
    Supports .xlsx (openpyxl), .xls (xlrd), .ods/.odf (OpenDocument), and .csv."""
    import csv, io
    rows = []
    if ext == ".csv":
        # Decode tolerantly: strip a BOM if present, fall back through encodings
        # so files exported from Excel/Windows/Mac all parse.
        text = None
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                text = content.decode(enc)
                break
            except Exception:
                continue
        if text is None:
            text = content.decode("utf-8", errors="replace")

        # Detect the delimiter — exports aren't always comma-separated
        # (semicolon in many locales, tab/pipe from some systems).
        sample = text[:8192]
        delim = ","
        try:
            delim = csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
        except Exception:
            first_line = next((ln for ln in sample.splitlines() if ln.strip()), "")
            if first_line:
                delim = max(",;\t|", key=first_line.count)

        raw_rows = list(csv.reader(io.StringIO(text), delimiter=delim))
        # Smart header detection: the header is the most "label-like" row within
        # the first 10 (skips title/blank rows that sit above real headers).
        header_idx = 0
        best_score = -1
        for idx, r in enumerate(raw_rows[:10]):
            non_empty = sum(1 for c in r if c is not None and str(c).strip())
            text_cells = sum(1 for c in r if c is not None and str(c).strip()
                             and not str(c).strip().replace(",", "").replace(".", "")
                             .replace("$", "").replace("-", "").isdigit())
            score = text_cells * 2 + non_empty
            if non_empty >= 2 and score > best_score:
                best_score = score
                header_idx = idx
        if raw_rows:
            headers = [str(c).strip() for c in raw_rows[header_idx]]
            for r in raw_rows[header_idx + 1:]:
                if any(c is not None and str(c).strip() for c in r):
                    rows.append(dict(zip(headers, r)))
    elif ext == ".xls":
        # Legacy Excel (BIFF) — use xlrd
        import xlrd
        wb = xlrd.open_workbook(file_contents=content)
        sheet_results = []
        for sidx in range(wb.nsheets):
            ws = wb.sheet_by_index(sidx)
            if ws.nrows < 2:
                continue
            # Smart header detection: best row in first 10 rows
            header_row_idx = 0
            best_score = 0
            for ri in range(min(10, ws.nrows)):
                row_vals = [ws.cell_value(ri, ci) for ci in range(ws.ncols)]
                non_empty = sum(1 for v in row_vals if v is not None and str(v).strip())
                text = sum(1 for v in row_vals if isinstance(v, str) and str(v).strip())
                score = text * 2 + non_empty
                if score > best_score and non_empty >= 3:
                    best_score = score
                    header_row_idx = ri
            hdrs = [str(ws.cell_value(header_row_idx, ci)).strip() for ci in range(ws.ncols)]
            valid_cols = [i for i, h in enumerate(hdrs) if h]
            if len(valid_cols) < 2:
                continue
            sheet_rows = []
            for ri in range(header_row_idx + 1, ws.nrows):
                row_vals = []
                for ci in range(ws.ncols):
                    cell = ws.cell(ri, ci)
                    # Convert xlrd date serial to Python datetime
                    if cell.ctype == xlrd.XL_CELL_DATE:
                        try:
                            import datetime as _dt
                            dt = xlrd.xldate_as_datetime(cell.value, wb.datemode)
                            row_vals.append(dt)
                        except Exception:
                            row_vals.append(cell.value)
                    else:
                        row_vals.append(cell.value)
                if any(v is not None and str(v).strip() for v in row_vals):
                    sheet_rows.append(dict(zip(hdrs, row_vals)))
            if sheet_rows:
                hdr_key = tuple(sorted(h.lower() for h in hdrs if h))
                sheet_results.append((hdr_key, hdrs, sheet_rows))
        if sheet_results:
            if combine_sheets:
                from collections import defaultdict
                groups = defaultdict(list)
                for hdr_key, hdrs, srows in sheet_results:
                    groups[hdr_key].extend(srows)
                rows = max(groups.values(), key=len)
            else:
                rows = max(sheet_results, key=lambda x: len(x[2]))[2]
    elif ext in (".ods", ".odf"):
        # OpenDocument spreadsheets — parse content.xml from zip package.
        import xml.etree.ElementTree as _et
        import zipfile as _zipfile

        ns = {
            "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
            "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
            "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
        }
        a_row_rep = "{urn:oasis:names:tc:opendocument:xmlns:table:1.0}number-rows-repeated"
        a_col_rep = "{urn:oasis:names:tc:opendocument:xmlns:table:1.0}number-columns-repeated"
        a_str = "{urn:oasis:names:tc:opendocument:xmlns:office:1.0}string-value"
        a_date = "{urn:oasis:names:tc:opendocument:xmlns:office:1.0}date-value"
        a_val = "{urn:oasis:names:tc:opendocument:xmlns:office:1.0}value"

        try:
            with _zipfile.ZipFile(io.BytesIO(content)) as zf:
                xml_bytes = zf.read("content.xml")
        except Exception as exc:
            raise ValueError(f"Cannot read OpenDocument file: {exc}") from exc

        root = _et.fromstring(xml_bytes)
        sheet_results = []

        for table in root.findall(".//table:table", ns):
            all_sheet_rows = []
            for tr in table.findall("table:table-row", ns):
                row_repeat = int(tr.attrib.get(a_row_rep, "1") or "1")
                row_vals = []
                for cell in tr:
                    if not (cell.tag.endswith("table-cell") or cell.tag.endswith("covered-table-cell")):
                        continue
                    col_repeat = int(cell.attrib.get(a_col_rep, "1") or "1")
                    parts = []
                    for p in cell.findall(".//text:p", ns):
                        txt = "".join(p.itertext()).strip()
                        if txt:
                            parts.append(txt)
                    v = " ".join(parts).strip()
                    if not v:
                        v = (cell.attrib.get(a_str) or cell.attrib.get(a_date) or cell.attrib.get(a_val) or "")
                    row_vals.extend([v] * max(1, col_repeat))
                if not row_vals:
                    continue
                for _ in range(max(1, row_repeat)):
                    all_sheet_rows.append(tuple(row_vals))

            if not all_sheet_rows:
                continue

            header_row_idx = 0
            best_header_score = 0
            for idx, row in enumerate(all_sheet_rows[:10]):
                non_empty = sum(1 for c in row if c is not None and str(c).strip())
                text_cells = sum(1 for c in row if c is not None and isinstance(c, str) and len(str(c).strip()) > 0)
                score = text_cells * 2 + non_empty
                if score > best_header_score and non_empty >= 3:
                    best_header_score = score
                    header_row_idx = idx

            if header_row_idx < len(all_sheet_rows):
                sheet_headers = [str(c).strip() if c is not None else "" for c in all_sheet_rows[header_row_idx]]
                valid_cols = [i for i, h in enumerate(sheet_headers) if h]
                if len(valid_cols) < 2:
                    continue
                sheet_rows = []
                for row in all_sheet_rows[header_row_idx + 1:]:
                    if any(c is not None and str(c).strip() for c in row):
                        sheet_rows.append(dict(zip(sheet_headers, row)))
                if sheet_rows:
                    hdr_key = tuple(sorted(h.lower() for h in sheet_headers if h))
                    sheet_results.append((hdr_key, sheet_headers, sheet_rows))

        if sheet_results:
            if combine_sheets:
                from collections import defaultdict
                groups = defaultdict(list)
                for hdr_key, hdrs, srows in sheet_results:
                    groups[hdr_key].extend(srows)
                rows = max(groups.values(), key=len)
            else:
                best = max(sheet_results, key=lambda x: len(x[2]))
                rows = best[2]
    else:
        # .xlsx — use openpyxl
        import openpyxl
        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        except Exception as exc:
            raise ValueError(f"Cannot read Excel file: {exc}. If this is a .xls file, rename to .xls extension.") from exc
        # Collect parsed rows per sheet with their headers
        sheet_results = []  # list of (headers_tuple, rows_list)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_sheet_rows = []
            for row in ws.iter_rows(values_only=True):
                all_sheet_rows.append(row)
            if not all_sheet_rows:
                continue
            # Smart header detection: find the best header row
            header_row_idx = 0
            best_header_score = 0
            for idx, row in enumerate(all_sheet_rows[:10]):
                if not row:
                    continue
                non_empty = sum(1 for c in row if c is not None and str(c).strip())
                text_cells = sum(1 for c in row if c is not None and isinstance(c, str) and len(str(c).strip()) > 0)
                # Require at least 3 non-empty cells to be a real header row
                score = text_cells * 2 + non_empty
                if score > best_header_score and non_empty >= 3:
                    best_header_score = score
                    header_row_idx = idx
            if header_row_idx < len(all_sheet_rows):
                sheet_headers = [str(c).strip() if c else "" for c in all_sheet_rows[header_row_idx]]
                # Filter out empty header columns
                valid_cols = [i for i, h in enumerate(sheet_headers) if h]
                if len(valid_cols) < 2:
                    continue
                sheet_rows = []
                for row in all_sheet_rows[header_row_idx + 1:]:
                    if any(c is not None and str(c).strip() for c in row):
                        sheet_rows.append(dict(zip(sheet_headers, row)))
                if sheet_rows:
                    # Use a normalized header key (sorted, lowered) for grouping
                    hdr_key = tuple(sorted(h.lower() for h in sheet_headers if h))
                    sheet_results.append((hdr_key, sheet_headers, sheet_rows))

        if sheet_results:
            if combine_sheets:
                # Group sheets by header structure and combine sheets with same headers
                from collections import defaultdict
                groups = defaultdict(list)
                for hdr_key, hdrs, srows in sheet_results:
                    groups[hdr_key].extend(srows)
                # Pick the group with the most total rows
                best_group = max(groups.values(), key=len)
                rows = best_group
            else:
                # Pick single sheet with most rows
                best = max(sheet_results, key=lambda x: len(x[2]))
                rows = best[2]
        wb.close()
    return rows


# ─── Hardcoded format templates ("the glasses") ───────────────────────────────
# Some recurring claim exports cannot be read by column-name matching because they
# have NO header row (the first row is data), or they are batch SUMMARIES rather
# than per-claim detail, or they are multi-sheet workbooks where the meaningful
# data lives on a non-first sheet in a fixed positional layout. Generic header
# matching physically cannot recognize these, so the daily uploads never moved the
# totals. These templates pin the EXACT shape of each known recurring file and emit
# rows already labeled with synthetic headers that CLAIMS_COLUMN_MAP understands
# (plus an injected status), so the rest of the import pipeline works unchanged.
#
# SAFETY: each template only fires when its fingerprint matches precisely. If the
# layout ever changes, the template simply does not match and the file falls back
# to the generic parser (i.e. behaves exactly as before — never wrong numbers).

_SVD_CLAIMNO_RE = re.compile(r"^[A-Za-z]{2,5}\d+-\d+$")   # e.g. SVD9322-52429
_CARC_RE = re.compile(r"^[A-Za-z]{1,3}\d{1,4}([, ]+[A-Za-z]{1,3}\d{1,4})*$")  # PR49 / CO22, CO50


def _cell_iso_date(v):
    """Best-effort ISO date string from a cell (datetime or 'YYYY-MM-DD ...')."""
    if v is None:
        return ""
    if isinstance(v, (datetime, date)):
        return v.strftime("%Y-%m-%d")
    s = str(v).strip()
    if not s:
        return ""
    if " " in s and len(s) > 10:
        s = s.split(" ")[0]
    return s


def _load_xlsx_sheets(content: bytes):
    """Return [(sheet_name, [row_tuples...]), ...] for an .xlsx workbook, or []."""
    import openpyxl
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception:
        return []
    out = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = [r for r in ws.iter_rows(values_only=True)]
        out.append((sn, rows))
    wb.close()
    return out


def _tpl_svd_denials(sheets):
    """SVD denial / AR report: NO header row, fixed positional layout, usually a
    MULTI-SHEET workbook. The export stacks TWO record shapes that both describe
    billed claims, and the A/R aging frequently spills onto a SECOND sheet:
      • Denial worklist  — col0 is the claim number (SVD9322-52429), col4 a CARC
        code (PR49 / CO22), col16 the charge. These were billed, then denied, and
        are reworked — so they count as Billed AND carry a denial.
      • A/R aging detail  — col0 is a numeric id, col26 is the claim number
        (SVD3166-47529), col28 the charge, col29 the allowed. Billed claims sitting
        in A/R. Both shapes are emitted as billed claims keyed on the real claim #.
    EVERY submitted / reworked line is emitted — nothing is collapsed. A claim that
    was worked in the denial queue AND tracked in A/R (often across two sheets and
    multiple aging snapshots) represents real, repeated billing activity, so each
    occurrence keeps its own row: the first sighting carries the clean claim #, and
    later sightings become claim#2 / claim#3 … (deterministic in file order, so a
    re-upload updates in place and never drifts). ALL sheets are scanned, so a
    workbook whose A/R lives on a later sheet is fully read instead of abandoned."""
    def _col(r, i):
        return r[i] if (r and len(r) > i and r[i] is not None) else ""
    all_out = []
    seq: dict = {}
    matched = False
    for sn, rows in sheets:
        if len(rows) < 2:
            continue
        # Fingerprint each sheet independently. A sheet qualifies if it carries the
        # denial worklist shape (claim # in col0 + CARC in col4) OR the A/R aging
        # shape (claim # in col26). The A/R-only test is what lets a second sheet
        # holding nothing but aging detail still be recognized and read.
        sample = [r for r in rows[:60] if r and any(c is not None and str(c).strip() for c in r)]
        if len(sample) < 3:
            continue
        key_hits = sum(1 for r in sample if _SVD_CLAIMNO_RE.match(str(_col(r, 0)).strip()))
        carc_hits = sum(1 for r in sample if _CARC_RE.match(str(_col(r, 4)).strip()))
        ar_hits = sum(1 for r in sample if _SVD_CLAIMNO_RE.match(str(_col(r, 26)).strip()))
        if not ((key_hits >= 2 and carc_hits >= 2) or ar_hits >= 2):
            continue
        matched = True
        for r in rows:
            if not r:
                continue
            c0 = str(_col(r, 0)).strip()
            c26 = str(_col(r, 26)).strip()
            base = ""
            rec = None
            if _SVD_CLAIMNO_RE.match(c0):
                # Denial worklist line — billed claim that was denied.
                first = str(_col(r, 10)).strip()
                last = str(_col(r, 11)).strip()
                name = (", ".join(p for p in (last, first) if p)).strip(", ")
                codes = str(_col(r, 4)).strip()
                reason = str(_col(r, 5)).strip()
                base = c0
                rec = {
                    "Claim Status": "Denied",
                    "Payor": str(_col(r, 1)).strip(),
                    "DOS": _cell_iso_date(_col(r, 2)),
                    "CPT Code": str(_col(r, 3)).strip(),
                    "Denial Reason": (codes + (" - " if codes and reason else "") + reason).strip(),
                    "Patient Name": name,
                    "Bill Date": _cell_iso_date(_col(r, 13)),
                    "Denied Date": _cell_iso_date(_col(r, 14)),
                    "Owner": str(_col(r, 15)).strip(),
                    "Charge Amount": _col(r, 16),
                    "Paid Amount": _col(r, 17),
                }
            elif _SVD_CLAIMNO_RE.match(c26):
                # A/R aging detail line — billed claim still outstanding.
                first = str(_col(r, 2)).strip()
                last = str(_col(r, 3)).strip()
                name = (", ".join(p for p in (last, first) if p)).strip(", ")
                bdate = _cell_iso_date(_col(r, 24)) or _cell_iso_date(_col(r, 6))
                base = c26
                rec = {
                    "Claim Status": "A/R Follow-Up",
                    "Payor": str(_col(r, 9)).strip(),
                    "DOS": _cell_iso_date(_col(r, 6)),
                    "CPT Code": str(_col(r, 15)).strip(),
                    "Patient Name": name,
                    "Bill Date": bdate,
                    "Charge Amount": _col(r, 28),
                    "Allowed Amount": _col(r, 29),
                    "Paid Amount": _col(r, 30),
                }
            if rec is None:
                continue
            # Unique key per submitted/reworked line so the gross billed total is
            # preserved instead of being deduped away. First sighting keeps the
            # clean claim #; repeats become claim#2, claim#3 … (stable file order).
            n = seq.get(base, 0) + 1
            seq[base] = n
            rec["Claim #"] = base if n == 1 else f"{base}#{n}"
            all_out.append(rec)
    return all_out if (matched and all_out) else None


def _tpl_lims_payments(sheets):
    """LIMS payments/ERA sheet: deposits posted. Header carries BATCH # / DEPOSIT
    DATE / PAYER NAME / AMOUNT / EFT NUMBER. Each row is a posted payment."""
    for sn, rows in sheets:
        if not rows:
            continue
        cells = [_norm_key(str(c)) if c is not None else "" for c in rows[0]]
        has_amount = any(c == "amount" for c in cells)
        has_payer = any("payer" in c for c in cells)
        has_eft = any("eft" in c for c in cells)
        has_deposit = any("deposit" in c for c in cells)
        if not (has_amount and (has_eft or has_deposit) and has_payer):
            continue
        idx = {}
        for i, c in enumerate(cells):
            if c == "amount":
                idx["amount"] = i
            elif "payer" in c:
                idx["payer"] = i
            elif "eft" in c:
                idx["eft"] = i
            elif "deposit date" in c or (("deposit" in c) and "date" in c):
                idx["date"] = i
        out = []
        for r in rows[1:]:
            if not r or not any(c is not None and str(c).strip() for c in r):
                continue
            try:
                amt = float(str(r[idx["amount"]]).replace("$", "").replace(",", "").strip())
            except Exception:
                amt = 0.0
            eft = str(r[idx["eft"]]).split(".")[0].strip() if "eft" in idx and len(r) > idx["eft"] and r[idx["eft"]] is not None else ""
            payer = str(r[idx["payer"]]).strip() if "payer" in idx and len(r) > idx["payer"] and r[idx["payer"]] is not None else ""
            pdate = _cell_iso_date(r[idx["date"]]) if "date" in idx and len(r) > idx["date"] else ""
            key = f"PMT-{eft}" if eft else f"PMT-{payer}-{pdate}-{amt}"
            out.append({
                "Claim #": key,
                "Claim Status": "Paid",
                "Payor": payer,
                "Paid Amount": amt,
                "Paid Date": pdate,
                "Bill Date": pdate,
            })
        if out:
            return out
    return None


# Priority order — the first template whose fingerprint matches wins for a given
# workbook, so a file contributes to exactly one bucket and never double counts.
# The SVD DAILY batch log (DATE / BATCH # / NUMBER OF CLAIMS / TOTAL BILLED) is
# Melissa & Susan's COLLECTIVE daily clearinghouse-transmission report -- the team
# confirmed it is the same ~$91K of submissions already itemized per-claim (with the
# real charge amounts) in Susan's "Claim Sent" register. The batch log only values
# each claim at a flat ~$141 estimate; it is a transmission VIEW of work already
# counted, NOT separate billing. It is therefore intentionally NOT in _CLAIM_TEMPLATES
# -- importing it would double-count Melissa + Susan to ~$176K instead of their true
# combined ~$91K. It is kept as a document and guarded by _is_batch_transmission_log;
# Melissa's distinct outputs are eligibility (LIMS) and payments (ERA'S STEDI).
_CLAIM_TEMPLATES = (
    ("svd_denials", _tpl_svd_denials),
    ("lims_payments", _tpl_lims_payments),
)


def _extract_templated_claim_rows(content: bytes, ext: str):
    """If the file matches a known recurring SV format, return (template_name,
    labeled_rows); otherwise None. Only .xlsx workbooks are templated."""
    if ext not in (".xlsx",):
        return None
    sheets = _load_xlsx_sheets(content)
    if not sheets:
        return None
    for name, fn in _CLAIM_TEMPLATES:
        try:
            rows = fn(sheets)
        except Exception:
            rows = None
        if rows:
            return name, rows
    return None


# Header tokens that mark a *claims* table (patient / charge / CPT / payor /
# claim no.). Used when reading PDFs and Word docs so a claims worklist exported
# to PDF or Word is recognized and read into claims — not just spreadsheet exports.
_CLAIM_DOC_HEADER_KEYWORDS = (
    "patient", "member", "subscriber", "account", "mrn", "dos",
    "date of service", "service date", "from date", "claim", "claim no",
    "claim number", "cpt", "procedure", "code", "charge", "charges", "billed",
    "amount", "balance", "allowed", "paid", "payment", "payor", "payer",
    "insurance", "carrier", "provider", "rendering", "npi", "status",
    "denial", "adjustment", "units", "modifier", "icd", "diagnosis", "dx",
)


def _load_claim_rows(content: bytes, ext: str):
    """Claim ingestion entry point — the UNIVERSAL reader.

    Reads claim rows out of whatever format the team actually sends: Excel/CSV
    exports, but also PDF daily worklists and Word documents. Spreadsheets try
    the hardcoded format templates first (headerless / batch-summary / multi-sheet
    SV exports), then the generic smart-header parser. PDFs and Word docs run
    through the table extractors with claim-aware header detection. Returns
    (rows, template_name|None). Never raises for PDF/Word — an unreadable doc
    yields [] so the caller treats it as a plain document, not an error."""
    e = (ext or "").lower()
    if e == ".pdf":
        try:
            return _parse_pdf_rows(content, header_keywords=_CLAIM_DOC_HEADER_KEYWORDS), None
        except Exception:
            return [], None
    if e in (".doc", ".docx"):
        try:
            return _parse_docx_rows(content, header_keywords=_CLAIM_DOC_HEADER_KEYWORDS), None
        except Exception:
            return [], None
    tpl = _extract_templated_claim_rows(content, ext)
    if tpl:
        return tpl[1], tpl[0]
    return _parse_excel_rows(content, ext), None


def _parse_pdf_rows(content: bytes, header_keywords: tuple = None) -> list[dict]:
    """Parse table-like PDF content into list[dict] using detected headers.

    Intended for controlled imports where PDFs are exported in tabular layout.
    `header_keywords` overrides the tokens used to locate the header row, so the
    same extractor reads production-log PDFs and claims-worklist PDFs alike.
    """
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise ValueError("PDF parsing dependency missing. Install 'pypdf'.") from exc

    import io as _io

    reader = PdfReader(_io.BytesIO(content))
    lines: list[str] = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        for raw in txt.splitlines():
            s = re.sub(r"\s+", " ", str(raw or "")).strip()
            if s:
                lines.append(s)

    if not lines:
        return []

    def _split_line(line: str) -> list[str]:
        if "|" in line:
            return [p.strip() for p in line.split("|") if p.strip()]
        if "\t" in line:
            return [p.strip() for p in line.split("\t") if p.strip()]
        parts = [p.strip() for p in re.split(r"\s{2,}", line) if p.strip()]
        if len(parts) >= 3:
            return parts
        return [line.strip()]

    header_idx = -1
    headers: list[str] = []
    if header_keywords is None:
        header_keywords = (
            "date", "work date", "task", "description", "hours", "qty", "quantity", "category", "user", "notes",
        )
    for i, line in enumerate(lines[:60]):
        parts = _split_line(line)
        if len(parts) < 3:
            continue
        joined = " ".join(parts).lower()
        hits = sum(1 for k in header_keywords if k in joined)
        if hits >= 2:
            header_idx = i
            headers = parts
            break

    if header_idx < 0:
        raise ValueError(
            "Could not detect a table header in PDF. "
            "Expected columns like Date, Task/Description, Category, Qty, Hours, Notes."
        )

    ncols = len(headers)
    rows: list[dict] = []
    for line in lines[header_idx + 1:]:
        if re.fullmatch(r"[-_=| ]{3,}", line):
            continue
        parts = _split_line(line)
        if len(parts) < 2:
            continue
        if len(parts) < ncols:
            parts = parts + [""] * (ncols - len(parts))
        elif len(parts) > ncols:
            parts = parts[: ncols - 1] + [" ".join(parts[ncols - 1 :])]
        row = dict(zip(headers, parts))
        if any(str(v).strip() for v in row.values()):
            rows.append(row)
    return rows


def _parse_docx_rows(content: bytes, header_keywords: tuple = None) -> list[dict]:
    """Parse a Word (.docx) document into list[dict].

    Prefers the document's first real table (header row + data rows). If the
    document has no usable table, falls back to reading paragraph text and
    detecting a tabular header the same way the PDF parser does.
    """
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency guard
        raise ValueError("Word parsing dependency missing. Install 'python-docx'.") from exc

    import io as _io

    doc = Document(_io.BytesIO(content))

    # 1) Prefer a real table with a header row + at least one data row.
    for table in doc.tables:
        trows = table.rows
        if len(trows) < 2:
            continue
        headers = [(_clean_val(c.text) or "").strip() for c in trows[0].cells]
        if sum(1 for h in headers if h) < 2:
            continue
        ncols = len(headers)
        out: list[dict] = []
        for tr in trows[1:]:
            cells = [(_clean_val(c.text) or "").strip() for c in tr.cells]
            if len(cells) < ncols:
                cells = cells + [""] * (ncols - len(cells))
            elif len(cells) > ncols:
                cells = cells[:ncols]
            row = dict(zip(headers, cells))
            if any(str(v).strip() for v in row.values()):
                out.append(row)
        if out:
            return out

    # 2) Fall back to paragraph text using the PDF line/header heuristics.
    text = "\n".join(p.text for p in doc.paragraphs)
    fake_pdf_lines = [re.sub(r"\s+", " ", ln).strip() for ln in text.splitlines() if ln.strip()]
    if not fake_pdf_lines:
        return []

    def _split_line(line: str) -> list[str]:
        if "|" in line:
            return [p.strip() for p in line.split("|") if p.strip()]
        if "\t" in line:
            return [p.strip() for p in line.split("\t") if p.strip()]
        parts = [p.strip() for p in re.split(r"\s{2,}", line) if p.strip()]
        return parts if len(parts) >= 3 else [line.strip()]

    if header_keywords is None:
        header_keywords = (
            "date", "work date", "task", "description", "hours", "qty", "quantity",
            "category", "user", "notes",
        )
    header_idx, headers = -1, []
    for i, line in enumerate(fake_pdf_lines[:60]):
        parts = _split_line(line)
        if len(parts) < 3:
            continue
        if sum(1 for k in header_keywords if k in " ".join(parts).lower()) >= 2:
            header_idx, headers = i, parts
            break
    if header_idx < 0:
        raise ValueError(
            "Could not detect a table in the Word document. "
            "Use a table with columns like Date, Task/Description, Category, Qty, Hours, Notes."
        )
    ncols = len(headers)
    rows: list[dict] = []
    for line in fake_pdf_lines[header_idx + 1:]:
        parts = _split_line(line)
        if len(parts) < 2:
            continue
        if len(parts) < ncols:
            parts = parts + [""] * (ncols - len(parts))
        elif len(parts) > ncols:
            parts = parts[: ncols - 1] + [" ".join(parts[ncols - 1:])]
        row = dict(zip(headers, parts))
        if any(str(v).strip() for v in row.values()):
            rows.append(row)
    return rows


def _norm_key(k):
    """Normalize an Excel header: lowercase, strip, collapse whitespace, remove _/-/# chars."""
    s = (k or "").strip().lower()
    s = s.replace("_", " ").replace("-", " ").replace("#", "").replace(".", " ")
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def _fuzzy_match_column(header, col_map):
    """Try exact match first, then substring/contains matching as fallback."""
    norm = _norm_key(header)
    if norm in col_map:
        return col_map[norm]
    # Try if any map key is contained in the header
    for map_key, db_col in col_map.items():
        if len(map_key) >= 3 and map_key in norm:
            return db_col
    return None


def _clean_val(val):
    """Convert a cell value to a clean string. Strips time from datetime objects."""
    if val is None:
        return ""
    if isinstance(val, datetime):
        return val.strftime("%Y-%m-%d")
    if isinstance(val, date):
        return val.isoformat()
    s = str(val).strip()
    # Strip trailing 00:00:00 from date strings like "2026-02-20 00:00:00"
    if len(s) > 10 and s[10:].strip() in ('00:00:00', '0:00:00', '00:00', '0:00'):
        s = s[:10]
    return s


def _import_credentialing_from_excel(content: bytes, ext: str, client_id: int, uploaded_by: str = ""):
    from app.client_db import get_db as _get_db
    from datetime import datetime as _dt_now

    # ── Status normalization — maps Excel values to dropdown filter values ──
    CRED_STATUS_NORMALIZE = {
        # Not Started
        "not started": "Not Started", "new": "Not Started", "none": "Not Started",
        "pending start": "Not Started", "not begun": "Not Started", "n/a": "Not Started",
        "to do": "Not Started", "open": "Not Started",
        # In Progress
        "in progress": "In Progress", "in-progress": "In Progress", "pending": "In Progress",
        "processing": "In Progress", "working": "In Progress", "active": "In Progress",
        "in process": "In Progress", "in review": "In Progress", "under review": "In Progress",
        "need action": "In Progress", "needs action": "In Progress", "action needed": "In Progress",
        "follow up": "In Progress", "follow-up": "In Progress", "followup": "In Progress",
        "waiting": "In Progress", "awaiting": "In Progress", "in queue": "In Progress",
        # Submitted
        "submitted": "Submitted", "sent": "Submitted", "filed": "Submitted",
        "application submitted": "Submitted", "app submitted": "Submitted",
        "mailed": "Submitted", "faxed": "Submitted", "uploaded": "Submitted",
        "received": "Submitted", "acknowledged": "Submitted",
        # Approved
        "approved": "Approved", "completed": "Approved", "credentialed": "Approved",
        "active - approved": "Approved", "accepted": "Approved", "enrolled": "Approved",
        "effective": "Approved", "live": "Approved", "done": "Approved",
        "granted": "Approved", "passed": "Approved",
        # Denied
        "denied": "Denied", "rejected": "Denied", "declined": "Denied",
        "not approved": "Denied", "failed": "Denied", "terminated": "Denied",
        "closed": "Denied", "cancelled": "Denied", "canceled": "Denied",
        # Expired
        "expired": "Expired", "lapsed": "Expired", "renewal needed": "Expired",
        "expiring": "Expired", "past due": "Expired", "overdue": "Expired",
        "renewal": "Expired", "recredentialing": "Expired",
    }

    VALID_CRED_STATUSES = {"Not Started", "In Progress", "Submitted", "Approved", "Denied", "Expired"}

    def _normalize_cred_status(raw):
        if not raw:
            return "Not Started"
        s = str(raw).strip()
        key = s.lower()
        if key in CRED_STATUS_NORMALIZE:
            return CRED_STATUS_NORMALIZE[key]
        for map_key, normalized in CRED_STATUS_NORMALIZE.items():
            if len(map_key) >= 4 and map_key in key:
                return normalized
        # If raw already matches a valid status (case-insensitive), use it
        for vs in VALID_CRED_STATUSES:
            if vs.lower() == key:
                return vs
        return "In Progress"  # Safe default for unrecognized statuses

    COL_MAP = {
        # Provider
        "provider": "ProviderName", "providername": "ProviderName", "provider name": "ProviderName",
        "rendering provider": "ProviderName", "doctor": "ProviderName", "physician": "ProviderName",
        "doctor name": "ProviderName", "physician name": "ProviderName", "practitioner": "ProviderName",
        "rendering": "ProviderName", "servicing provider": "ProviderName",
        "name": "ProviderName", "provider/doctor": "ProviderName",
        # Payor
        "payor": "Payor", "payer": "Payor", "insurance": "Payor",
        "insurance name": "Payor", "insurance company": "Payor", "plan": "Payor",
        "plan name": "Payor", "payer name": "Payor", "carrier": "Payor",
        "insurance plan": "Payor", "health plan": "Payor", "ins": "Payor",
        "primary insurance": "Payor", "primary payor": "Payor", "primary payer": "Payor",
        "ins name": "Payor", "ins company": "Payor",
        # Type / Subtask
        "type": "CredType", "credtype": "CredType", "cred type": "CredType",
        "credential type": "CredType", "credentialing type": "CredType",
        "application type": "CredType", "app type": "CredType",
        "subtask": "CredType", "sub task": "CredType", "task": "CredType",
        "cred subtask": "CredType", "task type": "CredType",
        # Status
        "status": "Status", "cred status": "Status", "credentialing status": "Status",
        "app status": "Status", "application status": "Status", "current status": "Status",
        # Dates
        "submitted": "SubmittedDate", "submitted date": "SubmittedDate", "submitteddate": "SubmittedDate",
        "date submitted": "SubmittedDate", "submission date": "SubmittedDate", "app submitted": "SubmittedDate",
        "application submitted": "SubmittedDate", "submit date": "SubmittedDate",
        "date": "SubmittedDate", "start date": "SubmittedDate",
        "date started": "SubmittedDate", "started": "SubmittedDate", "start": "SubmittedDate",
        "follow up": "FollowUpDate", "followupdate": "FollowUpDate", "follow up date": "FollowUpDate",
        "followup": "FollowUpDate", "followup date": "FollowUpDate", "next follow up": "FollowUpDate",
        "fu date": "FollowUpDate", "f/u date": "FollowUpDate", "f/u": "FollowUpDate",
        "next action date": "FollowUpDate", "due date": "FollowUpDate",
        "last follow up": "FollowUpDate", "last followup": "FollowUpDate",
        "last fu": "FollowUpDate", "last f/u": "FollowUpDate",
        "approved": "ApprovedDate", "approved date": "ApprovedDate", "approveddate": "ApprovedDate",
        "approval date": "ApprovedDate", "date approved": "ApprovedDate",
        "expiration": "ExpirationDate", "expires": "ExpirationDate", "expiration date": "ExpirationDate",
        "expirationdate": "ExpirationDate", "exp date": "ExpirationDate", "expiry": "ExpirationDate",
        "expiry date": "ExpirationDate", "renewal date": "ExpirationDate",
        "effective": "ApprovedDate", "effective date": "ApprovedDate",
        # Owner / Notes
        "owner": "Owner", "assigned to": "Owner", "assigned": "Owner", "coordinator": "Owner",
        "representative": "Owner", "rep": "Owner", "analyst": "Owner",
        "notes": "Notes", "comments": "Notes", "comment": "Notes", "remarks": "Notes",
        # Reference / Tracking
        "reference": "Notes", "tracking id": "Notes", "reference / tracking id": "Notes",
        "reference id": "Notes", "tracking": "Notes", "ref id": "Notes",
        # Sub-profile / LOB
        "sub profile": "sub_profile", "subprofile": "sub_profile", "sub_profile": "sub_profile",
        "lob": "sub_profile", "line of business": "sub_profile",
    }
    rows = _parse_excel_rows(content, ext)
    if not rows:
        return 0, ["No rows found"]
    first_row_keys = list(rows[0].keys()) if rows else []
    imported, errors = 0, []

    # Use a single connection for dedup + insert to avoid cross-connection visibility issues
    conn = _get_db()
    try:
        for i, row in enumerate(rows):
            mapped = {}
            for raw_key, val in row.items():
                db_col = _fuzzy_match_column(raw_key, COL_MAP)
                if db_col and val is not None:
                    mapped[db_col] = _clean_val(val)
            if not mapped.get("ProviderName") and not mapped.get("Payor"):
                continue
            # Normalize status to match filter dropdown values
            mapped["Status"] = _normalize_cred_status(mapped.get("Status", ""))
            try:
                existing = conn.execute(
                    "SELECT id FROM credentialing WHERE client_id=? AND ProviderName=? AND Payor=?",
                    (client_id, mapped.get("ProviderName", ""), mapped.get("Payor", ""))
                ).fetchone()
                if existing:
                    allowed = ["ProviderName","Payor","CredType","Status","SubmittedDate",
                               "FollowUpDate","ApprovedDate","ExpirationDate","Owner","Notes","sub_profile"]
                    parts, params = ["updated_at=?"], [_dt_now.now().isoformat()]
                    for f in allowed:
                        if f in mapped:
                            parts.append(f"{f}=?")
                            params.append(mapped[f])
                    if uploaded_by:
                        parts.append("uploaded_by=?")
                        params.append(uploaded_by)
                    params.append(existing["id"])
                    conn.execute(f"UPDATE credentialing SET {','.join(parts)} WHERE id=?", params)
                else:
                    conn.execute("""INSERT INTO credentialing
                        (client_id,ProviderName,Payor,CredType,Status,SubmittedDate,FollowUpDate,ApprovedDate,ExpirationDate,Owner,Notes,sub_profile,uploaded_by)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                        (client_id, mapped.get("ProviderName",""), mapped.get("Payor",""),
                         mapped.get("CredType","Initial"), mapped.get("Status","Not Started"),
                         mapped.get("SubmittedDate",""), mapped.get("FollowUpDate",""),
                         mapped.get("ApprovedDate",""), mapped.get("ExpirationDate",""),
                         mapped.get("Owner",""), mapped.get("Notes",""), mapped.get("sub_profile",""),
                         str(uploaded_by or "")))
                imported += 1
            except Exception as e:
                errors.append(f"Row {i+2}: {e}")
        conn.commit()
    finally:
        conn.close()

    if not imported and first_row_keys:
        errors.append(f"No rows matched. Excel headers found: {first_row_keys[:15]}")
        errors.append("Expected headers like: Provider, Payor/Insurance, Type, Status, Submitted, Follow Up, Approved, Expiration")
    return imported, errors


def _import_enrollment_from_excel(content: bytes, ext: str, client_id: int, uploaded_by: str = ""):
    from app.client_db import get_db as _get_db
    from datetime import datetime as _dt_now

    # ── Status normalization — maps Excel values to dropdown filter values ──
    ENROLL_STATUS_NORMALIZE = {
        # Not Started
        "not started": "Not Started", "new": "Not Started", "none": "Not Started",
        "pending start": "Not Started", "n/a": "Not Started", "to do": "Not Started",
        "open": "Not Started",
        # In Progress
        "in progress": "In Progress", "in-progress": "In Progress", "pending": "In Progress",
        "processing": "In Progress", "working": "In Progress", "active": "In Progress",
        "in process": "In Progress", "in review": "In Progress", "under review": "In Progress",
        "need action": "In Progress", "needs action": "In Progress", "action needed": "In Progress",
        "follow up": "In Progress", "follow-up": "In Progress", "followup": "In Progress",
        "waiting": "In Progress", "awaiting": "In Progress", "in queue": "In Progress",
        # Submitted
        "submitted": "Submitted", "sent": "Submitted", "filed": "Submitted",
        "application submitted": "Submitted", "app submitted": "Submitted",
        "mailed": "Submitted", "faxed": "Submitted", "uploaded": "Submitted",
        "received": "Submitted", "acknowledged": "Submitted",
        # Approved
        "approved": "Approved", "completed": "Approved", "credentialed": "Approved",
        "accepted": "Approved", "effective": "Approved", "live": "Approved",
        "done": "Approved", "granted": "Approved", "passed": "Approved",
        # Enrolled
        "enrolled": "Enrolled", "active - enrolled": "Enrolled", "participating": "Enrolled",
        "contracted": "Enrolled", "in network": "Enrolled", "in-network": "Enrolled",
        "par": "Enrolled",
        # Denied
        "denied": "Denied", "rejected": "Denied", "declined": "Denied",
        "not approved": "Denied", "failed": "Denied", "terminated": "Denied",
        "closed": "Denied", "cancelled": "Denied", "canceled": "Denied",
    }

    VALID_ENROLL_STATUSES = {"Not Started", "In Progress", "Submitted", "Approved", "Enrolled", "Denied"}

    def _normalize_enroll_status(raw):
        if not raw:
            return "Not Started"
        s = str(raw).strip()
        key = s.lower()
        if key in ENROLL_STATUS_NORMALIZE:
            return ENROLL_STATUS_NORMALIZE[key]
        for map_key, normalized in ENROLL_STATUS_NORMALIZE.items():
            if len(map_key) >= 4 and map_key in key:
                return normalized
        for vs in VALID_ENROLL_STATUSES:
            if vs.lower() == key:
                return vs
        return "In Progress"  # Safe default for unrecognized statuses

    COL_MAP = {
        # Provider
        "provider": "ProviderName", "providername": "ProviderName", "provider name": "ProviderName",
        "rendering provider": "ProviderName", "doctor": "ProviderName", "physician": "ProviderName",
        "doctor name": "ProviderName", "physician name": "ProviderName", "practitioner": "ProviderName",
        "rendering": "ProviderName", "servicing provider": "ProviderName",
        "name": "ProviderName",
        # Payor
        "payor": "Payor", "payer": "Payor", "insurance": "Payor",
        "insurance name": "Payor", "insurance company": "Payor", "plan": "Payor",
        "plan name": "Payor", "payer name": "Payor", "carrier": "Payor",
        "insurance plan": "Payor", "health plan": "Payor", "ins": "Payor",
        "primary insurance": "Payor", "primary payor": "Payor", "primary payer": "Payor",
        # Type
        "type": "EnrollType", "enrolltype": "EnrollType", "enroll type": "EnrollType",
        "enrollment type": "EnrollType", "application type": "EnrollType",
        # Status
        "status": "Status", "enrollment status": "Status", "enroll status": "Status",
        "app status": "Status", "application status": "Status", "current status": "Status",
        # Dates
        "submitted": "SubmittedDate", "submitted date": "SubmittedDate", "submitteddate": "SubmittedDate",
        "date submitted": "SubmittedDate", "submission date": "SubmittedDate", "submit date": "SubmittedDate",
        "application submitted": "SubmittedDate", "app submitted": "SubmittedDate",
        "follow up": "FollowUpDate", "followupdate": "FollowUpDate", "follow up date": "FollowUpDate",
        "followup": "FollowUpDate", "followup date": "FollowUpDate", "next follow up": "FollowUpDate",
        "fu date": "FollowUpDate", "f/u date": "FollowUpDate", "f/u": "FollowUpDate",
        "approved": "ApprovedDate", "approved date": "ApprovedDate", "approveddate": "ApprovedDate",
        "approval date": "ApprovedDate", "date approved": "ApprovedDate",
        "effective": "EffectiveDate", "effective date": "EffectiveDate", "effectivedate": "EffectiveDate",
        "eff date": "EffectiveDate", "start date": "EffectiveDate",
        # Owner / Notes
        "owner": "Owner", "assigned to": "Owner", "assigned": "Owner", "coordinator": "Owner",
        "notes": "Notes", "comments": "Notes", "comment": "Notes", "remarks": "Notes",
        "sub profile": "sub_profile", "subprofile": "sub_profile", "sub_profile": "sub_profile",
    }
    rows = _parse_excel_rows(content, ext)
    if not rows:
        return 0, ["No rows found"]
    first_row_keys = list(rows[0].keys()) if rows else []
    imported, errors = 0, []
    conn = _get_db()
    try:
        for i, row in enumerate(rows):
            mapped = {}
            for raw_key, val in row.items():
                db_col = _fuzzy_match_column(raw_key, COL_MAP)
                if db_col and val is not None:
                    mapped[db_col] = _clean_val(val)
            if not mapped.get("ProviderName") and not mapped.get("Payor"):
                continue
            # Normalize status to match filter dropdown values
            mapped["Status"] = _normalize_enroll_status(mapped.get("Status", ""))
            try:
                existing = conn.execute(
                    "SELECT id FROM enrollment WHERE client_id=? AND ProviderName=? AND Payor=?",
                    (client_id, mapped.get("ProviderName", ""), mapped.get("Payor", ""))
                ).fetchone()
                if existing:
                    allowed = ["ProviderName","Payor","EnrollType","Status","SubmittedDate",
                               "FollowUpDate","ApprovedDate","EffectiveDate","Owner","Notes","sub_profile"]
                    parts, params = ["updated_at=?"], [_dt_now.now().isoformat()]
                    for f in allowed:
                        if f in mapped:
                            parts.append(f"{f}=?")
                            params.append(mapped[f])
                    if uploaded_by:
                        parts.append("uploaded_by=?")
                        params.append(uploaded_by)
                    params.append(existing["id"])
                    conn.execute(f"UPDATE enrollment SET {','.join(parts)} WHERE id=?", params)
                else:
                    conn.execute("""INSERT INTO enrollment
                        (client_id,ProviderName,Payor,EnrollType,Status,SubmittedDate,FollowUpDate,ApprovedDate,EffectiveDate,Owner,Notes,sub_profile,uploaded_by)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                        (client_id, mapped.get("ProviderName",""), mapped.get("Payor",""),
                         mapped.get("EnrollType","Enrollment"), mapped.get("Status","Not Started"),
                         mapped.get("SubmittedDate",""), mapped.get("FollowUpDate",""),
                         mapped.get("ApprovedDate",""), mapped.get("EffectiveDate",""),
                         mapped.get("Owner",""), mapped.get("Notes",""), mapped.get("sub_profile",""),
                         str(uploaded_by or "")))
                imported += 1
            except Exception as e:
                errors.append(f"Row {i+2}: {e}")
        conn.commit()
    finally:
        conn.close()
    if not imported and first_row_keys:
        errors.append(f"No rows matched. Excel headers found: {first_row_keys[:15]}")
        errors.append("Expected headers like: Provider, Payor/Insurance, Type, Status, Submitted, Follow Up, Approved, Effective")
    return imported, errors


def _import_edi_from_excel(content: bytes, ext: str, client_id: int, uploaded_by: str = ""):
    from app.client_db import get_db as _get_db
    from datetime import datetime as _dt_now

    # ── Status normalization for EDI/ERA/EFT — maps to dropdown values ──
    EDI_STATUS_NORMALIZE = {
        # Not Started
        "not started": "Not Started", "new": "Not Started", "none": "Not Started",
        "pending start": "Not Started", "n/a": "Not Started", "to do": "Not Started",
        "open": "Not Started",
        # In Process
        "in process": "In Process", "in-process": "In Process", "in progress": "In Process",
        "in-progress": "In Process", "pending": "In Process", "processing": "In Process",
        "working": "In Process", "active": "In Process", "submitted": "In Process",
        "sent": "In Process", "filed": "In Process", "in review": "In Process",
        "under review": "In Process", "need action": "In Process", "needs action": "In Process",
        "waiting": "In Process", "awaiting": "In Process", "testing": "In Process",
        # Live
        "live": "Live", "completed": "Live", "active - live": "Live",
        "approved": "Live", "enrolled": "Live", "effective": "Live",
        "done": "Live", "connected": "Live", "enabled": "Live",
        "set up": "Live", "setup": "Live", "configured": "Live",
        "production": "Live", "go live": "Live", "go-live": "Live",
        # Failed
        "failed": "Failed", "denied": "Failed", "rejected": "Failed",
        "error": "Failed", "not accepted": "Failed", "terminated": "Failed",
        "closed": "Failed", "cancelled": "Failed", "canceled": "Failed",
        "expired": "Failed", "disconnected": "Failed",
    }

    VALID_EDI_STATUSES = {"Not Started", "In Process", "Live", "Failed"}

    def _normalize_edi_status(raw):
        if not raw:
            return "Not Started"
        s = str(raw).strip()
        key = s.lower()
        if key in EDI_STATUS_NORMALIZE:
            return EDI_STATUS_NORMALIZE[key]
        for map_key, normalized in EDI_STATUS_NORMALIZE.items():
            if len(map_key) >= 4 and map_key in key:
                return normalized
        for vs in VALID_EDI_STATUSES:
            if vs.lower() == key:
                return vs
        return "In Process"  # Safe default for unrecognized statuses

    COL_MAP = {
        # Provider
        "provider": "ProviderName", "providername": "ProviderName", "provider name": "ProviderName",
        "rendering provider": "ProviderName", "doctor": "ProviderName", "physician": "ProviderName",
        "doctor name": "ProviderName", "physician name": "ProviderName", "practitioner": "ProviderName",
        "name": "ProviderName",
        # Payor
        "payor": "Payor", "payer": "Payor", "insurance": "Payor",
        "insurance name": "Payor", "insurance company": "Payor", "plan": "Payor",
        "plan name": "Payor", "payer name": "Payor", "carrier": "Payor",
        "insurance plan": "Payor", "health plan": "Payor",
        # EDI-specific
        "payer id": "PayerID", "payerid": "PayerID", "payer_id": "PayerID",
        "edi": "EDIStatus", "edi status": "EDIStatus", "edistatus": "EDIStatus",
        "era": "ERAStatus", "era status": "ERAStatus", "erastatus": "ERAStatus",
        "eft": "EFTStatus", "eft status": "EFTStatus", "eftstatus": "EFTStatus",
        # Dates
        "submitted": "SubmittedDate", "submitted date": "SubmittedDate", "submitteddate": "SubmittedDate",
        "date submitted": "SubmittedDate", "submission date": "SubmittedDate",
        "go live": "GoLiveDate", "golivedate": "GoLiveDate", "go live date": "GoLiveDate",
        "go-live": "GoLiveDate", "golive": "GoLiveDate", "live date": "GoLiveDate",
        # Owner / Notes
        "owner": "Owner", "assigned to": "Owner", "assigned": "Owner", "coordinator": "Owner",
        "notes": "Notes", "comments": "Notes", "comment": "Notes", "remarks": "Notes",
        "sub profile": "sub_profile", "subprofile": "sub_profile", "sub_profile": "sub_profile",
    }
    rows = _parse_excel_rows(content, ext)
    if not rows:
        return 0, ["No rows found"]
    first_row_keys = list(rows[0].keys()) if rows else []
    imported, errors = 0, []
    conn = _get_db()
    try:
        for i, row in enumerate(rows):
            mapped = {}
            for raw_key, val in row.items():
                db_col = _fuzzy_match_column(raw_key, COL_MAP)
                if db_col and val is not None:
                    mapped[db_col] = _clean_val(val)
            if not mapped.get("ProviderName") and not mapped.get("Payor"):
                continue
            # Normalize all three EDI status fields to match dropdown values
            mapped["EDIStatus"] = _normalize_edi_status(mapped.get("EDIStatus", ""))
            mapped["ERAStatus"] = _normalize_edi_status(mapped.get("ERAStatus", ""))
            mapped["EFTStatus"] = _normalize_edi_status(mapped.get("EFTStatus", ""))
            try:
                existing = conn.execute(
                    "SELECT id FROM edi_setup WHERE client_id=? AND ProviderName=? AND Payor=?",
                    (client_id, mapped.get("ProviderName", ""), mapped.get("Payor", ""))
                ).fetchone()
                if existing:
                    allowed = ["ProviderName","Payor","EDIStatus","ERAStatus","EFTStatus",
                               "SubmittedDate","GoLiveDate","PayerID","Owner","Notes","sub_profile"]
                    parts, params = ["updated_at=?"], [_dt_now.now().isoformat()]
                    for f in allowed:
                        if f in mapped:
                            parts.append(f"{f}=?")
                            params.append(mapped[f])
                    if uploaded_by:
                        parts.append("uploaded_by=?")
                        params.append(uploaded_by)
                    params.append(existing["id"])
                    conn.execute(f"UPDATE edi_setup SET {','.join(parts)} WHERE id=?", params)
                else:
                    conn.execute("""INSERT INTO edi_setup
                        (client_id,ProviderName,Payor,EDIStatus,ERAStatus,EFTStatus,SubmittedDate,GoLiveDate,PayerID,Owner,Notes,sub_profile,uploaded_by)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                        (client_id, mapped.get("ProviderName",""), mapped.get("Payor",""),
                         mapped.get("EDIStatus","Not Started"), mapped.get("ERAStatus","Not Started"),
                         mapped.get("EFTStatus","Not Started"),
                         mapped.get("SubmittedDate",""), mapped.get("GoLiveDate",""),
                         mapped.get("PayerID",""), mapped.get("Owner",""), mapped.get("Notes",""),
                         mapped.get("sub_profile",""), str(uploaded_by or "")))
                imported += 1
            except Exception as e:
                errors.append(f"Row {i+2}: {e}")
        conn.commit()
    finally:
        conn.close()
    if not imported and first_row_keys:
        errors.append(f"No rows matched. Excel headers found: {first_row_keys[:15]}")
        errors.append("Expected headers like: Provider, Payor/Insurance, Payer ID, EDI Status, ERA Status, EFT Status")
    return imported, errors


def _import_claims_from_excel(content: bytes, ext: str, client_id: int, uploaded_by: str = "",
                              default_sub_profile: str = ""):
    """
    Parse an Excel/CSV claims report and upsert rows into claims_master.
    Flexible column matching — maps common header names to DB columns.
    Normalizes status values to match standard CLAIM_STATUSES.
    Returns (imported_count, error_list).
    """
    import csv, io
    from app.client_db import get_db
    from datetime import date as _date

    # Status normalization map — maps common Excel status values to standard statuses
    STATUS_NORMALIZE = {
        # Intake
        "intake": "Intake", "new": "Intake", "received": "Intake", "open": "Intake",
        "entered": "Intake", "created": "Intake", "registered": "Intake",
        # Verification
        "verification": "Verification", "verify": "Verification", "verifying": "Verification",
        "eligibility": "Verification", "elig check": "Verification", "auth": "Verification",
        "authorization": "Verification", "pre-auth": "Verification", "precert": "Verification",
        # Coding
        "coding": "Coding", "coded": "Coding", "code review": "Coding",
        "charge entry": "Coding", "charge review": "Coding",
        # Billed/Submitted
        "billed/submitted": "Billed/Submitted", "billed": "Billed/Submitted",
        "submitted": "Billed/Submitted", "filed": "Billed/Submitted", "sent": "Billed/Submitted",
        "pending": "Billed/Submitted", "in process": "Billed/Submitted",
        "in-process": "Billed/Submitted", "processing": "Billed/Submitted",
        "pending payment": "Billed/Submitted", "awaiting payment": "Billed/Submitted",
        "claim submitted": "Billed/Submitted", "billed to insurance": "Billed/Submitted",
        # Rejected
        "rejected": "Rejected", "reject": "Rejected", "returned": "Rejected",
        "kicked back": "Rejected", "not accepted": "Rejected", "error": "Rejected",
        "failed": "Rejected", "invalid": "Rejected",
        # Denied
        "denied": "Denied", "deny": "Denied", "denial": "Denied",
        "not covered": "Denied", "non-covered": "Denied",
        "denied - initial": "Denied", "initial denial": "Denied",
        # A/R Follow-Up
        "a/r follow-up": "A/R Follow-Up", "a/r follow up": "A/R Follow-Up",
        "ar follow up": "A/R Follow-Up", "ar follow-up": "A/R Follow-Up",
        "ar followup": "A/R Follow-Up", "follow up": "A/R Follow-Up",
        "follow-up": "A/R Follow-Up", "followup": "A/R Follow-Up",
        "in review": "A/R Follow-Up", "under review": "A/R Follow-Up",
        "pending review": "A/R Follow-Up", "working": "A/R Follow-Up",
        "in progress": "A/R Follow-Up",
        # Appeals
        "appeals": "Appeals", "appeal": "Appeals", "appealed": "Appeals",
        "appeal filed": "Appeals", "reconsideration": "Appeals",
        "corrected claim": "Appeals", "resubmitted": "Appeals",
        # Paid
        "paid": "Paid", "approved": "Paid", "finalized": "Paid",
        "payment received": "Paid", "closed - paid": "Paid",
        "settled": "Paid", "remitted": "Paid", "collected": "Paid",
        # Closed
        "closed": "Closed", "write off": "Closed", "write-off": "Closed",
        "written off": "Closed", "adjusted": "Closed", "void": "Closed",
        "voided": "Closed", "zero balance": "Closed", "closed - adjusted": "Closed",
    }

    # Statuses that precede billing — a claim in any of these has NOT gone out
    # the door yet, so it legitimately carries no Bill Date.
    _PRE_BILL_STATUSES = {"Intake", "Verification", "Coding"}

    def _normalize_status(raw):
        if not raw:
            return "Intake"
        s = str(raw).strip()
        key = s.lower()
        if key in STATUS_NORMALIZE:
            return STATUS_NORMALIZE[key]
        # Partial match: check if any known key is contained in the value
        for map_key, normalized in STATUS_NORMALIZE.items():
            if len(map_key) >= 4 and map_key in key:
                return normalized
        # If the raw value already matches a standard status (case-insensitive), use it
        from app.client_db import CLAIM_STATUSES
        for cs in CLAIM_STATUSES:
            if cs.lower() == key:
                return cs
        # Default: return as-is but log it
        return s

    # Use the module-level mapping parameters (single source of truth) so the
    # importer and the admin mapping diagnostic recognize columns identically.
    COLUMN_MAP = CLAIMS_COLUMN_MAP

    def _parse_float(v):
        try:
            return float(str(v).replace("$", "").replace(",", "").strip())
        except Exception:
            return 0.0

    def _parse_date(v):
        if not v:
            return ""
        s = str(v).strip()
        # Handle datetime objects from openpyxl directly
        from datetime import datetime, date as _dt_date
        if isinstance(v, (datetime, _dt_date)):
            return v.strftime("%Y-%m-%d")
        # Strip time component if present (e.g. "2025-07-07 00:00:00")
        if " " in s and len(s) > 10:
            s = s.split(" ")[0]
        for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m-%d-%Y", "%m/%d/%y", "%Y%m%d",
                     "%d-%b-%Y", "%d-%b-%y", "%b %d, %Y", "%B %d, %Y",
                     "%d/%m/%Y", "%Y/%m/%d"):
            try:
                return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
            except Exception:
                pass
        return s

    # Parse rows: try the hardcoded format templates first (so headerless /
    # batch-summary / multi-sheet SV exports are recognized), then fall back to the
    # shared smart parser (multi-sheet, smart header detection).
    rows, _tpl_used = _load_claim_rows(content, ext)

    if not rows:
        return 0, ["No rows found in file"]

    conn = get_db()
    cur = conn.cursor()
    today_str = business_today_iso()
    imported = 0
    errors = []
    counter = 1
    # When the file is a posted-payments export (LIMS deposits), each row is a
    # payment actually posted/deposited — not just a claim. We mirror those into
    # the payments table so the dashboard "Posted" bucket reflects the real money
    # posted (e.g. the ~$7k of EFT deposits), in addition to landing on the claim.
    _is_payment_template = (_tpl_used == "lims_payments")
    payment_rows = []

    # Pass 1: fuzzy-map every row, and tally how many times each file-provided
    # claim/account number appears. A single claim frequently spans several
    # service-line rows (one per CPT), all sharing the same claim number. The
    # upsert key is (client_id, ClaimKey), so without disambiguation every line
    # after the first overwrites the previous one — only the last line's charge
    # survives and the admin "billed" total is badly under-counted. Counting the
    # occurrences here lets Pass 2 keep each service line as its own row.
    mapped_rows = []
    base_key_counts = {}
    for row in rows:
        mapped = {}
        for raw_key, val in row.items():
            db_col = _fuzzy_match_column(raw_key, COLUMN_MAP)
            if db_col:
                mapped[db_col] = val
        if not mapped:
            continue
        base_key = str(mapped.get("ClaimKey", "")).strip()
        if base_key:
            base_key_counts[base_key] = base_key_counts.get(base_key, 0) + 1
        mapped_rows.append(mapped)

    try:
        for mapped in mapped_rows:
            base_key = str(mapped.get("ClaimKey", "")).strip()
            if base_key:
                # Multi-service-line claim: keep each line distinct by appending a
                # deterministic per-line suffix derived from the line's financial /
                # service fields. Re-uploading the same file regenerates identical
                # suffixes, so the upsert updates in place and totals never double
                # count. Single-line claims keep the bare claim number untouched so
                # the common case (and the claim number shown in the UI) is unchanged.
                mapped["ClaimKey"] = base_key
                if base_key_counts.get(base_key, 0) > 1:
                    line_sig = "|".join(
                        str(mapped.get(f, "")).strip().lower()
                        for f in ("CPTCode", "DOS", "ChargeAmount", "AllowedAmount",
                                  "AdjustmentAmount", "PaidAmount", "Description",
                                  "BillDate", "PaidDate", "DenialReason")
                    )
                    digest = hashlib.sha1(
                        f"{base_key}|{line_sig}".encode("utf-8")
                    ).hexdigest()[:8].upper()
                    mapped["ClaimKey"] = f"{base_key}-L{digest}"
            else:
                # Generate a stable ClaimKey when the source file has no claim/account
                # number. Use a deterministic hash of the claim's identifying fields so
                # (a) distinct claims never collide — even across several files uploaded
                # the same day — and (b) re-uploading the same claim updates the same row
                # instead of silently overwriting an unrelated one. The old behaviour used
                # a per-file row counter (IMP-<date>-0001 …), so the 2nd, 3rd … upload of
                # the day reused those exact keys and clobbered earlier claims — which is
                # why freshly imported data never moved the admin totals.
                sig = "|".join(
                    str(mapped.get(f, "")).strip().lower()
                    for f in ("PatientName", "PatientID", "Payor", "ProviderName",
                              "NPI", "DOS", "CPTCode", "ChargeAmount")
                )
                if sig.strip("|"):
                    digest = hashlib.sha1(sig.encode("utf-8")).hexdigest()[:12].upper()
                    mapped["ClaimKey"] = f"IMP-{digest}"
                else:
                    # Truly empty signature (blank row) — fall back to a unique key so
                    # multiple blank rows don't all collapse onto a single record.
                    mapped["ClaimKey"] = f"IMP-{today_str}-{uuid.uuid4().hex[:8].upper()}"
            counter += 1

            # Normalize claim status to standard values
            raw_status = mapped.get("ClaimStatus", "Intake")
            mapped["ClaimStatus"] = _normalize_status(raw_status)

            # Inherit the uploader's active sub-profile when the file itself
            # doesn't carry a sub_profile column. Without this, imported claims
            # default to '' and stay hidden whenever the user is viewing a
            # specific sub-profile in the Claims Queue ("ghost" claims).
            if not str(mapped.get("sub_profile", "")).strip() and str(default_sub_profile or "").strip():
                mapped["sub_profile"] = str(default_sub_profile).strip()

            # Derive the outstanding balance when the source file has no balance
            # column. Without this a claims report that lists charges/payments but
            # no explicit balance imported as AR=0, so the admin dashboard never
            # reflected the real outstanding amount. Balance = Charge - Adjustment
            # - Paid (never negative). Paid/Closed claims carry no AR.
            if "BalanceRemaining" not in mapped:
                if mapped["ClaimStatus"] in ("Paid", "Closed"):
                    mapped["BalanceRemaining"] = 0.0
                else:
                    derived = (_parse_float(mapped.get("ChargeAmount", 0))
                               - _parse_float(mapped.get("AdjustmentAmount", 0))
                               - _parse_float(mapped.get("PaidAmount", 0)))
                    mapped["BalanceRemaining"] = max(derived, 0.0)

            # A claim that has reached (or moved past) "Billed/Submitted" MUST
            # carry a Bill Date — every dated billed/production view (Billed
            # Activity, All-Time Billed, the Team Production "$ Billed" column,
            # AR aging) keys off BillDate. Many source files set the status to a
            # billed value but ship no bill-date column, so those claims imported
            # with BillDate='' and were invisible to all of those reports (e.g.
            # 671 Billed/Submitted claims showing $0 all-time billed). When a
            # billed claim has no parseable Bill Date, stamp one: prefer the
            # service date (DOS) so historical claims keep a realistic timeline,
            # and fall back to the import date only when DOS is missing too.
            _bill_date = _parse_date(mapped.get("BillDate", ""))
            if not _bill_date and mapped["ClaimStatus"] not in _PRE_BILL_STATUSES:
                _bill_date = _parse_date(mapped.get("DOS", "")) or today_str
            mapped["BillDate"] = _bill_date

            try:
                cur.execute("""
                    INSERT INTO claims_master
                    (client_id, ClaimKey, PatientID, PatientName, Payor, ProviderName, NPI,
                     DOS, CPTCode, Description, ChargeAmount, AllowedAmount, AdjustmentAmount,
                     PaidAmount, BalanceRemaining, ClaimStatus, BillDate, DeniedDate, PaidDate,
                     DenialCategory, DenialReason, Owner, StatusStartDate, LastTouchedDate, sub_profile,
                     uploaded_by)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(client_id, ClaimKey) DO UPDATE SET
                        PatientID=excluded.PatientID, PatientName=excluded.PatientName,
                        Payor=excluded.Payor, ProviderName=excluded.ProviderName, NPI=excluded.NPI,
                        DOS=excluded.DOS, CPTCode=excluded.CPTCode, Description=excluded.Description,
                        ChargeAmount=excluded.ChargeAmount, AllowedAmount=excluded.AllowedAmount,
                        AdjustmentAmount=excluded.AdjustmentAmount, PaidAmount=excluded.PaidAmount,
                        BalanceRemaining=excluded.BalanceRemaining, ClaimStatus=excluded.ClaimStatus,
                        BillDate=CASE WHEN TRIM(COALESCE(claims_master.BillDate,''))<>''
                                     THEN claims_master.BillDate ELSE excluded.BillDate END,
                        DeniedDate=excluded.DeniedDate, PaidDate=excluded.PaidDate,
                        DenialCategory=excluded.DenialCategory, DenialReason=excluded.DenialReason,
                        Owner=excluded.Owner, LastTouchedDate=excluded.LastTouchedDate,
                        sub_profile=excluded.sub_profile, uploaded_by=excluded.uploaded_by,
                        updated_at=CURRENT_TIMESTAMP
                """, (
                    client_id,
                    str(mapped.get("ClaimKey", "")),
                    str(mapped.get("PatientID", "")),
                    str(mapped.get("PatientName", "")),
                    str(mapped.get("Payor", "")),
                    str(mapped.get("ProviderName", "")),
                    str(mapped.get("NPI", "")),
                    _parse_date(mapped.get("DOS", "")),
                    str(mapped.get("CPTCode", "")),
                    str(mapped.get("Description", "")),
                    _parse_float(mapped.get("ChargeAmount", 0)),
                    _parse_float(mapped.get("AllowedAmount", 0)),
                    _parse_float(mapped.get("AdjustmentAmount", 0)),
                    _parse_float(mapped.get("PaidAmount", 0)),
                    _parse_float(mapped.get("BalanceRemaining", 0)),
                    str(mapped["ClaimStatus"]),
                    _parse_date(mapped.get("BillDate", "")),
                    _parse_date(mapped.get("DeniedDate", "")),
                    _parse_date(mapped.get("PaidDate", "")),
                    str(mapped.get("DenialCategory", "")),
                    str(mapped.get("DenialReason", "")),
                    str(mapped.get("Owner", "")),
                    today_str,
                    today_str,
                    str(mapped.get("sub_profile", "")),
                    str(uploaded_by or ""),
                ))
                imported += 1
                if _is_payment_template:
                    _pmt_amt = _parse_float(mapped.get("PaidAmount", 0))
                    if _pmt_amt > 0:
                        _ck = str(mapped.get("ClaimKey", ""))
                        payment_rows.append({
                            "ClaimKey": _ck,
                            "PostDate": _parse_date(mapped.get("PaidDate", "")),
                            "PaymentAmount": _pmt_amt,
                            "PayerType": "Primary",
                            "CheckNumber": _ck[4:] if _ck.startswith("PMT-") else "",
                            "sub_profile": str(mapped.get("sub_profile", "")),
                        })
            except Exception as e:
                errors.append(f"Row {counter}: {e}")

        # Mirror posted payments into the payments table so the "Posted" bucket
        # reflects real deposits. Delete-then-insert keyed on the deterministic
        # PMT-<eft> ClaimKey keeps this idempotent across reimports (no double
        # counting). PostedBy carries the uploader so attribution is preserved.
        if payment_rows:
            _pmt_keys = list({pr["ClaimKey"] for pr in payment_rows})
            cur.executemany(
                "DELETE FROM payments WHERE client_id=? AND ClaimKey=?",
                [(client_id, k) for k in _pmt_keys],
            )
            # Posting is captured per-claim in the notes log so Paid and Posted
            # stay the SAME number on the dashboard, while each claim still shows
            # whether (and when) its payment was posted — flagged "recently
            # posted" when the deposit landed in the last 7 days. Module='Payment'
            # auto-notes are cleared first so reimports never stack duplicates.
            cur.executemany(
                "DELETE FROM notes_log WHERE client_id=? AND ClaimKey=? AND Module='Payment' AND Author='system'",
                [(client_id, k) for k in _pmt_keys],
            )
            _recent_cutoff = (business_today() - timedelta(days=7)).isoformat()
            for pr in payment_rows:
                cur.execute(
                    """INSERT INTO payments
                       (client_id, ClaimKey, PostDate, PaymentAmount, PayerType,
                        CheckNumber, PostedBy, sub_profile)
                       VALUES (?,?,?,?,?,?,?,?)""",
                    (client_id, pr["ClaimKey"], pr["PostDate"], pr["PaymentAmount"],
                     pr["PayerType"], pr["CheckNumber"], str(uploaded_by or ""),
                     pr["sub_profile"]),
                )
                _pdate = str(pr["PostDate"] or "").strip()
                _recent = bool(_pdate) and _pdate >= _recent_cutoff
                _prefix = "Recently posted" if _recent else "Posted"
                _when = f" on {_pdate}" if _pdate else ""
                _eft = pr.get("CheckNumber") or ""
                _note = f"{_prefix}{_when}: ${pr['PaymentAmount']:,.2f}" + (f" (EFT {_eft})" if _eft else "")
                cur.execute(
                    """INSERT INTO notes_log (client_id, ClaimKey, Module, RefID, Note, Author)
                       VALUES (?,?,?,?,?,?)""",
                    (client_id, pr["ClaimKey"], "Payment", 0, _note, "system"),
                )

        conn.commit()
    finally:
        conn.close()
    # Report unmapped headers as info for debugging
    if rows:
        first_row_keys = list(rows[0].keys())
        unmapped = [k for k in first_row_keys if not _fuzzy_match_column(k, COLUMN_MAP)]
        if unmapped:
            errors.append(f"Unmapped Excel columns (ignored): {unmapped[:10]}")
    return imported, errors


@router.get("/files/{file_id}/download")
def download_file(file_id: int, hub_session: Optional[str] = Cookie(None)):
    """Download the original uploaded file."""
    user = _require_user(hub_session)
    if user.get("role") in ("admin", "staff"):
        rec = get_file_record(file_id, _client_scope(user))
    else:
        rec = get_file_record(file_id, None)
        if rec and int(rec.get("client_id") or 0) not in set(_doc_account_ids(user)):
            rec = None
    if not rec:
        raise HTTPException(404, "File not found")
    path = os.path.join(UPLOAD_DIR, rec["filename"])
    if not os.path.isfile(path):
        # The DB still references the file but the bytes are gone from disk. This
        # happens when uploads are written to non-persistent storage (e.g. the
        # UPLOAD_DIR is not a mounted persistent volume) and the server restarts.
        # Surface a clear message instead of a generic 404 so the cause is obvious.
        log.warning("Download failed — file row %s (%s) missing on disk at %s",
                    file_id, rec.get("original_name"), path)
        raise HTTPException(
            410,
            "This file is no longer available in storage. It was uploaded earlier "
            "but the stored copy is missing — please re-upload it. (If this keeps "
            "happening, uploads are not being saved to persistent storage.)",
        )
    from fastapi.responses import FileResponse
    return FileResponse(
        path,
        filename=rec["original_name"],
        media_type="application/octet-stream",
    )


@router.post("/files/{file_id}/replace")
async def replace_file(
    file_id: int,
    file: UploadFile = FastAPIFile(...),
    hub_session: Optional[str] = Cookie(None),
):
    """Replace an existing uploaded file with a new version.
    The old file remains until the new upload is validated and persisted.
    If it's an Excel in a data category, the data is re-imported."""
    user = _require_user(hub_session)
    scope = _client_scope(user)
    rec = get_file_record(file_id, scope)
    if not rec:
        raise HTTPException(404, "File not found")

    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in (".xlsx", ".xls", ".csv", ".ods", ".odf", ".pdf", ".doc", ".docx"):
        raise HTTPException(400, "Unsupported file type")

    content = await file.read()
    file_size = len(content)

    # Count rows for Excel/CSV
    row_count = 0
    file_type = "excel" if ext in (".xlsx", ".xls", ".csv", ".ods", ".odf") else "pdf"
    if file_type == "excel":
        try:
            import csv as _csv, io as _io
            if ext == ".csv":
                reader = _csv.reader(_io.StringIO(content.decode("utf-8", errors="replace")))
                row_count = max(0, sum(1 for _ in reader) - 1)
            else:
                row_count = len(_parse_excel_rows(content, ext, combine_sheets=True))
        except Exception:
            pass

    # Hard intercept (before replacing old file): never allow ambiguous excel routing.
    category = rec.get("category", "")
    effective_category = category
    category_source = "existing"
    infer_debug = None
    if file_type == "excel" and category not in DATA_IMPORT_CATEGORIES:
        inferred, infer_debug = _infer_excel_category(content, ext, file.filename or "", rec.get("description", "") or "")
        if inferred in DATA_IMPORT_CATEGORIES:
            effective_category = inferred
            category_source = "auto"
        else:
            # Can't map to a data section — keep existing document category, no import.
            effective_category = category
            category_source = "document"

    # Save new file after validation succeeds.
    new_unique = f"{uuid.uuid4().hex}{ext}"
    new_path = os.path.join(UPLOAD_DIR, new_unique)
    with open(new_path, "wb") as f:
        f.write(content)

    # Update DB record
    try:
        update_file_record(file_id, {
            "filename": new_unique,
            "original_name": file.filename or rec["original_name"],
            "file_type": file_type,
            "file_size": file_size,
            "row_count": row_count,
            "uploaded_by": user["username"],
            "status": "Replaced",
            "category": effective_category,
        }, scope)
    except Exception:
        if os.path.isfile(new_path):
            os.remove(new_path)
        raise

    # Delete old file from disk only after DB update succeeds.
    old_path = os.path.join(UPLOAD_DIR, rec["filename"])
    if os.path.isfile(old_path):
        os.remove(old_path)

    # Auto re-import if data category
    imported = 0
    import_errors = []
    if effective_category in DATA_IMPORT_CATEGORIES and file_type in ("excel", "pdf", "document"):
        _uploader = user.get("username") or ""
        try:
            if effective_category == "Claims":
                imported, import_errors = _import_claims_from_excel(content, ext, scope, uploaded_by=_uploader)
            elif file_type == "excel" and effective_category == "Credentialing":
                imported, import_errors = _import_credentialing_from_excel(content, ext, scope, uploaded_by=_uploader)
            elif file_type == "excel" and effective_category == "Enrollment":
                imported, import_errors = _import_enrollment_from_excel(content, ext, scope, uploaded_by=_uploader)
            elif file_type == "excel" and effective_category == "EDI":
                imported, import_errors = _import_edi_from_excel(content, ext, scope, uploaded_by=_uploader)
        except Exception as e:
            import_errors = [str(e)]

    notify_activity(user["username"], "replaced file", "Documents",
                    f"{rec['original_name']} → {file.filename}")

    return {
        "ok": True,
        "file_id": file_id,
        "original_name": file.filename,
        "effective_category": effective_category,
        "category_source": category_source,
        "category_inference": infer_debug,
        "imported": imported,
        "import_errors": import_errors[:5],
    }


def auto_import_pending_claim_files(client_id: Optional[int] = None) -> dict:
    """Seamlessly import any uploaded spreadsheets that look like claims but were
    saved under a non-data category (e.g. "General") and never imported.

    This is what makes ingestion zero-touch: the team can upload whatever they
    have under any category, and this sweep picks up anything claim-shaped,
    imports it into claims_master, and re-files it under "Claims" so it counts
    toward billed/submitted totals — with no manual click. The claims importer
    upserts by (client_id, ClaimKey), so re-running this is idempotent and never
    double-counts. Runs on startup (whole DB) and per-account when Documents or
    the Dashboard load.

    Returns {"files": n_files_imported, "rows": n_rows_imported}."""
    from .client_db import get_db as _get_db
    conn = _get_db()
    try:
        cur = conn.cursor()
        cats = ",".join("?" * len(DATA_IMPORT_CATEGORIES))
        params = list(DATA_IMPORT_CATEGORIES)
        where = f"file_type IN ('excel','pdf','document') AND category NOT IN ({cats})"
        if client_id is not None:
            where += " AND client_id=?"
            params.append(int(client_id))
        rows = cur.execute(
            f"SELECT id, client_id, filename, original_name, uploaded_by "
            f"FROM client_files WHERE {where}",
            params,
        ).fetchall()
    finally:
        conn.close()

    files_done = 0
    rows_done = 0
    for r in rows:
        path = os.path.join(UPLOAD_DIR, r["filename"])
        if not os.path.isfile(path):
            continue
        SWEEP_EXTS = (".xlsx", ".xls", ".csv", ".ods", ".odf", ".pdf", ".doc", ".docx")
        ext = os.path.splitext(r["original_name"] or "")[1].lower()
        if ext not in SWEEP_EXTS:
            ext = os.path.splitext(r["filename"])[1].lower()
        if ext not in SWEEP_EXTS:
            continue
        spreadsheet = ext in (".xlsx", ".xls", ".csv", ".ods", ".odf")
        try:
            with open(path, "rb") as fh:
                content = fh.read()
            # Recognize via hardcoded templates OR generic structural match first
            # (fast path / headerless SV exports). When neither matches we do NOT
            # skip — that silent skip is exactly how daily worklists with unusual
            # headers got shelved as inert "General" documents and never computed.
            # Unless the sheet clearly belongs to another data section, ATTEMPT the
            # claims import and let the importer be the judge: it skips any row that
            # doesn't map to claim columns, so a genuinely non-claims sheet imports
            # 0 rows and is left untouched below. PDFs / Word docs skip the
            # spreadsheet-only structural pre-check and go straight to extraction.
            if spreadsheet:
                templated = bool(_extract_templated_claim_rows(content, ext))
                if not templated:
                    parsed = _parse_excel_rows(content, ext, combine_sheets=True)
                    headers = list(parsed[0].keys()) if parsed else []
                    if _is_clearinghouse_ack(headers):
                        # Submission acknowledgement recap (echoes an already-billed
                        # register) — importing it double-bills, so never sweep it in.
                        continue
                    if _is_batch_transmission_log(headers) or _is_svd_batch_workbook(content, ext):
                        # SVD DAILY collective transmission recap -- the same claims as
                        # the per-claim register; importing it double-bills Melissa+Susan.
                        continue
                    if not _claims_structural_match(headers)["is_claims"]:
                        inferred, _dbg = _infer_excel_category(
                            content, ext, r["original_name"] or "", "")
                        if inferred in ("Credentialing", "Enrollment", "EDI"):
                            continue  # belongs to another section — don't import as claims
            imported, _errors = _import_claims_from_excel(
                content, ext, int(r["client_id"]),
                uploaded_by=str(r["uploaded_by"] or ""))
        except Exception as e:
            log.warning("auto-import sweep: file %s failed: %s", r["id"], e)
            continue
        if imported:
            files_done += 1
            rows_done += imported
            try:
                update_file_record(int(r["id"]),
                                   {"category": "Claims", "status": "Imported"},
                                   int(r["client_id"]))
            except Exception:
                pass
    if files_done:
        log.info("Auto-import sweep ingested %s row(s) from %s file(s)",
                 rows_done, files_done)
    return {"files": files_done, "rows": rows_done}


@router.post("/files/{file_id}/import-claims")
def import_stored_file_as_claims(file_id: int, hub_session: Optional[str] = Cookie(None)):
    """Import an already-uploaded file into claims_master.

    Recovers daily work that was saved under a non-data category (e.g. "General")
    and therefore never auto-imported. Reads the stored file from disk — Excel,
    CSV, PDF or Word — runs it through the universal Claims reader, and re-files it
    under "Claims" so it's no longer flagged as a pending import."""
    user = _require_user(hub_session)
    if user.get("role") in ("admin", "staff"):
        rec = get_file_record(file_id, _client_scope(user))
    else:
        rec = get_file_record(file_id, None)
        if rec and int(rec.get("client_id") or 0) not in set(_doc_account_ids(user)):
            rec = None
    if not rec:
        raise HTTPException(404, "File not found")

    path = os.path.join(UPLOAD_DIR, rec["filename"])
    if not os.path.isfile(path):
        raise HTTPException(
            410,
            "The stored copy of this file is missing — please re-upload it.",
        )

    ext = os.path.splitext(rec.get("original_name") or "")[1].lower()
    if ext not in (".xlsx", ".xls", ".csv", ".ods", ".odf", ".pdf", ".doc", ".docx"):
        ext = os.path.splitext(rec["filename"])[1].lower()
    with open(path, "rb") as f:
        content = f.read()

    client_id = int(rec["client_id"])
    uploader = user.get("username") or ""
    try:
        imported, import_errors = _import_claims_from_excel(
            content, ext, client_id, uploaded_by=uploader)
    except Exception as e:
        raise HTTPException(400, f"Import failed: {e}")

    # Re-file under Claims so it reads as imported going forward.
    try:
        update_file_record(file_id, {"category": "Claims", "status": "Imported"}, client_id)
    except Exception:
        pass

    try:
        notify_activity(uploader, "imported claims", "Documents",
                        f'{rec.get("original_name")} → {imported} claim row(s)')
        if imported:
            notify_bulk_activity(uploader, "imported", "Claims", imported,
                                 f'from stored file "{rec.get("original_name")}"')
    except Exception as _e:
        log.warning("import-claims notify failed: %s", _e)

    return {
        "ok": True,
        "file_id": file_id,
        "original_name": rec.get("original_name"),
        "imported": imported,
        "import_errors": import_errors[:5],
    }


@router.delete("/files/{file_id}")
def delete_file(file_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    # Also delete the physical file from disk
    rec = get_file_record(file_id, scope)
    if rec:
        path = os.path.join(UPLOAD_DIR, rec["filename"])
        if os.path.isfile(path):
            os.remove(path)
    delete_file_record(file_id, scope)
    notify_activity(user["username"], "deleted file", "Documents",
                    rec["original_name"] if rec else "")
    return {"ok": True}


# ─── AI Report Generation (OpenAI GPT) ───────────────────────────────────────

@router.post("/report/{client_id}/ai-narrative")
async def generate_ai_narrative(client_id: int, hub_session: Optional[str] = Cookie(None)):
    """Send dashboard/report data to OpenAI GPT and return a professional narrative."""
    user = _require_reporting_access(hub_session)
    from app.config import OPENAI_API_KEY
    from app.client_db import get_db
    from datetime import date

    # Gather all report data
    conn = get_db()
    conn.row_factory = sqlite3.Row
    client_row = conn.execute("SELECT company,contact_name,email,phone,practice_type,specialty FROM clients WHERE id=?", (client_id,)).fetchone()
    client_info = dict(client_row) if client_row else {}
    practice_type = client_info.get("practice_type", "") or ""

    try:
        overall = _build_section_data(conn, client_id)
        sub_profiles = {}
        if practice_type == "MHP+OMT":
            for sp_name in ["OMT", "MHP"]:
                sub_profiles[sp_name] = _build_section_data(conn, client_id, sub_profile=sp_name)
    finally:
        conn.close()

    production_snapshot = get_user_production_snapshot()
    prod_users = production_snapshot.get("user_stats", [])
    prod_total_entries = production_snapshot.get("total_entries", 0)
    prod_total_hours = production_snapshot.get("total_hours", 0)
    prod_by_user_lines = chr(10).join(
        f"  - {u.get('username','?')}: {u.get('entry_count',0)} entries, {u.get('total_hours',0)}h, {u.get('total_qty',0)} qty"
        for u in prod_users
    ) or "  None"

    if not OPENAI_API_KEY:
        # Rule-based narrative — no API key needed
        try:
            cl = overall.get("claims", {})
            cred = overall.get("credentialing", {})
            enr = overall.get("enrollment", {})
            edi = overall.get("edi", {})
            pay = overall.get("payments", {})
            company = client_info.get("company", "the practice")
            charged = cl.get("total_charged", 0)
            paid = cl.get("total_paid", 0)
            balance = cl.get("total_balance", 0)
            total_claims = cl.get("total", 0)
            coll_rate = round((paid / charged) * 100, 1) if charged else 0
            denials = cl.get("top_denials", [])
            cred_count = len(cred.get("detail", []))
            enr_count = len(enr.get("detail", []))
            edi_count = len(edi.get("detail", []))

            health = "healthy" if coll_rate >= 90 else "moderate" if coll_rate >= 70 else "needs attention"
            narrative = (
                f"<b>Executive Summary:</b> {company} shows a {health} revenue cycle with a collection rate of <b>{coll_rate}%</b> "
                f"on <b>{total_claims}</b> total claims. Total charges stand at <b>${charged:,.2f}</b> against payments of <b>${paid:,.2f}</b>, "
                f"leaving an outstanding A/R balance of <b>${balance:,.2f}</b>.\n\n"
            )
            if denials:
                top = denials[0]
                narrative += (
                    f"<b>Denial Management:</b> The leading denial category is <b>{top.get('category','Unknown')}</b> "
                    f"({top.get('count', 0)} claims). Addressing this category represents the highest-leverage action "
                    f"to recover revenue. "
                )
                if len(denials) > 1:
                    narrative += f"Additional denial categories include: {', '.join(d.get('category','?') for d in denials[1:4])}. "
                narrative += "\n\n"
            if cred_count or enr_count or edi_count:
                narrative += (
                    f"<b>Operational Status:</b> Credentialing shows <b>{cred_count}</b> active records, "
                    f"enrollment <b>{enr_count}</b> records, and EDI connectivity is configured for <b>{edi_count}</b> connections. "
                    "Ensure all pending credentialing items are resolved to avoid future payment delays.\n\n"
                )
            narrative += (
                f"<b>User Production Analysis:</b> Team members logged <b>{prod_total_entries}</b> production entries totaling "
                f"<b>{prod_total_hours}</b> hours today. "
            )
            if prod_users:
                top_user = prod_users[0]
                narrative += (
                    f"Top contributor was <b>{top_user.get('username','')}</b> with "
                    f"<b>{top_user.get('total_hours',0)}h</b> across <b>{top_user.get('entry_count',0)}</b> entries.\n\n"
                )
            else:
                narrative += "No team production entries were logged for today.\n\n"
            if coll_rate < 80:
                narrative += (
                    "<b>Recommended Actions:</b> (1) Work down the top denial category immediately. "
                    "(2) Review claim submission timely filing windows. "
                    "(3) Confirm all providers are enrolled with all active payors. "
                    "(4) Audit EDI/ERA/EFT setup for any inactive connections.\n\n"
                )
            else:
                narrative += (
                    "<b>Recommended Actions:</b> Maintain current billing cadence. "
                    "Continue monitoring denial trends weekly and confirm all credentialing is current.\n\n"
                )
            narrative += f"<b>Outlook:</b> With continued focus on denial resolution and timely claim submission, {company} is positioned to improve net collections."
            return {"narrative": narrative, "model": "rule-based", "company": client_info.get("company", "")}
        except Exception:
            return {"narrative": "Narrative generation unavailable — set OPENAI_API_KEY for AI narratives.", "model": "none", "company": client_info.get("company", "")}

    # Build concise data summary for GPT
    cl = overall.get("claims", {})
    cred = overall.get("credentialing", {})
    enr = overall.get("enrollment", {})
    edi = overall.get("edi", {})
    pay = overall.get("payments", {})

    data_summary = f"""
PRACTICE: {client_info.get('company', 'Unknown')}
SPECIALTY: {client_info.get('specialty', 'N/A')}
PRACTICE TYPE: {practice_type or 'Standard'}
REPORT DATE: {business_today_iso()}

CLAIMS OVERVIEW:
- Total Claims: {cl.get('total', 0)}
- Total Charged: ${cl.get('total_charged', 0):,.2f}
- Total Paid: ${cl.get('total_paid', 0):,.2f}
- Outstanding A/R: ${cl.get('total_balance', 0):,.2f}
- Collection Rate: {round((cl.get('total_paid',0) / cl.get('total_charged',1)) * 100, 1) if cl.get('total_charged') else 0}%

CLAIMS BY STATUS:
{chr(10).join(f"  - {s.get('status','?')}: {s.get('count',0)} claims, ${s.get('charged',0):,.2f} charged, ${s.get('paid',0):,.2f} paid" for s in cl.get('by_status', []))}

TOP DENIAL CATEGORIES:
{chr(10).join(f"  - {d.get('category','?')}: {d.get('count',0)}" for d in cl.get('top_denials', [])) or '  None'}

CREDENTIALING: {len(cred.get('detail', []))} records
{chr(10).join(f"  - {c.get('provider','?')} / {c.get('payor','?')}: {c.get('status','?')}" for c in cred.get('detail', [])[:10])}

ENROLLMENT: {len(enr.get('detail', []))} records
{chr(10).join(f"  - {e.get('provider','?')} / {e.get('payor','?')}: {e.get('status','?')}" for e in enr.get('detail', [])[:10])}

EDI SETUP: {len(edi.get('detail', []))} connections
{chr(10).join(f"  - {e.get('provider','?')} / {e.get('payor','?')}: EDI={e.get('edi','?')}, ERA={e.get('era','?')}, EFT={e.get('eft','?')}" for e in edi.get('detail', [])[:10])}

PAYMENTS: {pay.get('count', 0)} payments totaling ${pay.get('total', 0):,.2f}

USER PRODUCTION ANALYSIS ({production_snapshot.get('work_date','')}):
- Total Production Entries Logged: {prod_total_entries}
- Total Production Hours: {prod_total_hours}
- Production by Team Member:
{prod_by_user_lines}
"""

    # Sub-profile data
    if sub_profiles:
        for sp_name, sp_data in sub_profiles.items():
            sc = sp_data.get("claims", {})
            data_summary += f"""
SUB-PROFILE: {sp_name}
  Claims: {sc.get('total', 0)} | Charged: ${sc.get('total_charged', 0):,.2f} | Paid: ${sc.get('total_paid', 0):,.2f} | AR: ${sc.get('total_balance', 0):,.2f}
  Credentialing: {len(sp_data.get('credentialing', {}).get('detail', []))} | Enrollment: {len(sp_data.get('enrollment', {}).get('detail', []))} | EDI: {len(sp_data.get('edi', {}).get('detail', []))}
"""

    system_prompt = """You are a senior Revenue Cycle Management (RCM) analyst at MedPharma SC, a healthcare credentialing and billing company. 
Write a detailed, professional narrative report that a healthcare practice owner can read and understand.

Your report should include:
1. EXECUTIVE SUMMARY — 2-3 sentences on overall account health
2. FINANCIAL PERFORMANCE — Analyze charges, collections, A/R, collection rate with context
3. CLAIMS ANALYSIS — Break down claim statuses, flag concerns, note denials
4. DENIAL MANAGEMENT — If denials exist, explain significance and recommend actions
5. CREDENTIALING STATUS — Summarize progress, flag pending items
6. ENROLLMENT STATUS — Summarize payor enrollment position
7. EDI CONNECTIVITY — Note setup status
8. SUB-PROFILE COMPARISON — If multiple sub-profiles exist, compare performance
9. USER PRODUCTION ANALYSIS — Summarize team productivity and outliers using production logs
10. RECOMMENDED ACTIONS — Specific, prioritized action items
11. OUTLOOK — Brief forward-looking statement

Write in a professional medical billing tone. Use specific numbers from the data.
Do NOT use markdown headers or bullets — write flowing paragraphs separated by blank lines, with key figures in bold (use <b> tags).
Keep it concise but thorough — aim for 400-600 words."""

    try:
        import openai
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Generate the narrative report based on this data:\n\n{data_summary}"}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        narrative = response.choices[0].message.content
        return {"narrative": narrative, "model": "gpt-4o-mini", "company": client_info.get("company", "")}
    except Exception as e:
        logging.getLogger(__name__).exception("AI narrative generation failed")
        raise HTTPException(500, "AI generation failed. Please try again later.")


# ─── PDF Report Generation ───────────────────────────────────────────────────

@router.api_route("/report/{client_id}/pdf", methods=["GET", "POST"])
async def download_report_pdf(client_id: int, period: str = "all", sub_profile: Optional[str] = None,
                              hub_session: Optional[str] = Cookie(None),
                              request: Request = None):
    """Generate and return a branded PDF report."""
    user = _require_reporting_access(hub_session)
    from app.client_db import get_db
    from datetime import date
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from fastapi.responses import StreamingResponse

    # Extract narrative from POST body if available
    narrative = None
    if request and request.method == "POST":
        try:
            body = await request.json()
            narrative = body.get("narrative")
        except Exception:
            pass

    conn = get_db()
    conn.row_factory = sqlite3.Row

    client_row = conn.execute("SELECT company,contact_name,email,phone,practice_type,specialty FROM clients WHERE id=?", (client_id,)).fetchone()
    client_info = dict(client_row) if client_row else {}
    practice_type = client_info.get("practice_type", "") or ""
    company = client_info.get("company", "Client")

    try:
        overall = _build_section_data(conn, client_id, sub_profile=sub_profile, period=period)
        sub_profiles_data = {}
        if practice_type == "MHP+OMT" and not sub_profile:
            for sp_name in ["OMT", "MHP"]:
                sub_profiles_data[sp_name] = _build_section_data(conn, client_id, sub_profile=sp_name, period=period)
    finally:
        conn.close()

    cl = overall.get("claims", {})
    cred = overall.get("credentialing", {})
    enr = overall.get("enrollment", {})
    edi_data = overall.get("edi", {})
    pay = overall.get("payments", {})

    # Build PDF
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.6*inch, rightMargin=0.6*inch)

    styles = getSampleStyleSheet()
    blue = HexColor("#0d47a1")
    light_blue = HexColor("#e3f2fd")
    dark = HexColor("#1a1a2e")
    gray = HexColor("#6b7280")
    green = HexColor("#059669")
    red = HexColor("#dc2626")
    white = HexColor("#ffffff")

    # Custom styles
    styles.add(ParagraphStyle('ReportTitle', parent=styles['Title'], fontSize=22, textColor=blue, spaceAfter=4, alignment=TA_LEFT))
    styles.add(ParagraphStyle('ReportSubtitle', parent=styles['Normal'], fontSize=10, textColor=gray, spaceAfter=16))
    styles.add(ParagraphStyle('SectionHead', parent=styles['Heading2'], fontSize=13, textColor=blue, spaceBefore=18, spaceAfter=8,
                               borderWidth=0, leftIndent=0))
    styles.add(ParagraphStyle('BodyText2', parent=styles['Normal'], fontSize=10, textColor=dark, leading=14, alignment=TA_JUSTIFY, spaceAfter=6))
    styles.add(ParagraphStyle('KPILabel', parent=styles['Normal'], fontSize=8, textColor=gray, alignment=TA_CENTER))
    styles.add(ParagraphStyle('KPIValue', parent=styles['Normal'], fontSize=16, textColor=blue, alignment=TA_CENTER, leading=20))
    styles.add(ParagraphStyle('SmallGray', parent=styles['Normal'], fontSize=8, textColor=gray))
    styles.add(ParagraphStyle('NarrativeText', parent=styles['Normal'], fontSize=10, textColor=dark, leading=15, alignment=TA_JUSTIFY, spaceAfter=8))

    story = []
    period_label = {"all": "All Time", "mtd": "Month to Date", "ytd": "Year to Date"}.get(period, period)

    # ── Header ──
    story.append(Paragraph(f"MedPharma SC", styles['ReportTitle']))
    story.append(Paragraph(f"Revenue Cycle Management & Credentialing Report", styles['ReportSubtitle']))
    story.append(HRFlowable(width="100%", thickness=2, color=blue, spaceAfter=12))
    story.append(Paragraph(f"<b>{company}</b> — {period_label} Report  |  Generated: {business_today().strftime('%B %d, %Y')}", styles['BodyText2']))
    if practice_type:
        story.append(Paragraph(f"Practice Type: {practice_type}  |  Specialty: {client_info.get('specialty', 'N/A')}", styles['SmallGray']))
    story.append(Spacer(1, 12))

    # ── KPI Summary Table ──
    coll_rate = round((cl.get('total_paid', 0) / cl.get('total_charged', 1)) * 100, 1) if cl.get('total_charged') else 0
    kpi_data = [
        [Paragraph('<b>Total Claims</b>', styles['KPILabel']),
         Paragraph('<b>Total Charged</b>', styles['KPILabel']),
         Paragraph('<b>Total Paid</b>', styles['KPILabel']),
         Paragraph('<b>Outstanding A/R</b>', styles['KPILabel']),
         Paragraph('<b>Collection Rate</b>', styles['KPILabel'])],
        [Paragraph(f"<font size='16' color='#0d47a1'><b>{cl.get('total', 0)}</b></font>", styles['KPIValue']),
         Paragraph(f"<font size='16' color='#7c3aed'><b>${cl.get('total_charged', 0):,.0f}</b></font>", styles['KPIValue']),
         Paragraph(f"<font size='16' color='#059669'><b>${cl.get('total_paid', 0):,.0f}</b></font>", styles['KPIValue']),
         Paragraph(f"<font size='16' color='#dc2626'><b>${cl.get('total_balance', 0):,.0f}</b></font>", styles['KPIValue']),
         Paragraph(f"<font size='16' color='#d97706'><b>{coll_rate}%</b></font>", styles['KPIValue'])]
    ]
    kpi_table = Table(kpi_data, colWidths=[doc.width/5]*5)
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), light_blue),
        ('BOX', (0, 0), (-1, -1), 1, blue),
        ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor("#bfdbfe")),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 16))

    # ── AI Narrative (if provided) ──
    if narrative:
        story.append(Paragraph("Overall Account Summary", styles['SectionHead']))
        # Clean HTML tags for reportlab compatibility
        clean = narrative.replace('\n\n', '<br/><br/>').replace('\n', '<br/>')
        clean = clean.replace('<b>', '<b>').replace('</b>', '</b>')
        story.append(Paragraph(clean, styles['NarrativeText']))
        story.append(Spacer(1, 8))

    # ── Claims by Status ──
    by_status = cl.get('by_status', [])
    if by_status:
        story.append(Paragraph("Claims by Status", styles['SectionHead']))
        tdata = [['Status', 'Count', 'Charged', 'Paid']]
        for s in by_status:
            tdata.append([s.get('status', ''), str(s.get('count', 0)),
                         f"${s.get('charged', 0):,.2f}", f"${s.get('paid', 0):,.2f}"])
        tdata.append(['TOTAL', str(cl.get('total', 0)),
                      f"${cl.get('total_charged', 0):,.2f}", f"${cl.get('total_paid', 0):,.2f}"])
        t = Table(tdata, colWidths=[doc.width*0.35, doc.width*0.15, doc.width*0.25, doc.width*0.25])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), blue),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BACKGROUND', (0, -1), (-1, -1), light_blue),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor("#e5e7eb")),
            ('BOX', (0, 0), (-1, -1), 1, blue),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── Credentialing ──
    cred_detail = cred.get('detail', [])
    if cred_detail:
        story.append(Paragraph("Credentialing", styles['SectionHead']))
        tdata = [['Provider', 'Payor', 'Type', 'Status', 'Submitted', 'Approved']]
        for r in cred_detail[:20]:
            tdata.append([r.get('provider', '')[:25], r.get('payor', '')[:25], r.get('type', ''),
                         r.get('status', ''), r.get('submitted', '-'), r.get('approved', '-')])
        cw = [doc.width*0.2, doc.width*0.2, doc.width*0.12, doc.width*0.15, doc.width*0.16, doc.width*0.17]
        t = Table(tdata, colWidths=cw)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), blue), ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor("#e5e7eb")),
            ('BOX', (0, 0), (-1, -1), 1, blue),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── Enrollment ──
    enr_detail = enr.get('detail', [])
    if enr_detail:
        story.append(Paragraph("Enrollment", styles['SectionHead']))
        tdata = [['Provider', 'Payor', 'Type', 'Status', 'Submitted', 'Effective']]
        for r in enr_detail[:20]:
            tdata.append([r.get('provider', '')[:25], r.get('payor', '')[:25], r.get('type', ''),
                         r.get('status', ''), r.get('submitted', '-'), r.get('effective', '-')])
        enr_cw = [doc.width*0.2, doc.width*0.2, doc.width*0.12, doc.width*0.15, doc.width*0.16, doc.width*0.17]
        t = Table(tdata, colWidths=enr_cw)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), blue), ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor("#e5e7eb")),
            ('BOX', (0, 0), (-1, -1), 1, blue),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── EDI Setup ──
    edi_detail = edi_data.get('detail', [])
    if edi_detail:
        story.append(Paragraph("EDI Setup", styles['SectionHead']))
        tdata = [['Provider', 'Payor', 'Payer ID', 'EDI', 'ERA', 'EFT', 'Go-Live']]
        for r in edi_detail[:20]:
            tdata.append([r.get('provider', '')[:20], r.get('payor', '')[:20], r.get('payer_id', ''),
                         r.get('edi', ''), r.get('era', ''), r.get('eft', ''), r.get('go_live', '-')])
        edw = [doc.width*0.16]*7
        t = Table(tdata, colWidths=edw)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), blue), ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor("#e5e7eb")),
            ('BOX', (0, 0), (-1, -1), 1, blue),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── Sub-Profile Comparison ──
    if sub_profiles_data:
        story.append(Paragraph("Sub-Profile Comparison", styles['SectionHead']))
        tdata = [['Sub-Profile', 'Claims', 'Charged', 'Paid', 'A/R', 'Coll. Rate']]
        for sp_name, sp_d in sub_profiles_data.items():
            sc = sp_d.get('claims', {})
            sr = round((sc.get('total_paid', 0) / sc.get('total_charged', 1)) * 100, 1) if sc.get('total_charged') else 0
            tdata.append([sp_name, str(sc.get('total', 0)),
                         f"${sc.get('total_charged', 0):,.2f}", f"${sc.get('total_paid', 0):,.2f}",
                         f"${sc.get('total_balance', 0):,.2f}", f"{sr}%"])
        spw = [doc.width*0.18, doc.width*0.12, doc.width*0.2, doc.width*0.18, doc.width*0.18, doc.width*0.14]
        t = Table(tdata, colWidths=spw)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), blue), ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, HexColor("#e5e7eb")),
            ('BOX', (0, 0), (-1, -1), 1, blue),
            ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── Payments ──
    story.append(Paragraph("Payments Summary", styles['SectionHead']))
    story.append(Paragraph(f"Total Payments: <b>${pay.get('total', 0):,.2f}</b>  |  Payment Count: <b>{pay.get('count', 0)}</b>", styles['BodyText2']))

    # ── Footer ──
    story.append(Spacer(1, 24))
    story.append(HRFlowable(width="100%", thickness=1, color=gray, spaceAfter=8))
    story.append(Paragraph(f"<i>This report was generated by MedPharma SC — Revenue Cycle Management & Credentialing Solutions</i>",
                           styles['SmallGray']))
    story.append(Paragraph(f"<i>Confidential — For internal use only  |  {business_today().strftime('%B %d, %Y')}</i>", styles['SmallGray']))

    doc.build(story)
    buf.seek(0)
    safe_name = company.replace(" ", "_").replace("/", "-")
    filename = f"{safe_name}_Report_{business_today_iso()}.pdf"
    return StreamingResponse(buf, media_type="application/pdf",
                             headers={"Content-Disposition": f"attachment; filename={filename}"})


# ─── Bulk Claim Status Update ─────────────────────────────────────────────────

class BulkStatusIn(BaseModel):
    claim_ids: list[int]
    ClaimStatus: Optional[str] = None
    Owner: Optional[str] = None
    NextAction: Optional[str] = None
    NextActionDueDate: Optional[str] = None


@router.post("/claims/bulk-status")
def bulk_status_update(body: BulkStatusIn, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = _client_scope(user)
    data = {}
    if body.ClaimStatus:
        data["ClaimStatus"] = body.ClaimStatus
    if body.Owner:
        data["Owner"] = body.Owner
    if body.NextAction:
        data["NextAction"] = body.NextAction
    if body.NextActionDueDate:
        data["NextActionDueDate"] = body.NextActionDueDate
    updated = bulk_update_claims(body.claim_ids, data, scope)
    # Audit log
    log_audit(scope, user.get("username", ""), "bulk_status_update",
              "claims", None, f"Updated {updated} claims: {data}")
    notify_bulk_activity(user["username"], "bulk updated", "Claims", updated,
                         f"Status: {body.ClaimStatus or 'N/A'}, Owner: {body.Owner or 'N/A'}")
    return {"ok": True, "updated": updated}


# ─── Global Search ────────────────────────────────────────────────────────────

@router.get("/search")
def search(q: str = "", client_id: Optional[int] = None,
           hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    if not q or len(q) < 2:
        return {"results": []}
    results = global_search(q, scope)
    return {"results": results}


# ─── Alerts & Notifications ──────────────────────────────────────────────────

@router.get("/alerts")
def alerts_endpoint(client_id: Optional[int] = None,
                    hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    # Auto-flag SLA breaches before returning alerts
    auto_flag_sla(scope)
    alert_list = get_alerts(scope)
    return {"alerts": alert_list, "count": len(alert_list)}


# ── In-app notification inbox (chat invites, chat messages, EOD, welcome) ──
# These are written by the same code paths that send emails, so the recipient
# sees them even if SendGrid/SMTP isn't configured.

@router.get("/notifications")
def notifications_list(unread: int = 0, limit: int = 50,
                       hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    items = list_notifications(int(user["id"]),
                               unread_only=bool(int(unread or 0)),
                               limit=max(1, min(int(limit or 50), 200)))
    unread_count = count_unread_notifications(int(user["id"]))
    return {"items": items, "unread": unread_count}


@router.get("/notifications/unread-count")
def notifications_unread_count(hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    return {"unread": count_unread_notifications(int(user["id"]))}


@router.post("/notifications/{nid}/read")
def notifications_mark_read(nid: int,
                            hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    ok = mark_notification_read(int(user["id"]), int(nid))
    return {"ok": ok}


@router.post("/notifications/read-all")
def notifications_mark_all_read(hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    n = mark_all_notifications_read(int(user["id"]))
    return {"ok": True, "marked": n}


@router.delete("/notifications/{nid}")
def notifications_delete(nid: int,
                         hub_session: Optional[str] = Cookie(None)):
    """Permanently remove a single notification for this user."""
    user = _require_user(hub_session)
    ok = delete_notification(int(user["id"]), int(nid))
    return {"ok": ok}


@router.post("/notifications/clear")
def notifications_clear(kind: Optional[str] = None, read_only: int = 0,
                        hub_session: Optional[str] = Cookie(None)):
    """Bulk-delete this user's notifications. Optional ``kind`` filter (e.g.
    'chat_message') and ``read_only=1`` to keep unread ones."""
    user = _require_user(hub_session)
    n = delete_notifications(int(user["id"]), kind=(kind or None),
                             read_only=bool(int(read_only or 0)))
    return {"ok": True, "deleted": n}


# ── EOD report history (persisted reports, viewable in-app) ────────────────

@router.get("/reports/eod/history")
def eod_history(limit: int = 30,
                hub_session: Optional[str] = Cookie(None)):
    """List archived EOD reports. Admin/staff only."""
    _require_admin(hub_session)
    return {"reports": list_eod_reports(limit=max(1, min(int(limit or 30), 365)))}


@router.get("/reports/eod/archive/{report_id}")
def eod_archive_view(report_id: int,
                     hub_session: Optional[str] = Cookie(None)):
    """Fetch a single archived EOD report (admin/staff only)."""
    _require_admin(hub_session)
    rec = get_eod_report(int(report_id))
    if not rec:
        raise HTTPException(404, "EOD report not found")
    return rec


@router.get("/reports/eod/rollup")
def eod_rollup(bucket: str = "day", count: int = 0,
               client_id: Optional[int] = None,
               hub_session: Optional[str] = Cookie(None)):
    """Team activity rolled up per day / per week / per month.

    Aggregates the same work the nightly EOD report summarizes, bucketed across
    time so admins can spot trends. Admin/staff only. `bucket` is one of
    day|week|month. `count` defaults to a sensible number per bucket when 0.
    """
    _require_admin(hub_session)
    b = (bucket or "day").lower().strip()
    if b not in ("day", "week", "month"):
        b = "day"
    if not count or int(count) <= 0:
        count = {"day": 14, "week": 8, "month": 6}[b]
    count = max(1, min(int(count), 366))
    return get_team_activity_rollup(bucket=b, count=count, client_id=client_id)


# ─── Audit Log ────────────────────────────────────────────────────────────────

@router.get("/audit-log")
def audit_log_endpoint(client_id: Optional[int] = None, limit: int = 100,
                       hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    entries = get_audit_log(scope, limit)
    return {"entries": entries}


# ─── Export to CSV / Excel / PDF ─────────────────────────────────────────────

_SECTION_LABELS = {
    "claims": "Claims",
    "credentialing": "Credentialing",
    "enrollment": "Enrollment",
    "edi_setup": "EDI Setup",
    "providers": "Providers",
    "production": "Team Production",
}


def _rows_to_csv_bytes(rows: list[dict], headers: list[str]) -> bytes:
    import csv, io
    out = io.StringIO()
    writer = csv.DictWriter(out, fieldnames=headers, extrasaction="ignore")
    writer.writeheader()
    for r in rows:
        writer.writerow({h: ("" if r.get(h) is None else r.get(h)) for h in headers})
    return out.getvalue().encode("utf-8")


def _rows_to_xlsx_bytes(rows: list[dict], headers: list[str], sheet_title: str) -> bytes:
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    ws = wb.active
    ws.title = (sheet_title or "Sheet1")[:31]
    header_fill = PatternFill("solid", fgColor="0D47A1")
    header_font = Font(bold=True, color="FFFFFF")
    ws.append(headers)
    for c in ws[1]:
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="left", vertical="center")
    for r in rows:
        ws.append([("" if r.get(h) is None else r.get(h)) for h in headers])
    # Best-effort column auto-width
    for i, h in enumerate(headers, start=1):
        col_letter = ws.cell(row=1, column=i).column_letter
        max_len = len(str(h))
        for r in rows[:500]:
            v = r.get(h)
            if v is None:
                continue
            ln = len(str(v))
            if ln > max_len:
                max_len = ln
        ws.column_dimensions[col_letter].width = min(60, max(10, max_len + 2))
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _rows_to_pdf_bytes(rows: list[dict], headers: list[str], title: str) -> bytes:
    from io import BytesIO
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter),
                            topMargin=0.4 * inch, bottomMargin=0.4 * inch,
                            leftMargin=0.4 * inch, rightMargin=0.4 * inch)
    styles = getSampleStyleSheet()
    blue = HexColor("#0d47a1")
    light = HexColor("#e3f2fd")
    grey = HexColor("#6b7280")
    styles.add(ParagraphStyle("ExportTitle", parent=styles["Title"], fontSize=16,
                              textColor=blue, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle("ExportSub", parent=styles["Normal"], fontSize=9,
                              textColor=grey, spaceAfter=12))
    story = [
        Paragraph(title or "Export", styles["ExportTitle"]),
        Paragraph(
            f"Generated {business_now().strftime('%B %d, %Y %I:%M %p')} — "
            f"{len(rows)} row(s)",
            styles["ExportSub"],
        ),
    ]
    if not rows:
        story.append(Paragraph(
            "No records to export for the current selection.",
            styles["Normal"],
        ))
    else:
        # Limit columns to keep the PDF readable
        max_cols = 10
        cols = headers[:max_cols]
        data = [cols]
        for r in rows[:1000]:
            data.append([str(r.get(h, "") if r.get(h) is not None else "")[:80] for h in cols])
        col_w = max(0.7 * inch, (doc.width / max(1, len(cols))))
        table = Table(data, repeatRows=1, colWidths=[col_w] * len(cols))
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
            ("FONTSIZE", (0, 0), (-1, 0), 8),
            ("FONTSIZE", (0, 1), (-1, -1), 7),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), light]),
            ("GRID", (0, 0), (-1, -1), 0.25, HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(table)
        if len(rows) > 1000:
            story.append(Spacer(1, 8))
            story.append(Paragraph(
                f"Showing first 1,000 of {len(rows)} rows — use CSV/Excel for full data.",
                styles["ExportSub"],
            ))
    doc.build(story)
    return buf.getvalue()


def _resolve_export_rows(section: str, scope: Optional[int], sub_profile: Optional[str]) -> list[dict]:
    if section == "claims":
        return export_claims(scope, sub_profile)
    if section in ("credentialing", "enrollment", "edi_setup", "providers"):
        return export_table(section, scope)
    if section == "production":
        return list_production_logs(scope)
    raise HTTPException(400, f"Unknown section: {section}")


_SECTION_FALLBACK_HEADERS = {
    "claims": [
        "id", "client_id", "ClaimKey", "PatientID", "PatientName", "Payor",
        "ProviderName", "NPI", "DOS", "CPTCode", "Description", "ChargeAmount",
        "AllowedAmount", "PaidAmount", "Balance", "ClaimStatus", "Owner",
        "NextAction", "NextActionDueDate", "sub_profile",
    ],
    "credentialing": [
        "id", "client_id", "ProviderName", "Payor", "CredType", "Status",
        "SubmittedDate", "ApprovedDate", "ExpirationDate", "Owner", "Notes",
    ],
    "enrollment": [
        "id", "client_id", "ProviderName", "Payor", "EnrollType", "Status",
        "SubmittedDate", "EffectiveDate", "Owner", "Notes",
    ],
    "edi_setup": [
        "id", "client_id", "ProviderName", "Payor", "PayerID", "EDIStatus",
        "ERAStatus", "EFTStatus", "GoLiveDate", "Owner", "Notes",
    ],
    "providers": [
        "id", "client_id", "ProviderName", "NPI", "Specialty", "Email",
        "Phone", "Notes",
    ],
    "production": [
        "id", "client_id", "work_date", "username", "category",
        "task_description", "quantity", "time_spent", "notes",
    ],
}


@router.get("/export/{section}")
def export_section(section: str, client_id: Optional[int] = None,
                   sub_profile: Optional[str] = None,
                   format: str = "csv",
                   hub_session: Optional[str] = Cookie(None)):
    """Export a section as CSV / Excel / PDF. Empty data returns a file with
    only headers (or an empty PDF) instead of a 404 so the UI buttons always
    succeed."""
    from fastapi.responses import StreamingResponse
    import io

    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    fmt = (format or "csv").lower().strip()
    if fmt not in ("csv", "xlsx", "pdf"):
        raise HTTPException(400, "format must be csv, xlsx, or pdf")

    rows = _resolve_export_rows(section, scope, sub_profile)
    headers = (
        list(rows[0].keys()) if rows
        else _SECTION_FALLBACK_HEADERS.get(section, ["id"])
    )
    label = _SECTION_LABELS.get(section, section.title())
    stamp = business_now().strftime("%Y%m%d_%H%M")
    base_name = f"{section}_{stamp}"

    log_audit(scope, user.get("username", ""), "export",
              section, None, f"Exported {len(rows)} rows as {fmt}")

    if fmt == "csv":
        payload = _rows_to_csv_bytes(rows, headers)
        return StreamingResponse(
            io.BytesIO(payload),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={base_name}.csv"},
        )
    if fmt == "xlsx":
        payload = _rows_to_xlsx_bytes(rows, headers, sheet_title=label)
        return StreamingResponse(
            io.BytesIO(payload),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base_name}.xlsx"},
        )
    # pdf
    payload = _rows_to_pdf_bytes(rows, headers, title=f"{label} Export")
    return StreamingResponse(
        io.BytesIO(payload),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={base_name}.pdf"},
    )


# ─── Dashboard with Date Filters ─────────────────────────────────────────────

@router.get("/dashboard/filtered")
def dashboard_filtered(client_id: Optional[int] = None,
                       period: str = "all",
                       start_date: Optional[str] = None,
                       end_date: Optional[str] = None,
                       sub_profile: Optional[str] = None,
                       hub_session: Optional[str] = Cookie(None)):
    """Dashboard with date filtering support."""
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    # Run SLA auto-flagging on dashboard load
    auto_flag_sla(scope)
    data = get_dashboard(scope, sub_profile=sub_profile,
                         date_from=start_date, date_to=end_date)
    data["user"] = user
    data["alerts"] = get_alerts(scope)
    return data

# ─── Sharefile Links ──────────────────────────────────────────────────────────

class SharefileLinkIn(BaseModel):
    label: str
    url: str
    category: str = "General"


@router.get("/sharefile-links")
def get_sharefile_links(client_id: Optional[int] = None, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id if client_id is not None else _doc_scope(user)
    return {"links": list_sharefile_links(scope)}


@router.post("/sharefile-links")
def create_sharefile_link(
    payload: SharefileLinkIn,
    client_id: Optional[int] = None,
    hub_session: Optional[str] = Cookie(None),
):
    user = _require_user(hub_session)
    if client_id:
        scope = client_id
    elif user.get("role") in ("admin", "staff"):
        scope = _client_scope(user)
    else:
        # Client sub-users add the link to the shared account they belong to.
        scope = _client_upload_account(user)
    if not scope:
        raise HTTPException(400, "Select a client account first, then add the link")
    if not payload.label.strip():
        raise HTTPException(400, "Label is required")
    if not payload.url.strip().startswith(("http://", "https://")):
        raise HTTPException(400, "URL must start with http:// or https://")
    new_id = add_sharefile_link(
        client_id=scope,
        label=payload.label.strip(),
        url=payload.url.strip(),
        category=payload.category.strip() or "General",
        added_by=user.get("username", ""),
    )
    return {"id": new_id, "ok": True}


@router.delete("/sharefile-links/{link_id}")
def remove_sharefile_link(link_id: int, client_id: Optional[int] = None, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    scope = client_id or _client_scope(user)
    delete_sharefile_link(link_id, scope)
    return {"ok": True}


# ─── Chat: Admin-managed rooms with member management ───────────────────────

class ChatMessageIn(BaseModel):
    body: str


class ChatRoomCreate(BaseModel):
    name: str
    description: Optional[str] = ""
    client_id: Optional[int] = None
    member_user_ids: Optional[list[int]] = None


class ChatRoomUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    client_id: Optional[int] = None
    archived: Optional[int] = None


class ChatMemberIn(BaseModel):
    user_id: int
    role: Optional[str] = "member"


class DirectMessageOpen(BaseModel):
    user_id: int


def _is_admin_user(user: dict) -> bool:
    # Only full admins can see/manage every chat room. Staff are regular
    # members — they only see rooms they’ve been explicitly added to. This
    # keeps internal MedPharma rooms private from staff who weren’t invited
    # and stops external client rooms from being visible to unrelated staff.
    return (user or {}).get("role") == "admin"


def _is_internal_user(user: dict) -> bool:
    """Internal MedPharma team members (admin/staff/bizdev) — they can see the
    full team roster to start chats. Clients are NOT internal."""
    return (user or {}).get("role") in ("admin", "staff", "bizdev")


def _require_room_access(user: dict, room_id: int) -> dict:
    room = get_room(room_id)
    if not room:
        raise HTTPException(404, "Chat room not found")
    if not user_can_access_room(room_id, int(user.get("id") or 0), _is_admin_user(user)):
        raise HTTPException(403, "You are not a member of this chat room")
    return room


@router.get("/chat/rooms")
def chat_list_rooms(include_archived: int = 0,
                    hub_session: Optional[str] = Cookie(None)):
    """List chat rooms visible to the current user."""
    user = _require_user(hub_session)
    rooms = list_rooms_for_user(
        int(user["id"]),
        is_admin=_is_admin_user(user),
        include_archived=bool(int(include_archived or 0)),
    )
    return {"rooms": rooms, "is_admin": _is_admin_user(user)}


@router.post("/chat/rooms")
def chat_create_room(body: ChatRoomCreate, request: Request, hub_session: Optional[str] = Cookie(None)):
    """Admin/staff only: create a chat room and seed members.

    Sends each new member a "you were added to chat" email with a
    deep link directly to the room. Returns a per-user delivery report
    so the UI can warn the operator if any emails didn't go out.
    """
    user = _require_chat_manager(hub_session)
    name = (body.name or "").strip()
    if not name:
        raise HTTPException(400, "Room name is required")
    member_ids = body.member_user_ids or []
    room_id = create_room(
        name=name,
        description=(body.description or "").strip(),
        client_id=body.client_id,
        created_by=user.get("username", ""),
        member_user_ids=member_ids,
        creator_user_id=int(user.get("id") or 0),
    )
    log_audit(body.client_id, user.get("username", ""), "chat_room_create",
              "chat_rooms", room_id,
              f"Created room '{name}' with {len(member_ids)} member(s)")
    notify_activity(user.get("username", ""), "created", "Chat Room",
                    f"room #{room_id} '{name}'")
    # In-app notification for every invited member — works even if email
    # is unconfigured.
    try:
        inviter_display = (user.get("contact_name")
                           or user.get("username") or "Your team").strip()
        fanout_notification(
            user_ids=[int(m) for m in member_ids],
            kind="chat_invite",
            title=f"You were added to chat: {name}",
            body=f"{inviter_display} added you to a chat room.",
            link=f"/hub?chat={room_id}",
            related_type="chat_room",
            related_id=room_id,
            skip_user_id=int(user.get("id") or 0),
        )
    except Exception:
        log.exception("chat invite in-app notification failed for room %s",
                      room_id)
    invite_report: list[dict] = []
    try:
        invite_report = _send_chat_invite_emails(
            request=request,
            room_id=room_id,
            room_name=name,
            user_ids=member_ids,
            inviter_name=(user.get("contact_name") or user.get("username") or ""),
            skip_user_id=int(user.get("id") or 0),
        )
    except Exception:
        log.exception("chat invite email failed for room %s", room_id)
    return {"ok": True, "id": room_id, "invites": invite_report}


@router.get("/chat/rooms/{room_id}")
def chat_get_room(room_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    room = _require_room_access(user, room_id)
    return {
        "room": room,
        "members": list_room_members(room_id),
        "is_admin": _is_admin_user(user),
    }


@router.put("/chat/rooms/{room_id}")
def chat_update_room(room_id: int, body: ChatRoomUpdate,
                     hub_session: Optional[str] = Cookie(None)):
    user = _require_admin(hub_session)
    if not get_room(room_id):
        raise HTTPException(404, "Chat room not found")
    fields = {k: v for k, v in body.model_dump().items() if v is not None}
    if not fields:
        return {"ok": True, "updated": 0}
    update_room(room_id, fields)
    log_audit(None, user.get("username", ""), "chat_room_update",
              "chat_rooms", room_id, f"Updated: {','.join(sorted(fields))}")
    return {"ok": True}


@router.delete("/chat/rooms/{room_id}")
def chat_delete_room(room_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    _require_room_access(user, room_id)
    delete_room(room_id)
    log_audit(None, user.get("username", ""), "chat_room_delete",
              "chat_rooms", room_id, "Deleted room")
    notify_activity(user.get("username", ""), "deleted", "Chat Room",
                    f"room #{room_id}")
    return {"ok": True}


@router.get("/chat/rooms/{room_id}/members")
def chat_room_members(room_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    _require_room_access(user, room_id)
    return {"members": list_room_members(room_id)}


@router.post("/chat/rooms/{room_id}/members")
def chat_add_member(room_id: int, body: ChatMemberIn, request: Request,
                    hub_session: Optional[str] = Cookie(None)):
    """Admin/staff only: add a user to a chat room.

    Sends the new member a "you were added to chat" email with a deep
    link. Returns delivery status so the UI can confirm or warn.
    """
    user = _require_chat_manager(hub_session)
    room = get_room(room_id)
    if not room:
        raise HTTPException(404, "Chat room not found")
    add_room_member(room_id, int(body.user_id),
                    role=(body.role or "member"),
                    added_by=user.get("username", ""))
    log_audit(None, user.get("username", ""), "chat_member_add",
              "chat_rooms", room_id, f"Added user #{body.user_id}")
    try:
        inviter_display = (user.get("contact_name")
                           or user.get("username") or "Your team").strip()
        create_notification(
            user_id=int(body.user_id),
            kind="chat_invite",
            title=f"You were added to chat: {room.get('name','') or 'a room'}",
            body=f"{inviter_display} added you to a chat room.",
            link=f"/hub?chat={room_id}",
            related_type="chat_room",
            related_id=room_id,
        )
    except Exception:
        log.exception("chat add-member in-app notification failed for "
                      "room %s user %s", room_id, body.user_id)
    invite_report: list[dict] = []
    try:
        invite_report = _send_chat_invite_emails(
            request=request,
            room_id=room_id,
            room_name=room.get("name", "") or "",
            user_ids=[int(body.user_id)],
            inviter_name=(user.get("contact_name") or user.get("username") or ""),
            skip_user_id=int(user.get("id") or 0),
        )
    except Exception:
        log.exception("chat invite email failed for room %s user %s",
                      room_id, body.user_id)
    invite = invite_report[0] if invite_report else {}
    return {"ok": True, "invite": invite}


@router.delete("/chat/rooms/{room_id}/members/{user_id}")
def chat_remove_member(room_id: int, user_id: int,
                       hub_session: Optional[str] = Cookie(None)):
    """Admin/staff only: remove a user from a chat room."""
    actor = _require_admin(hub_session)
    if not get_room(room_id):
        raise HTTPException(404, "Chat room not found")
    removed = remove_room_member(room_id, int(user_id))
    log_audit(None, actor.get("username", ""), "chat_member_remove",
              "chat_rooms", room_id,
              f"Removed user #{user_id} (existed={removed})")
    return {"ok": True, "removed": removed}


@router.get("/chat/rooms/{room_id}/messages")
def chat_get_messages(room_id: int, limit: int = 200,
                      before_id: Optional[int] = None,
                      mark_read: int = 1,
                      hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    _require_room_access(user, room_id)
    msgs = list_room_messages(
        room_id,
        limit=max(1, min(int(limit or 200), 1000)),
        before_id=before_id,
    )
    last_read = 0
    if mark_read and not before_id:
        last_read = mark_room_read(room_id, int(user["id"]))
    return {
        "room_id": room_id,
        "messages": msgs,
        "last_read_message_id": last_read,
        "reads": list_room_read_state(room_id),
    }


@router.post("/chat/rooms/{room_id}/messages")
def chat_post_message(room_id: int, body: ChatMessageIn,
                      hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    room = _require_room_access(user, room_id)
    text = (body.body or "").strip()
    if not text:
        raise HTTPException(400, "Message body is required")
    if len(text) > 4000:
        raise HTTPException(400, "Message too long (max 4000 characters)")
    try:
        msg_id = add_room_message(
            room_id=room_id,
            sender_id=int(user.get("id") or 0),
            sender_name=user.get("username", ""),
            sender_role=user.get("role", "member"),
            body=text,
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    # HIPAA: never put message body content into audit logs / activity feeds /
    # email notifications. The full body lives encrypted in chat_messages; the
    # only thing we record outside that table is a length-only marker.
    try:
        from app.security import phi_safe_preview
        body_marker = phi_safe_preview(text)
    except Exception:
        body_marker = f"[chat message • {len(text)} chars]"
    log_audit(room.get("client_id"), user.get("username", ""), "chat_message",
              "chat_rooms", room_id, body_marker)
    # In-app notification fanout to every other room member. PHI-safe:
    # we only store the length marker, never the real text.
    try:
        members = list_room_members(room_id) or []
        member_ids = [int(m.get("user_id") or m.get("id") or 0)
                      for m in members]
        member_ids = [uid for uid in member_ids if uid > 0]
        sender_display = (user.get("contact_name")
                          or user.get("username") or "Someone").strip()
        fanout_notification(
            user_ids=member_ids,
            kind="chat_message",
            title=f"New message from {sender_display}",
            body=f"in '{room.get('name','room')}' · {body_marker}",
            link=f"/hub?chat={room_id}",
            related_type="chat_room",
            related_id=room_id,
            skip_user_id=int(user.get("id") or 0),
        )
    except Exception:
        log.exception("chat message in-app notification failed for room %s",
                      room_id)
    try:
        notify_activity(user.get("username", ""), "sent message",
                        "Chat", f"room '{room.get('name','?')}' {body_marker}")
        # NOTE: we intentionally DO NOT email on every message anymore. That
        # buried real signal in noise. Instead, a scheduled job
        # (send_chat_unread_reminders) emails a one-time nudge only to people
        # who were @mentioned and still haven't read the message after 2 hours.
        # Read messages never trigger an email. The in-app unread badge above
        # handles the live "you have a new message" indicator.
    except Exception:
        log.exception("chat notify failed")
    return {"ok": True, "id": msg_id}


@router.post("/chat/rooms/{room_id}/mark-read")
def chat_mark_read_endpoint(room_id: int, hub_session: Optional[str] = Cookie(None)):
    user = _require_user(hub_session)
    _require_room_access(user, room_id)
    last_id = mark_room_read(room_id, int(user["id"]))
    return {"ok": True, "last_read_message_id": last_id}


@router.get("/chat/unread-count")
def chat_unread_count_endpoint(hub_session: Optional[str] = Cookie(None)):
    """Total unread messages for the current user across all visible rooms."""
    user = _require_user(hub_session)
    n = chat_unread_total(int(user["id"]), is_admin=_is_admin_user(user))
    return {"unread": n}


@router.get("/chat/users")
def chat_eligible_users(hub_session: Optional[str] = Cookie(None)):
    """List users that can be added to a chat room.

    Admin/staff see everyone. Clients see only the staff/admin users that
    have been granted access to their own account (so they can DM their team).
    """
    user = _require_user(hub_session)
    if _is_internal_user(user):
        return {"users": list_chat_eligible_users()}
    # Client view: limit to staff/admin assigned to this client.
    uid = int(user.get("id") or 0)
    assigned = list_client_access(uid)
    # Filter to active staff/admin only.
    out = [
        u for u in assigned
        if (u.get("role") or "").lower() in ("admin", "staff")
    ]
    return {"users": out}


@router.post("/chat/dm")
def chat_open_dm(body: DirectMessageOpen, hub_session: Optional[str] = Cookie(None)):
    """Open (or create) a private 1:1 direct-message room with another user.

    Returns the DM room id. Idempotent — the same pair of users always
    resolves to the same room regardless of who opens it first."""
    user = _require_user(hub_session)
    try:
        other_id = int(body.user_id or 0)
    except (TypeError, ValueError):
        raise HTTPException(400, "user_id is required")
    if other_id <= 0:
        raise HTTPException(400, "user_id is required")
    me_id = int(user.get("id") or 0)
    if other_id == me_id:
        raise HTTPException(400, "You cannot message yourself")
    # The other user must be someone this user is allowed to chat with.
    eligible_ids = {int(u["id"]) for u in (chat_eligible_users(hub_session) or {}).get("users", [])}
    if other_id not in eligible_ids:
        raise HTTPException(403, "You cannot start a direct message with this user")
    room_id = get_or_create_dm_room(me_id, other_id, created_by=user.get("username", ""))
    return {"ok": True, "id": room_id}


# ─── Client Seed Backup ───────────────────────────────────────────────────────

@router.get("/clients/export-seed")
def export_client_seed(hub_session: Optional[str] = Cookie(None)):
    """Admin: download clients_seed.json — commit this to repo for persistence."""
    user = _require_user(hub_session)
    if user.get("role") != "admin":
        raise HTTPException(403, "Admin only")
    from fastapi.responses import JSONResponse
    seed = _load_clients_seed()
    return JSONResponse(
        content=seed,
        headers={"Content-Disposition": "attachment; filename=clients_seed.json"},
    )


# ─── Team Tracking / Productivity (ActivTrak-style) ──────────────────────────

@router.post("/track/heartbeat")
def track_heartbeat(request: Request, hub_session: Optional[str] = Cookie(None)):
    """Frontend pings this every ~60s while the tab is focused. Updates the
    user's last-seen timestamp and contributes to their daily active time."""
    user = _require_user(hub_session)
    log_activity(
        user["username"], "heartbeat",
        client_id=user.get("id"),
        ip=(request.client.host if request.client else ""),
        user_agent=request.headers.get("user-agent", ""),
    )
    return {"ok": True, "ts": datetime.now().isoformat(timespec="seconds")}


@router.get("/track/activity")
def track_activity(
    username: Optional[str] = None,
    start: Optional[str] = None,
    end: Optional[str] = None,
    event_type: Optional[str] = None,
    limit: int = 500,
    hub_session: Optional[str] = Cookie(None),
):
    """List timestamped activity events. Non-admin users can only see their own."""
    user = _require_user(hub_session)
    if user.get("role") not in ("admin", "staff"):
        username = user["username"]
    rows = list_activity_events(
        username=username, start=start, end=end,
        event_type=event_type, limit=min(int(limit or 500), 5000),
    )
    return {"ok": True, "count": len(rows), "events": rows}


@router.get("/track/live")
def track_live(within: int = 300, hub_session: Optional[str] = Cookie(None)):
    """Who's online right now (seen in the last `within` seconds)."""
    _require_user(hub_session)
    users = get_live_users(within_seconds=max(30, min(int(within or 300), 3600)))
    return {"ok": True, "within_seconds": within, "users": users}


@router.get("/track/productivity")
def track_productivity(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    username: Optional[str] = None,
    hub_session: Optional[str] = Cookie(None),
):
    """Per-user-per-day productivity rollup (ActivTrak-style)."""
    user = _require_user(hub_session)
    if user.get("role") not in ("admin", "staff"):
        username = user["username"]
    if not start_date:
        start_date = (business_now() - timedelta(days=14)).strftime("%Y-%m-%d")
    if not end_date:
        end_date = business_today_iso()
    return get_productivity_report(start_date=start_date, end_date=end_date, username=username)
